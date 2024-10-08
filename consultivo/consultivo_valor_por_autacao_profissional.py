import pandas as pd
import streamlit as st
from st_pages import add_indentation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import locale
import time
from num2words import num2words
import os
from tempfile import NamedTemporaryFile
from utils.funcoes import format_paragraph, add_formatted_text, format_title_centered, \
    format_title_justified, num_extenso, data_extenso, fonte_name_and_size, add_section,\
    num_extenso_percentual, set_table_borders, obter_texto_parcelas


try:
    locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
except locale.Error as e:
    print(f"Erro ao definir a localidade: {e}")


add_indentation() 


# Store the initial value of widgets in session state
if "visibility" not in st.session_state:
    st.session_state.visibility = "visible"
    st.session_state.disabled = False
    st.session_state.horizontal = False

#####################################################################################

lista_numerada = ['a)', 'b)', 'c)', 'd)', 'e)', 'f)', 'g)', 'h)', 'i)', 'j)', 'k)', 'l)', 'm)', 'n)', 'o)', 'p)', 'q)', 'r)', 's)', 't)']

#####################################################################################

recuo = "&nbsp;" * 24

# Inicializar DataFrame vazio
df_inputs = pd.DataFrame(columns=[
    'objeto', 
    'total-de-horas', 
    'valor-aplicado', 
    'valor-formatado',
    'valor_por_extenso',
    'subtotal'
    'Subtotal-extenso'
    ])


dados, desenvolvimento = st.columns([2,3])
#Dicionário das informações
perguntas_respostas = {}

with dados:

    ###########Streamlit
    st.write('**Informação para a proposta**')
    # Carregando a lista de clientes
    lista_clientes = pd.read_csv('clientes.csv')
    lista_clientes = lista_clientes.sort_values(by='Nome')['Nome'].unique().tolist()

    # Adiciona uma opção para cadastrar novo cliente
    lista_clientes.append("--Novo cliente--")
    lista_clientes = sorted(lista_clientes)

    nome_cliente = st.selectbox(
            'Cliente',
            (lista_clientes),
            index=None,
            placeholder='Selecione o cliente'
        )   

    if nome_cliente == '--Novo cliente--':
        form = st.form('Novo Cliente')
        nome_cliente = form.text_input('Cadastrar novo cliente')
        form.form_submit_button("Cadastrar")
    
    st.divider()

########### Python-docx
    # Abrir documento com papel timbrado da RKP
    document = Document(r"docx/RKP-PapelTimbrado.docx")

    # Definir fonte e tamanho do documento
    fonte_name_and_size(document, 'Arial', 12)


    # Adicionar uma sessão ao documento
    section = add_section(document, 4,2.5,2,3)


    # Data
    data = datetime.now()
    paragraph_date = document.add_paragraph(f'Brasília/DF, {data_extenso(data)}')
    paragraph_format = paragraph_date.paragraph_format
    paragraph_format.alignment = 2  # a direita

    
    #############################################################################
    # Adicionar título ao doc
    title = document.add_heading(level=1)
    title.add_run('PROPOSTA PARA PRESTAÇÃO \nDE SERVIÇOS ADVOCATÍCIOS').bold=True
    format_title_centered(title)
    space_format = title.paragraph_format
    space_format.space_before = Pt(48)


    p_de = document.add_paragraph()
    p_de.add_run('DE: ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS S/S').bold = True
    p_de_format = p_de.paragraph_format
    p_de_format.line_spacing = Pt(18)
    p_de_format.space_before = Pt(148)
    p_de_format.space_after = Pt(8)

    paragraph_para = document.add_paragraph()
    paragraph_para.add_run(f'PARA: {nome_cliente} - Interessado(a)').bold = True
    paragraph_format = paragraph_para.paragraph_format
    paragraph_format.line_spacing = Pt(18)
    paragraph_format.space_after = Pt(48)

    paragraph_ref = document.add_paragraph()
    paragraph_ref.add_run('Referência: PROPOSTA DE HONORÁRIOS ADVOCATÍCIOS').bold=True
    paragraph_format = paragraph_ref.paragraph_format
    paragraph_format.space_before = Pt(18)
    paragraph_format.space_after = Pt(96)
    paragraph_format.line_spacing = Pt(18)


    paragraph = document.add_paragraph()
    format_paragraph(paragraph, 3, 1.4764,0,48,16,20)
    full_text= "ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS S/S, com sede no SIG - Quadra 01, Lote 495, Edifício Barão do Rio Branco, sala 244, Brasília-DF, CEP 70.610-410, telefones 3321-7043 e 3226-0137, inscrita no CNPJ sob o nº 03.899.920/0001- 81, registro na Ordem dos Advogados do Brasil – OAB/DF sob o número 616/00 – RS, endereço eletrônico www.khouriadvocacia.com.br, vem, mui respeitosamente, apresentar PROPOSTA DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS, nas condições a seguir."
    # Texto que será negritado
    bold_text = "ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS S/S"
    # Adicionar o texto formatado ao parágrafo
    add_formatted_text(paragraph, full_text, bold_text)

    # I - DOS SERVIÇOS A SEREM DESENVOLVIDOS
###########Streamlit
    
    # objeto da proposta
    input_objeto = st.text_area(label="Objeto(s) da proposta", placeholder='Conforme solicitação, apresentamos proposta de honorários prestação de serviços advocatícios específicos de consultoria jurídica cível, que consistirá...')
    resumo_objeto = st.text_area(label="Resumo do(s) objeto(s) (ENTER para quebra de linha)")
    
        
    perguntas_respostas = {
        'nome_cliente': nome_cliente,
        '[objeto_texto]': input_objeto,
        '[resumo_objeto]': resumo_objeto,}



########### Python-docx
    title_one = document.add_heading('I - DOS SERVIÇOS A SEREM DESENVOLVIDOS', level=2)
    format_title_justified(title_one)

    
    # Objeto da proposta
    desdobramentos = perguntas_respostas['[objeto_texto]'].split("\n")
    textos_paragrafos = []
    texto_padrao = []
    if len(desdobramentos) > 1:
        for p in desdobramentos:
            paragrafo_ = document.add_paragraph(f'Conforme solicitação, apresentamos proposta de honorários prestação de serviços advocatícios específicos de consultoria jurídica cível, que consistirão {p}.')
            format_paragraph(paragrafo_, 3, 1.5748,0, 18,18,18)
            textos_paragrafos.append(paragrafo_.text)
    else:
        paragrah_padrao = document.add_paragraph(f"Conforme solicitação, apresentamos proposta de honorários prestação de serviços advocatícios específicos de consultoria jurídica cível, que consistirá {desdobramentos[0]}.")
        format_paragraph(paragrah_padrao,3, 1.5748,0, 18,18,18)

    paragraph_atividades = document.add_paragraph('A atuação desse Jurídico compreenderá as seguintes atividades:')
    format_paragraph(paragraph_atividades, 3, 1.4764,0,18,18,18)
    
        
    #atuação
    itens_atuacao = [
        "Providências preliminares de levantamento e análise de todas as informações e documentos relativos ao objeto da presente proposta, a fim de propiciar o embasamento jurídico necessário",
        "Participações em reuniões e eventuais discussões a respeito do contrato, incluindo em entendimentos entre as partes, caso seja necessário",
    ]

    atuacao = perguntas_respostas['[resumo_objeto]'].split("\n")
    if len(atuacao)> 1:
        for atuar in atuacao:
            itens_atuacao.append(atuar)
    else:
        itens_atuacao.append(resumo_objeto)

    # Cria uma nova lista com a mesma quantidade de elementos de servicos, preenchida com itens de lista_numerada
    lista_numerada_servicos = lista_numerada[:len(itens_atuacao)]
    
    # for i in range(len(itens_atuacao)):
    #     paragraph_itens_atuacao = document.add_paragraph(f'{lista_numerada[i]} {itens_atuacao[i]};')
    #     format_paragraph(paragraph_itens_atuacao, 3, 0,1.77165, 18, 18, 18)
    for i in range(len(itens_atuacao)):
        # Adiciona o ponto e vírgula por padrão
        punct = ";"
        
        # Se for o último item, muda o ponto final
        if i == len(itens_atuacao) - 1:
            punct = "."
        
        # Adiciona o parágrafo com a pontuação adequada
        paragraph_itens_atuacao = document.add_paragraph(f'{lista_numerada[i]} {itens_atuacao[i]}{punct}')
    
        # Aplica o formato ao parágrafo
        format_paragraph(paragraph_itens_atuacao, 3, 0, 1.77165, 18, 18, 18)

        
    format_paragraph(paragraph_itens_atuacao, 3, 0, 1.77165, 18, 18, 18)



    #paragrafo I-IV
    paragraph_four = document.add_paragraph()
    format_paragraph(paragraph_four,3, 1.5748,0, 18,18,18)
    paragraph_four.text ='Para o cumprimento dos serviços, o escritório disponibilizará sua equipe técnica, sendo que haverá advogado responsável pelo acompanhamento direto da demanda.'

    #paragrafo I-V
    paragraph_five = document.add_paragraph()
    format_paragraph(paragraph_five,3, 1.5748, 0, 18,18,18)
    paragraph_five.text ='A Roque Khouri & Pinheiro Advogados Associados alerta que qualquer ação judicial implica em risco, estando o interessado ciente principalmente da possibilidade de condenação em honorários advocatícios, conforme previsto no Código de Processo Civil (10% a 20% sobre o valor da causa atualizado), principalmente diante das especificidades da fase atual em que o processo se encontra. Alerta também se tratar de análise e de confecção de peças a serem elaboradas com base no direito, jurisprudência atual e principalmente das informações e documentos que serão sempre fornecidos pela Interessada.'


    #############################################################################
    # II - DA POLÍTICA GERAL DE VALORES - HONORÁRIOS
    #titulo II
    title_two = document.add_heading('II - DA POLÍTICA GERAL DE VALORES - HONORÁRIOS', level=2)
    format_title_justified(title_two)
    #Paragrafo II-I
    paragraph_two_one = document.add_paragraph()
    format_paragraph(paragraph_two_one,3, 1.5748, 0,18,18,18)
    paragraph_two_one.text = "Faz parte integrante de todas as nossas propostas de honorários os itens abaixo, componentes da nossa Política de Honorários de consultoria:" #\nTaxas horárias de honorários para projetos. Para projetos, nós cobramos valores de honorários de acordo com as seguintes taxas horárias:

    #Paragrafo II-II
    paragraph_two_two = document.add_paragraph()
    format_paragraph(paragraph_two_two,3, 1.5748, 0, 18,18,18)
    full_text= "Taxas horárias de honorários para projetos. Para projetos, nós cobramos valores de honorários de acordo com as seguintes taxas horárias:"
    # Texto que será negritado
    bold_text = "Taxas horárias de honorários para projetos."
    # Adicionar o texto formatado ao parágrafo
    add_formatted_text(paragraph_two_two, full_text, bold_text)

    # Adicionar tabela
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style =  None 
    # Definir bordas da tabela
    set_table_borders(table)
    # get table data -------------
    items = (
        ('Sócio Majoritário - Dr. Paulo Roque', 'R$1.150,00'),
        ('Sócia Nominal - Dra. Ângela Pinheiro', 'R$850,00'),
        ('Advogado Sênior', 'R$680,00'),
        ('Advogado Pleno', 'R$580,00'),
        ('Advogado Júnior', 'R$490,00'),
        ('Paralegal/Estagiário', 'R$290,00')
    )

    # populate header row --------
    heading_cells = table.rows[0].cells
    heading_cells[0].text = 'Profissional'
    heading_cells[1].text = 'Valor'


    # Alinhar texto na primeira linha ao centro verticalmente
    for cell in heading_cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Definir cor de fundo para a primeira linha
    shading_elm_0 = OxmlElement('w:shd')
    shading_elm_0.set(qn('w:fill'), 'B4FF00')  # RGB (180, 255, 0)
    heading_cells[0]._element.get_or_add_tcPr().append(shading_elm_0)
    shading_elm_1 = OxmlElement('w:shd')
    shading_elm_1.set(qn('w:fill'), 'B4FF00')  # RGB (180, 255, 0)
    heading_cells[1]._element.get_or_add_tcPr().append(shading_elm_1)


    # add a data row for each item
    for item in items:
        cells = table.add_row().cells
        cells[0].text = item[0]
        cells[1].text = item[1]
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Alinhar texto à direita na segunda coluna

    # Definir bordas da tabela
    set_table_borders(table)


    #Paragrafo II-III
    paragraph_two_three = document.add_paragraph()
    format_paragraph(paragraph_two_three, 3, 1.5748,0, 18,18,18)
    full_text= "Reembolso de Despesas. As despesas incorridas no desenvolvimento dos trabalhos, como, por exemplo, despesas com ligações telefônicas, correios, couriers e outros meios de envio de documentos, com impressão de cópias e digitalização de documentos, com taxas governamentais, com viagens, táxis e outros deslocamentos, e, se aplicável, despesas com custas processuais e outras despesas relativas a processos arbitrais, judiciais e administrativos, e honorários de advogados correspondentes, serão reembolsadas, mediante a apresentação de planilha discriminada, e, se solicitado, dos respectivos comprovantes. Nenhuma despesa superior a R$ 1.000,00 (um mil Reais) será incorrida sem sua prévia aprovação por escrito."
    # Texto que será negritado
    bold_text = "Reembolso de Despesas."
    # Adicionar o texto formatado ao parágrafo
    add_formatted_text(paragraph_two_three, full_text, bold_text)


    #Paragrafo II-IV
    paragraph_two_four = document.add_paragraph()
    format_paragraph(paragraph_two_four, 3, 1.5748,0, 18,18,18)
    full_text= "Interrupção ou Suspensão dos Trabalhos. Se por qualquer motivo os trabalhos forem eventualmente interrompidos ou suspensos, faremos o levantamento das horas trabalhadas e o valor de honorários pagos até então e, se houver saldo a ser pago, faremos o faturamento correspondente. Caso haja honorários a serem restituídos, a restituição será feita no mês seguinte ao da interrupção ou suspensão dos trabalhos e do valor a ser restituído serão descontados os tributos correspondentes pagos ou a pagar."
    # Texto que será negritado
    bold_text = "Interrupção ou Suspensão dos Trabalhos."
    # Adicionar o texto formatado ao parágrafo
    add_formatted_text(paragraph_two_four, full_text, bold_text)

    #############################################################################

    # III - DOS HONORÁRIOS ESPECÍFICOS
    #titulo III
    title_three = document.add_heading('III - DOS HONORÁRIOS ESPECÍFICOS', level=2)
    format_title_justified(title_three)

    #paragrafo III-I
    paragraph_three_one = document.add_paragraph('Com o intuito de manter a proporcionalidade entre prestação de serviços e pagamento, os honorários advocatícios devidos em consequência da presente prestação de serviços seriam cobrados por meio do sistema de horas, ou seja, cada ato praticado, esse Jurídico seria remunerado de acordo com o tempo necessário para praticá-lo.')
    format_paragraph(paragraph_three_one, 3, 1.5748, 0, 18,18,18)

    #paragrafo III-II
    paragraph_three_two = document.add_paragraph('Para a prestação de serviços advocatícios listada no Tópico I, a Roque Khouri & Pinheiro estima os seguintes valores: ')
    format_paragraph(paragraph_three_two, 3, 1.5748,0, 18,18,18)

    ###########Streamlit
    #tipo de cobrança por atuação x por profissional
    tipo_cobranca = st.selectbox('Tipo de cobrança de honorário', ['Atuação', 'Profissional'], index=None)
    st.divider()
    soma = 0
    #split do resumo do objeto
    atuacao = perguntas_respostas['[resumo_objeto]'].split("\n")
    textos_paragrafos = []
    texto_padrao = []
    hora_total_objeto = []
    valor_total_formatado = ''
    valor_total_proposta = 0
    if tipo_cobranca == 'Atuação':

        if len(atuacao) > 1:
            st.markdown('**Preencher para cada objeto**')
            for p in atuacao:
                st.write(f'**{p}**')
                hora_total = st.number_input(label='Total de horas:', step=10, key=f'hora_total_{p}')
                hora_total = int(hora_total)
                hora_total_objeto.append(hora_total)
                valor_aplicado = st.selectbox("Valor aplicado",
                                        (1150.00, 850.00, 680.00, 580.00, 490.00, 290.00), key=f'valor_aplicado_{p}')
                # Arredondar o valor para duas casas decimais e formatar como string
                valor_formatado = "{:.2f}".format(round(valor_aplicado, 2))
                valor_por_extenso = num_extenso(valor_formatado)
                subtotal = (hora_total * valor_aplicado)
                subtotal_formatado = "{:.2f}".format(round(subtotal, 2))
                subtotal_extenso = num_extenso(subtotal_formatado)
                st.write(f'*Subtotal R${subtotal}*')
                valor_total_proposta +=subtotal
                df_inputs = df_inputs.append(
                {
                    'objeto': p,
                    'total-de-horas': hora_total,
                    'valor-aplicado': valor_aplicado,
                    'valor_por_extenso': valor_por_extenso,
                    'valor-formatado': valor_por_extenso,
                    'subtotal': subtotal,
                    'subtotal-extenso': subtotal_extenso
                }, ignore_index = True
            )
            
                ###########Python-docx
                block_three_atuacao = document.add_paragraph(f"{p}")
                format_paragraph(block_three_atuacao, 3, 1.5748, 0, 18,18,18)
                block_three_hora = document.add_paragraph(f'{hora_total}h estimada para a confecção e revisão')
                format_paragraph(block_three_hora, 3, 1.5748,0, 18,18,18)
                block_three_valor_aplicado = document.add_paragraph(f"Valor da hora aplicada: R${valor_aplicado} ({num_extenso(valor_formatado)})")
                format_paragraph(block_three_valor_aplicado, 3, 1.5748,0, 18,18,18)
                block_three_subtotal = document.add_paragraph(f'R${subtotal_formatado} ({subtotal_extenso}) estimada para a confecção e revisão')
                format_paragraph(block_three_subtotal, 3, 1.5748, 0,18,18,18)
            
            # valor_total_formatado = "{:.2f}".format(round(valor_total_proposta, 2))
            # block_three_valor_final = document.add_paragraph(f'Valor total: R${valor_total_proposta} {num_extenso(valor_total_formatado)}')
            # format_paragraph(block_three_valor_final, 3, 1.5748, 0, 18,18,18)

        else:
            hora_total = st.number_input(label='Total de horas:', step=10, key='hora_total')
            hora_total = int(hora_total)
            hora_total_objeto.append(hora_total)
            valor_aplicado = st.selectbox("Valor aplicado",
                                    (1150.00, 850.00, 680.00, 580.00, 490.00, 290.00), key='valor_aplicado')
            # Arredondar o valor para duas casas decimais e formatar como string
            valor_formatado = "{:.2f}".format(round(valor_aplicado, 2))
            valor_por_extenso = num_extenso(valor_formatado)
            valor_total = (hora_total * valor_aplicado)
            valor_total_formatado = "{:.2f}".format(round(valor_total, 2))
            subtotal_extenso = num_extenso(valor_total_formatado)
            st.write(f'**Total R${valor_total}**')
            valor_total_proposta +=valor_total
            df_inputs = df_inputs.append(
                {
                    'objeto': atuacao,
                    'total-de-horas': hora_total,
                    'valor-aplicado': valor_aplicado,
                    'valor_por_extenso': valor_por_extenso,
                    'valor-formatado': valor_por_extenso,
                    'subtotal': valor_total,
                    'subtotal-extenso': subtotal_extenso
                }, ignore_index = True
            )

            block_three_atuacao = document.add_paragraph(atuacao)
            format_paragraph(block_three_atuacao, 3, 1.5748, 0,18,18,0)
            block_three_hora = document.add_paragraph(f'{hora_total}h estimada para a confecção e revisão')
            format_paragraph(block_three_hora, 3, 1.5748, 0,18,18,0)
            block_three_valor_aplicado = document.add_paragraph(f"Valor da hora aplicada: R${valor_total_formatado} ({num_extenso(valor_total_formatado)})")
            format_paragraph(block_three_valor_aplicado, 3, 1.5748,0, 18,18,0)
            block_three_subtotal = document.add_paragraph(f'R${valor_total_formatado} ({subtotal_extenso}) estimada para a confecção e revisão')
            format_paragraph(block_three_subtotal, 3, 1.5748, 0,18,18,0)

# ####Streamlit
        #valor da proposta proposto
        valor_total_proposta_formatado = "{:.2f}".format(round(valor_total_proposta, 2))
        st.write(f'**Valor da proposta: R${valor_total_proposta_formatado}**')
        
    elif tipo_cobranca == 'Profissional':
        # atuacao = perguntas_respostas['[resumo_objeto]'].split("\n")
        # textos_paragrafos = []
        # texto_padrao = []
        # hora_total_objeto = []
        # valor_total_proposta = 0
        qde_profissionais = st.number_input(label='Quantidade', min_value=1, key='qde_profissionais')
        st.divider()
        
        for q in range(qde_profissionais):
            st.write(f'**Profissional {q+1}**')
            hora_total_adv = st.number_input(label='Total de horas', step=10, key=f'hora_total_{q}')
            hora_total = int(hora_total_adv)
            hora_total_objeto.append(hora_total)
            valor_aplicado_adv = st.selectbox("Valor aplicado",
                                    (1150.00, 850.00, 680.00, 580.00, 490.00, 290.00), key=f'valor_aplicado_adv_{q}')
            # Arredondar o valor para duas casas decimais e formatar como string
            valor_formatado = "{:.2f}".format(round(valor_aplicado_adv, 2))
            valor_por_extenso = num_extenso(valor_formatado)
            valor_total = (hora_total * valor_aplicado_adv)
            valor_total_formatado = "{:.2f}".format(round(valor_total, 2))
            subtotal_extenso = num_extenso(valor_total_formatado)
            st.write(f'Subtotal R${valor_total_formatado}')
            st.divider()
            valor_total_proposta += valor_total

            ###########Python-docx
            if valor_aplicado_adv == 1150.00:
                block_three_atuacao = document.add_paragraph("Sócio majoritário")
                format_paragraph(block_three_atuacao, 3, 1.5748, 0, 18,0,0)
                block_three_hora = document.add_paragraph(f'{hora_total}h estimada para a confecção e revisão')
                format_paragraph(block_three_hora, 3, 1.5748,0, 18,18,18)
                block_three_valor_aplicado = document.add_paragraph(f"Valor da hora aplicada: R${valor_formatado} ({num_extenso(valor_formatado)})")
                format_paragraph(block_three_valor_aplicado, 3, 1.5748,0, 18,0,0)
                block_three_subtotal = document.add_paragraph(f'R${valor_total_formatado} ({num_extenso(valor_total_formatado)}) estimada para a confecção e revisão')
                format_paragraph(block_three_subtotal, 3, 1.5748, 0,18,0,18)
            elif valor_aplicado_adv == 850.00:
                block_three_atuacao = document.add_paragraph("Sócia Nominal")
                format_paragraph(block_three_atuacao, 3, 1.5748, 0, 18,0,0)
                block_three_hora = document.add_paragraph(f'{hora_total}h estimada para a confecção e revisão')
                format_paragraph(block_three_hora, 3, 1.5748,0, 18,0,0)
                block_three_valor_aplicado = document.add_paragraph(f"Valor da hora aplicada: R${valor_formatado} ({num_extenso(valor_formatado)})")
                format_paragraph(block_three_valor_aplicado, 3, 1.5748,0, 18,0,0)
                block_three_subtotal = document.add_paragraph(f'R${valor_total_formatado} ({num_extenso(valor_total_formatado)}) estimada para a confecção e revisão')
                format_paragraph(block_three_subtotal, 3, 1.5748, 0,18,0,0)
            elif valor_aplicado_adv == 680.00:
                block_three_atuacao = document.add_paragraph("Advogado sênior")
                format_paragraph(block_three_atuacao, 3, 1.5748, 0, 18,0,18)
                block_three_hora = document.add_paragraph(f'{hora_total}h estimada para a confecção e revisão')
                format_paragraph(block_three_hora, 3, 1.5748,0, 18,0,0)
                block_three_valor_aplicado = document.add_paragraph(f"Valor da hora aplicada: R${valor_formatado} ({num_extenso(valor_formatado)})")
                format_paragraph(block_three_valor_aplicado, 3, 1.5748,0, 18,0,0)
                block_three_subtotal = document.add_paragraph(f'R${valor_total_formatado} ({num_extenso(valor_total_formatado)}) estimada para a confecção e revisão')
                format_paragraph(block_three_subtotal, 3, 1.5748, 0,18,0,0)
            elif valor_aplicado_adv == 580.00:
                block_three_atuacao = document.add_paragraph("Advogado pleno")
                format_paragraph(block_three_atuacao, 3, 1.5748, 0, 18,0,0)
                block_three_hora = document.add_paragraph(f'{hora_total}h estimada para a confecção e revisão')
                format_paragraph(block_three_hora, 3, 1.5748,0, 18,0,0)
                block_three_valor_aplicado = document.add_paragraph(f"Valor da hora aplicada: R${valor_formatado} ({num_extenso(valor_formatado)})")
                format_paragraph(block_three_valor_aplicado, 3, 1.5748,0, 18,0,0)
                block_three_subtotal = document.add_paragraph(f'R${valor_total_formatado} ({num_extenso(valor_total_formatado)}) estimada para a confecção e revisão')
                format_paragraph(block_three_subtotal, 3, 1.5748, 0,18,0,0)
            elif valor_aplicado_adv == 490.00:
                block_three_atuacao = document.add_paragraph("Advogado júnior")
                format_paragraph(block_three_atuacao, 3, 1.5748, 0, 18,0,0)
                block_three_hora = document.add_paragraph(f'{hora_total}h estimada para a confecção e revisão')
                format_paragraph(block_three_hora, 3, 1.5748,0, 18,0,0)
                block_three_valor_aplicado = document.add_paragraph(f"Valor da hora aplicada: R${valor_formatado} ({num_extenso(valor_formatado)})")
                format_paragraph(block_three_valor_aplicado, 3, 1.5748,0, 18,0,0)
                block_three_subtotal = document.add_paragraph(f'R${valor_total_formatado} ({num_extenso(valor_total_formatado)}) estimada para a confecção e revisão')
                format_paragraph(block_three_subtotal, 3, 1.5748, 0,18,0,0)
            else:
                block_three_atuacao = document.add_paragraph("Paralegal/estagiário")
                format_paragraph(block_three_atuacao, 3, 1.5748, 0, 18,18,0)
                block_three_hora = document.add_paragraph(f'{hora_total}h estimada para a confecção e revisão')
                format_paragraph(block_three_hora, 3, 1.5748,0, 18,18,0)
                block_three_valor_aplicado = document.add_paragraph(f"Valor da hora aplicada: R${valor_formatado} ({num_extenso(valor_formatado)})")
                format_paragraph(block_three_valor_aplicado, 3, 1.5748,0, 18,0,0)
                block_three_subtotal = document.add_paragraph(f'R${valor_total_formatado} ({num_extenso(valor_total_formatado)}) estimada para a confecção e revisão')
                format_paragraph(block_three_subtotal, 3, 1.5748, 0,18,18,0)

        valor_total_formatado = "{:.2f}".format(round(valor_total_proposta, 2))
        block_three_valor_final = document.add_paragraph(f'Valor total: R${valor_total_formatado} ({num_extenso(valor_total_formatado)})')
        format_paragraph(block_three_valor_final, 3, 1.5748, 0, 18,18,18)
            
        # for profissional in qde_profissionais:
        st.write(f'**Total da proposta: R$ {valor_total_formatado}**')

#     else:
#         while True:
#             if tipo_cobranca:
#                 break
#             time.sleep(2)        

    #desconto
    st.divider()
    desconto = st.number_input('Percentual de desconto (%)', min_value=0.0, max_value=100.0, key='consultivo_desc_geral')
    desconto_percentual = float("{:.2f}".format(desconto))
    desconto_percentual_formatado = "{:.2f}".format(round(desconto, 2))

    total_final = 0.0
    total_final_formatado = ''
    if desconto > 0.0:
        total_final = valor_total_proposta*((100.00-desconto_percentual)/100)
        total_final_formatado = "{:.2f}".format(round(total_final, 2))
        st.write(f"Valor da proposta com desconto: R${total_final_formatado}")
    elif desconto is None:
        # st.write('o que acontece?')
        total_final = valor_total_proposta
        total_final_formatado = "{:.2f}".format(round(total_final, 2))
    total_final_formatado = total_final_formatado
    st.write(total_final_formatado)
    st.divider()    

    parcelamento = st.selectbox('Parcelamento', ['Regular', 'Entrada + parcelas'], index=None)
    numero_parcelas_formatado = ''
    valor_entrada_formatado = ''
    parcelamento_restante = ''
    numero_parcelas = 0
    parcelamento_restante = 0

    if parcelamento != None:
        if parcelamento == 'Regular':
            numero_parcelas = st.selectbox('nº de parcelas', options=range(2,25))
            numero_parcelas_formatado = "{:.2f}".format(round(numero_parcelas, 2))
            valor_parcelamento = total_final / numero_parcelas
            valor_parcelamento_formatado = "{:.2f}".format(round(valor_parcelamento, 2))
            st.write(f'O valor do parcelamento é de R$ {valor_parcelamento}')
        else:
            valor_entrada = st.number_input('Valor da entrada (R$)', min_value=1000)
            valor_entrada_formatado = "{:.2f}".format(round(valor_entrada, 2))
            saldo = total_final - valor_entrada
            st.write(f'*O saldo é de R$ {saldo}*')
            parcelamento_restante = st.selectbox('nº de parcelas', options=range(2,25))
            valor_parcelamento = saldo / parcelamento_restante
            valor_parcelamento_formatado = "{:.2f}".format(round(valor_parcelamento, 2))
            st.write(f'*O valor do parcelamento é de R$ {valor_parcelamento_formatado}*')

#####################################################################################
    # paragrafo III-III
    # Inicializar a variável 'parcelas_texto' de acordo com 'numero_parcelas' e 'parcelamento_restante'
    parcelas_texto = obter_texto_parcelas(numero_parcelas)

    # Atualizar 'parcelas_texto' caso o parcelamento seja 'Entrada + parcelas'
    if parcelamento == 'Entrada + parcelas':
        parcelas_texto = obter_texto_parcelas(parcelamento_restante)

    # Verificar e aplicar o desconto
    if desconto > 0:
        paragraph_three_three = document.add_paragraph()
        if parcelamento is None:
            paragraph_three_three.add_run("DESCONTO").bold = True
            paragraph_three_three.add_run(
                f""": Tendo em vista a parceria para com o cliente, a Roque Khouri & Pinheiro, por mera liberalidade e apenas no trabalho específico, concede o desconto de {desconto_percentual_formatado}% ({num_extenso_percentual(desconto_percentual_formatado)}) em todos os valores descritos, totalizando assim, R$ {total_final_formatado} ({num_extenso(total_final_formatado)}) pela prestação de serviços contratados, a ser pago no ato da assinatura desta proposta.""".strip()
            )
            format_paragraph(paragraph_three_three, 3, 1.5748, 0, 18, 18, 18)
        
        elif parcelamento == 'Regular':
            paragraph_three_three.add_run("DESCONTO").bold = True
            paragraph_three_three.add_run(
                f""": Tendo em vista a parceria para com o cliente, a Roque Khouri & Pinheiro, por mera liberalidade e apenas no trabalho específico, concede o desconto de {desconto_percentual_formatado}% ({num_extenso_percentual(desconto_percentual_formatado)}) em todos os valores descritos, totalizando assim, R$ {total_final_formatado} ({num_extenso(total_final_formatado)}) pela prestação de serviços contratados, a ser pagos em {parcelas_texto} iguais de R$ {valor_parcelamento_formatado} ({num_extenso(valor_parcelamento_formatado)}).""".strip()
            )
            format_paragraph(paragraph_three_three, 3, 1.5748, 0, 18, 18, 18)
        else: #parcelamento == 'Entrada + parcelas':
            paragraph_three_three.add_run("DESCONTO").bold = True
            paragraph_three_three.add_run(
                f""": Tendo em vista a parceria para com o cliente, a Roque Khouri & Pinheiro, por mera liberalidade e apenas no trabalho específico, concede o desconto de {desconto_percentual_formatado}% ({num_extenso_percentual(desconto_percentual_formatado)}) em todos os valores descritos, totalizando assim, R$ {total_final_formatado} ({num_extenso(total_final_formatado)}) pela prestação de serviços contratados, a ser pagos com entrada de R$ {valor_entrada_formatado} ({num_extenso(valor_entrada_formatado)}) e o restante dividido em {parcelas_texto} de R$ {valor_parcelamento_formatado} ({num_extenso(valor_parcelamento_formatado)}).""".strip()
            )
            format_paragraph(paragraph_three_three, 3, 1.5748, 0, 18, 18, 18)
        
    else:
        paragraph_three_three = document.add_paragraph()
        if parcelamento == 'Regular':
            # paragraph_three_three.add_run("PAGAMENTO").bold = True
            paragraph_three_three.add_run(
                f"""Para a prestação de serviços advocatícios listada no Tópico I, a Roque Khouri & Pinheiro Advogados Associados estima o pagamento de {parcelas_texto} mensais de R$ {valor_parcelamento_formatado} ({num_extenso(valor_parcelamento_formatado)}).""".strip()
            )
            format_paragraph(paragraph_three_three, 3, 1.5748, 0, 18, 18, 18)
        elif parcelamento == 'Entrada + parcelas':
            paragraph_three_three.add_run(
                f"""Para a prestação de serviços advocatícios listada no Tópico I, a Roque Khouri & Pinheiro Advogados Associados estima o pagamento de R$ {valor_entrada_formatado} ({num_extenso(valor_entrada_formatado)}) no ato da assinatura da proposta e o restante dividos em {parcelas_texto} de R$ {valor_parcelamento_formatado} ({num_extenso(valor_parcelamento_formatado)})""".strip()
            )
            format_paragraph(paragraph_three_three, 3, 1.5748, 0, 18, 18, 18)
        # else:
        #     paragraph_three_three.add_run(
        #         f"""Para a prestação de serviços advocatícios listada no Tópico I, a Roque Khouri & Pinheiro Advogados Associados estima o pagamento de R$ {total_final_formatado} ({num_extenso(total_final_formatado)}) no ato da assinatura da proposta)""".strip()
        #     )
        #     format_paragraph(paragraph_three_three, 3, 1.5748, 0, 18, 18, 18)

    #paragrafo III-IV
    paragraph_three_four = document.add_paragraph('Não estão incluídos na proposta ora apresentada eventuais custos com a contratação de advogados correspondentes fora de Brasília, bem como as despesas a serem incorridas em virtude da execução dos serviços, tais como, cópias reprográficas, custas judiciais, honorários periciais, emolumentos com autenticação de cópias e reconhecimento de firmas, obtenção de certidões, motoboys e deslocamentos à razão de R$ 1,00/km, entre outras despesas, as quais serão pagas diretamente por V.Sa. ou reembolsadas mediante a apresentação dos respectivos comprovantes.')
    format_paragraph(paragraph_three_four, 3, 1.5748,0, 18,18,18)

    #paragrafo III-IV.I
    paragraph_three_four_one = document.add_paragraph("Eventuais despesas relativas a custas judiciais e extrajudiciais, como cópias, tributos, honorários periciais, bem como despesas com o eventual deslocamento e hospedagem de pessoal da Roque Khouri & Pinheiro Advogados Associados para fora de Brasília em razão da prestação de serviços serão de responsabilidade dos Interessados. Qualquer outro serviço ou indagação, incluindo também contatos informais por aplicativo de mensagens, também serão devidamente remunerados de acordo com as horas efetivamente trabalhadas.")
    format_paragraph(paragraph_three_four_one, 3,  1.5748,0, 18,18,18)


    #paragrafo III-V
    paragraph_three_five = document.add_paragraph('Qualquer outro serviço ou indagação que não aqueles previstos no tópico I, serão estabelecidos os honorários de acordo com as horas efetivamente trabalhadas, mediante aprovação preliminar do interessado. ')
    format_paragraph(paragraph_three_five, 3, 1.5748,0, 18,18,18)


    #############################################################################
    # IV - DA CONFIDENCIALIDADE
    #Tituolo
    title_iv = document.add_heading('IV - DA CONFIDENCIALIDADE', level=2)
    format_title_justified(title_iv)
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, 3, 1.5748, 0, 18,18,18)
    paragraph.text = "O escritório e seus profissionais comprometem-se a: (i) tratar todas as informações que tiverem acesso por meio deste trabalho de forma confidencial durante o prazo de realização das atividades; e (ii) não utilizar qualquer informação confidencial para qualquer fim que não a realização dos trabalhos. Excetua-se do conceito de informação confidencial aquela que já for divulgada ou disponibilizada publicamente pelo interessado."

    #paragrafo IV-I
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, 3, 1.5748, 0, 18,18,18) 
    paragraph.text = "Atenciosamente,"


    # Adicionar imagem centralizada
    image_paragraph = document.add_paragraph()
    run = image_paragraph.add_run()
    run.add_picture(r"img/arp.png", width=Inches(2.0))
    image_paragraph.alignment = 1  # Centralizado
    image_paragraph.space_before = Pt(64)

    # Adicionar parágrafo centralizado
    paragraph = document.add_paragraph()
    paragraph.add_run('Roque Khouri & Pinheiro Advogados Associados \nÂngela Ramos Pinheiro\nOAB/DF 31.608').bold = True
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = 1  # Centralizado
    paragraph_format.space_before = Pt(4)

    # Adicionar parágrafo para "De acordo:"
    paragraph = document.add_paragraph()
    paragraph.add_run("De acordo:________________").bold = True
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(40)


    # Adicionar parágrafo para "Data:"
    paragraph = document.add_paragraph()
    paragraph.add_run("Data:_____________________").bold = True
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(32)
    paragraph = document.add_paragraph()

############################################################################
if nome_cliente:
    with desenvolvimento:
        st.write(paragraph_date.text)
        st.markdown(f"""
            <div style="text-align: right;">
                {paragraph_date.text}
            </div>
                {title.text}
            <div>
            </div>
            """, unsafe_allow_html=True)
        st.write(f'**{paragraph_para.text}**')
        st.write('')
        st.write(paragraph_ref.text)
        st.write('*texto padrao apresentação do escritorio*')
        st.write(title_one.text)
        
        # Loop até que input_objeto não esteja vazio
        while True:
            if input_objeto:
                break
            time.sleep(2)  
        if len(desdobramentos) > 1:
            for texto in textos_paragrafos:
                st.markdown(f"""
                        <div style="text-align: justify;">
                            {texto}
                        </div>
                        """, unsafe_allow_html=True)                        
        else:
            st.markdown(f"""
                        <div style="text-align: justify;">
                            {paragrah_padrao.text}
                        </div>
                        """, unsafe_allow_html=True)
        while True:
            if resumo_objeto:
                break
            time.sleep(2)


        st.write("")
        st.markdown(f"""
                        <div style="text-align: justify;">
                            {paragraph_atividades.text}
                        </div>
                        """, unsafe_allow_html=True)
        
        st.write("")    
        for item in itens_atuacao:
            st.markdown(f"""
                    <div style="text-align: justify;">
                        {recuo}-  {item};
                    </div>
                    """, unsafe_allow_html=True)
        st.markdown(f"""
                    <div style="text-align: justify;">
                        {paragraph_four.text}
                    </div>
                    """, unsafe_allow_html=True)
        st.write("")
        st.markdown(f"""
                        <div style="text-align: justify;">
                            {paragraph_five.text}
                        </div>
                        """, unsafe_allow_html=True)

        st.write("")
        st.markdown(f"""
                        <div style="text-align: justify;">
                            {title_two.text}
                        </div>
                        """, unsafe_allow_html=True)
        st.write("")
        st.write('*Texto padrão*')
        st.write("")
        st.write(title_three.text)
        st.markdown(f"""
                        <div style="text-align: justify;">
                            {paragraph_three_one.text}
                        </div>
                        """, unsafe_allow_html=True)

        st.write("")
        st.markdown(f"""
                        <div style="text-align: justify;">
                            {paragraph_three_two.text}
                        </div>
                        """, unsafe_allow_html=True)
        # while True:
        #     if valor_aplicado:
        #         break
        #     time.sleep(3)
        if tipo_cobranca == "Atuação":
            if len(atuacao) > 1:
                # st.write(df_inputs)
                for idx, row in df_inputs.iterrows():
                    f"{recuo}{idx+1}) {df_inputs['objeto'][idx]}"
                    f"{recuo}{df_inputs['total-de-horas'][idx]}h estimada para a confecção e revisão"
                    f"{recuo}Valor da hora aplicada: R$ {df_inputs['valor-aplicado'][idx]} ({df_inputs['valor-formatado'][idx]})"
                    f"{recuo}R$ {df_inputs['subtotal'][idx]} ({df_inputs['subtotal-extenso'][idx]}) estimada para a confecção e revisão"
            else:
                f"{recuo}{df_inputs['objeto']}"
                f"{recuo}{df_inputs['total-de-horas']}h estimada para a confecção e revisão"
                f"{recuo}Valor da hora aplicada: R$ {df_inputs['valor-aplicado']} ({df_inputs['valor-formatado']})"
                f"{recuo}R$ {df_inputs['subtotal']} ({df_inputs['subtotal-extenso']}) estimada para a confecção e revisão"
        
        if tipo_cobranca == 'Profissional':
                if valor_aplicado_adv == 1150.00:
                    st.write(f'{recuo}{block_three_atuacao.text}')
                    st.write(f'{recuo}{block_three_hora.text}')
                    st.write(f'{recuo}{block_three_valor_aplicado.text}')
                    st.write(f'{recuo}{block_three_subtotal.text}')
                elif valor_aplicado_adv == 850.00:
                    st.write(f'{recuo}{block_three_atuacao.text}')
                    st.write(f'{recuo}{block_three_hora.text}')
                    st.write(f'{recuo}{block_three_valor_aplicado.text}')
                    st.write(f'{recuo}{block_three_subtotal.text}')
                elif valor_aplicado_adv == 680.00:
                    st.write(f'{recuo}{block_three_atuacao.text}')
                    st.write(f'{recuo}{block_three_hora.text}')
                    st.write(f'{recuo}{block_three_valor_aplicado.text}')
                    st.write(f'{recuo}{block_three_subtotal.text}')
                elif valor_aplicado_adv == 580.00:
                    st.write(f'{recuo}{block_three_atuacao.text}')
                    st.write(f'{recuo}{block_three_hora.text}')
                    st.write(f'{recuo}{block_three_valor_aplicado.text}')
                    st.write(f'{recuo}{block_three_subtotal.text}')
                elif valor_aplicado_adv == 490.00:
                    st.write(f'{recuo}{block_three_atuacao.text}')
                    st.write(f'{recuo}{block_three_hora.text}')
                    st.write(f'{recuo}{block_three_valor_aplicado.text}')
                    st.write(f'{recuo}{block_three_subtotal.text}')
                else:
                    st.write(f'{recuo}{block_three_atuacao.text}')
                    st.write(f'{recuo}{block_three_hora.text}')
                    st.write(f'{recuo}{block_three_valor_aplicado.text}')
                    st.write(f'{recuo}{block_three_subtotal.text}')
        st.write("")
        if desconto >0:
            st.markdown(f"""
                            <div style="text-align: justify;">
                                {paragraph_three_three.text}
                            </div>
                            """, unsafe_allow_html=True)
        st.write("")
        # st.markdown(f"""
        #                 <div style="text-align: justify;">
        #                     {paragraph_three_four.text}
        #                 </div>
        #                 """, unsafe_allow_html=True)
        # st.write("")
        st.write("*texto padrão sobre a constratação de advogados correspondentes fora de Brasília e honorários sobre outros serviços*")
        # st.markdown(f"""
        #                 <div style="text-align: justify;">
        #                     {paragraph_three_five.text}
        #                 </div>
        #                 """, unsafe_allow_html=True)
        st.write("")
        st.write(title_iv.text)
        st.write("*texto padrão*")

        
        # Salvar documento em arquivo temporário e permitir download
        with NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            document.save(tmp_file.name)
            st.download_button(
                label="Baixar Documento",
                data=open(tmp_file.name, 'rb').read(),
                file_name=f'proposta_consultivo_{nome_cliente}.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
