import pandas as pd
import numpy as np
import streamlit as st
from streamlit_gsheets import GSheetsConnection
from st_pages import Page, Section, show_pages, add_page_title,add_indentation
from docx import Document
from docx.shared import Inches, Pt, Length, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from datetime import datetime
import locale
import time
import webbrowser
from utils.funcoes import format_paragraph, add_formatted_text, format_title_centered, \
    format_title_justified, num_extenso, data_extenso, fonte_name_and_size, add_section,\
    create_paragraph, atualizar_base_dados, num_extenso_percentual



st.set_page_config(layout="wide")

# Define o local para português do Brasil
locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')

add_indentation() 
# Expande a largura da tela


# Store the initial value of widgets in session state
if "visibility" not in st.session_state:
    st.session_state.visibility = "visible"
    st.session_state.disabled = False
    st.session_state.horizontal = False


#####################################################################################
dados, desenvolvimento = st.columns([2,3])


#Dicionário das informações
perguntas_respostas = {}

# Carregando a lista de clientes pela primeira vez
lista_clientes = atualizar_base_dados()

with dados:
    st.write('**Informação para a proposta**')

    
    # Carregando a lista de clientes pela primeira vez
    conn = st.connection("gsheets", type=GSheetsConnection)
    existing_data = conn.read(worksheet="cliente", ttls=5, usecols=[1])
    lista_clientes = existing_data.sort_values(by='Nome')['Nome'].unique().tolist()

    # Adiciona uma opção para cadastrar novo cliente
    lista_clientes.append("--Novo cliente--")
    lista_clientes = sorted(lista_clientes)

    nome_cliente = st.selectbox(
            'Cliente',
            (lista_clientes),
            index=None,
            placeholder='Selecione o cliente'
        )
    

    if nome_cliente == '--Novo cliente--':
        form = st.form('Novo Cliente')
        nome_cliente = form.text_input('Cadastrar novo cliente')
        form.form_submit_button("Cadastrar")
        lista_clientes.append(nome_cliente)

    input_objeto = st.text_area(label="Insira o(s) objeto(s) da proposta: (ENTER para quebra de linha) ")
    resumo_objeto = st.text_area(label="Resuma cada objeto(s) em uma frase. Dê enter para separá-los: ")
    hora_total = st.number_input(label='Total de horas:', step=10)
    hora_total = int(hora_total)
    valor_aplicado = st.selectbox("Valor aplicado",
                         (1150.00, 850.00, 680.00, 580.00, 490.00, 290.00),)
    # Arredondar o valor para duas casas decimais e formatar como string
    valor_formatado = "{:.2f}".format(round(valor_aplicado, 2))
    
    #Valor da proposta
    valor_total = 0.00
    if hora_total:
        valor_total = hora_total * valor_aplicado

    # Formatar o valor_total como string com duas casas decimais
    valor_total_formatado = "{:.2f}".format(round(valor_total, 2))
    st.write(f'**Valor total: R$ {valor_total_formatado}**')

    #desconto
    desconto = st.radio("Há desconto?", 
            ['Sim', 'Não'],
            key='desconto',
            label_visibility=st.session_state.visibility,
            disabled=st.session_state.disabled,
            horizontal=st.session_state.horizontal,
            index=None
           
            )
    #Percentual do desconto
    desconto_percentual = 0.00
    if desconto == 'Sim':
        desconto_percentual = st.number_input('Percentual do desconto ao cliente:')
    desconto_percentual = float("{:.2f}".format(desconto_percentual))
    desconto_percentual_formatado = "{:.2f}".format(round(desconto_percentual, 2))
    
    
    total_final = valor_total*((100.00-desconto_percentual)/100)

    # Formatar o total_final como string com duas casas decimais
    total_final_formatado = "{:.2f}".format(round(total_final, 2))
    if desconto == 'Sim':
        st.write(f'**Valor final com desconto: R$ {total_final_formatado}**')

    perguntas_respostas = {
        'nome_cliente': nome_cliente,
        '[objeto_texto]': input_objeto,
        '[resumo_objeto]': resumo_objeto,
        '[hora_total]': hora_total,
        '[valor_aplicado]': valor_aplicado,
        '[valor_total]': valor_total_formatado,
        '[desconto]': desconto,
        '[desconto_percentual]': desconto_percentual,
        '[total_final]': total_final
        }
    

#####################################################################################
# Criar documento em branco
document = Document()

# Definir fonte e tamanho do documento
fonte_name_and_size(document, 'Arial', 12)


# Adicionar uma sessão ao documento
section = add_section(document, 4,2.5,2,3)


# Data
data = datetime.now()
paragraph_date = document.add_paragraph(f'Brasília/DF, {data_extenso(data)}')
paragraph_format = paragraph_date.paragraph_format
paragraph_format.alignment = 2  # a direita


#############################################################################
# Adicionar título ao doc
title = document.add_heading('PROPOSTA PARA PRESTAÇÃO \nDE SERVIÇOS ADVOCATÍCIOS')
format_title_centered(title)


p_de = document.add_paragraph()
p_de.add_run('DE: ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS S/S').bold = True
p_de_format = p_de.paragraph_format
p_de_format.line_spacing = Pt(18)
p_de_format.space_before = Pt(124)
p_de_format.space_after = Pt(8)


paragraph_para = document.add_paragraph()
paragraph_para.add_run(f'PARA: {nome_cliente}  (Interessado(a)/(os)/(as))').bold = True
paragraph_format = paragraph_para.paragraph_format
paragraph_format.line_spacing = Pt(18)
paragraph_format.space_after = Pt(48)

paragraph_ref = document.add_paragraph()
paragraph_ref.text = 'Referência: PROPOSTA DE HONORÁRIOS ADVOCATÍCIOS'
paragraph_format = paragraph_ref.paragraph_format
paragraph_format.space_before = Pt(18)
paragraph_format.space_after = Pt(96)
paragraph_format.line_spacing = Pt(18)



paragraph = document.add_paragraph()
format_paragraph(paragraph, 3, 1.4764,48,16,20)
full_text= "ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS S/S, com sede no SIG - Quadra 01, Lote 495, Edifício Barão do Rio Branco, sala 244, Brasília-DF, CEP 70.610-410, telefones 3321-7043 e 3226-0137, inscrita no CNPJ sob o nº 03.899.920/0001- 81, registro na Ordem dos Advogados do Brasil – OAB/DF sob o número 616/00 – RS, endereço eletrônico www.khouriadvocacia.com.br, vem, mui respeitosamente, apresentar PROPOSTA DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS, nas condições a seguir."
# Texto que será negritado
bold_text = "ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS S/S"
# Adicionar o texto formatado ao parágrafo
add_formatted_text(paragraph, full_text, bold_text)

# I - DOS SERVIÇOS A SEREM DESENVOLVIDOS
title_one = document.add_heading('I - DOS SERVIÇOS A SEREM DESENVOLVIDOS', level=2)
format_title_justified(title_one)


# Objeto da proposta
desdobramentos = perguntas_respostas['[objeto_texto]'].split("\n")
textos_paragrafos = []
texto_padrao = []
if len(desdobramentos) > 1:
    for p in desdobramentos:
        paragrafo_ = document.add_paragraph(p)
        format_paragraph(paragrafo_, 3, 1.5748, 18,18,18)
        textos_paragrafos.append(paragrafo_.text)
else:
    paragrah_padrao = document.add_paragraph(f"Conforme solicitação, apresentamos proposta de honorários para confecção de contrato {desdobramentos[0]}")
    format_paragraph(paragrah_padrao,3, 1.5748, 18,18,18)

paragraph = document.add_paragraph('A atuação desse Jurídico compreenderá as seguintes atividades:')
format_paragraph(paragraph, 3, 1.4764,18,18,18)

#atuação

itens_atuacao = [
    "Providências preliminares de levantamento e análise de todas as informações e documentos relativos ao objeto da presente proposta, a fim de propiciar o embasamento jurídico necessário;",
    "[resumo_objeto];",
    "Participações em reuniões e eventuais discussões a respeito do contrato, incluindo em entendimentos entre as partes, caso seja necessário;"
]


itens_para_mostrar = []
atuacao = perguntas_respostas['[resumo_objeto]'].split("\n")
if len(atuacao) > 1:
    for atuar in atuacao:
        itens_atuacao.append(atuar)

    else:   
        # Substituir [resumo_objeto] pelo valor de resumo_objeto
        itens_atuacao = [item.replace("[resumo_objeto]", resumo_objeto) for item in itens_atuacao]
        itens_para_mostrar.append(itens_atuacao)

for item in itens_atuacao:
    paragraph_itens_atuacao = document.add_paragraph(style='List Number')
    # Definir o recuo apenas na primeira linha
    paragraph_itens_atuacao.paragraph_format.left_indent = Inches(1.77165)
    # Definir o alinhamento do parágrafo
    paragraph_itens_atuacao.alignment = 3
    paragraph_itens_atuacao.add_run(item)


#paragrafo I-IV
paragraph_four = document.add_paragraph()
format_paragraph(paragraph_four,3, 1.5748, 18,18,18)
paragraph_four.text ='Para o cumprimento dos serviços, o escritório disponibilizará sua equipe técnica, sendo que haverá advogado responsável pelo acompanhamento direto da demanda.'

#paragrafo I-V
paragraph_five = document.add_paragraph()
format_paragraph(paragraph_five,3, 1.5748, 18,18,18)
paragraph_five.text ='A Roque Khouri & Pinheiro Advogados Associados alerta que a análise e confecção de contrato é realizada com base no direito aplicável, jurisprudência atual e principalmente nas  informações e documentos que serão sempre fornecidos pela Interessada.'

#############################################################################
# II - DA POLÍTICA GERAL DE VALORES - HONORÁRIOS
#titulo II
title_two = document.add_heading('II - DA POLÍTICA GERAL DE VALORES - HONORÁRIOS', level=2)
format_title_justified(title_two)
#Paragrafo II-I
paragraph_two_one = document.add_paragraph()
format_paragraph(paragraph_two_one,3, 1.5748, 18,18,18)
paragraph_two_one.text = "Faz parte integrante de todas as nossas propostas de honorários os itens abaixo, componentes da nossa Política de Honorários de consultoria:" #\nTaxas horárias de honorários para projetos. Para projetos, nós cobramos valores de honorários de acordo com as seguintes taxas horárias:

#Paragrafo II-II
paragraph_two_two = document.add_paragraph()
format_paragraph(paragraph_two_two,3, 1.5748, 18,18,18)
full_text= "Taxas horárias de honorários para projetos. Para projetos, nós cobramos valores de honorários de acordo com as seguintes taxas horárias:"
# Texto que será negritado
bold_text = "Taxas horárias de honorários para projetos."
# Adicionar o texto formatado ao parágrafo
add_formatted_text(paragraph_two_two, full_text, bold_text)

# Adicionar tabela
table = document.add_table(rows=1, cols=2)
table.style =  'LightShading-Accent3' #None
# get table data -------------
items = (
    ('Sócio Majoritário - Dr. Paulo Roque', 'R$1.150,00'),
    ('Sócia Nominal - Dra. Ângela Pinheiro', 'R$850,00'),
    ('Advogado Sênior', 'R$680,00'),
    ('Advogado Pleno', 'R$580,00'),
    ('Advogado Júnior', 'R$490,00'),
    ('Paralegal/Estagiário', 'R$290,00')
)

# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].text = 'Profissional'
heading_cells[1].text = 'Valor'

# add a data row for each item
for item in items:
    cells = table.add_row().cells
    cells[0].text = item[0]
    cells[1].text = item[1]


#Paragrafo II-III
paragraph_two_three = document.add_paragraph()
format_paragraph(paragraph_two_three, 3, 1.5748, 18,18,18)
full_text= "Reembolso de Despesas. As despesas incorridas no desenvolvimento dos trabalhos, como, por exemplo, despesas com ligações telefônicas, correios, couriers e outros meios de envio de documentos, com impressão de cópias e digitalização de documentos, com taxas governamentais, com viagens, táxis e outros deslocamentos, e, se aplicável, despesas com custas processuais e outras despesas relativas a processos arbitrais, judiciais e administrativos, e honorários de advogados correspondentes, serão reembolsadas, mediante a apresentação de planilha discriminada, e, se solicitado, dos respectivos comprovantes. Nenhuma despesa superior a R$ 1.000,00 (um mil Reais) será incorrida sem sua prévia aprovação por escrito."
# Texto que será negritado
bold_text = "Reembolso de Despesas."
# Adicionar o texto formatado ao parágrafo
add_formatted_text(paragraph_two_three, full_text, bold_text)


#Paragrafo II-IV
paragraph_two_four = document.add_paragraph()
format_paragraph(paragraph_two_four, 3, 1.5748, 18,18,18)
full_text= "Interrupção ou Suspensão dos Trabalhos. Se por qualquer motivo os trabalhos forem eventualmente interrompidos ou suspensos, faremos o levantamento das horas trabalhadas e o valor de honorários pagos até então e, se houver saldo a ser pago, faremos o faturamento correspondente. Caso haja honorários a serem restituídos, a restituição será feita no mês seguinte ao da interrupção ou suspensão dos trabalhos e do valor a ser restituído serão descontados os tributos correspondentes pagos ou a pagar."
# Texto que será negritado
bold_text = "Interrupção ou Suspensão dos Trabalhos."
# Adicionar o texto formatado ao parágrafo
add_formatted_text(paragraph_two_four, full_text, bold_text)
#############################################################################

# III - DOS HONORÁRIOS ESPECÍFICOS
#titulo III
title_three = document.add_heading('III - DOS HONORÁRIOS ESPECÍFICOS', level=2)
format_title_justified(title_three)

#paragrafo III-I
paragraph_three_one = document.add_paragraph('Com o intuito de manter a proporcionalidade entre prestação de serviços e pagamento, os honorários advocatícios devidos em consequência da presente prestação de serviços seriam cobrados por meio do sistema de horas, ou seja, cada ato praticado, esse Jurídico seria remunerado de acordo com o tempo necessário para praticá-lo.')
format_paragraph(paragraph_three_one, 3, 1.5748, 18,18,18)

#paragrafo III-II
paragraph_three_two = document.add_paragraph('Para a prestação de serviços advocatícios listada no Tópico I, a Roque Khouri & Pinheiro estima os seguintes valores: ')
format_paragraph(paragraph_three_two, 3, 1.5748, 18,18,18)

#Bloco dos valores
block_three_two_resumo = document.add_paragraph(f'{resumo_objeto}')
format_paragraph(block_three_two_resumo, 3, 1.5748, 18,18,18)

block_three_two_hora = document.add_paragraph(f'{hora_total}h estimada para a confeccção e revisão')
format_paragraph(block_three_two_hora, 3, 1.5748, 18,18,18)

block_three_two_valor = document.add_paragraph(f'Valor da hora aplicada: R${valor_formatado} ({num_extenso(valor_formatado)})')
format_paragraph(block_three_two_valor, 3, 1.5748, 18,18,18)

block_three_two_valor_total = document.add_paragraph()
block_three_two_valor_total.add_run(f'Valor total: R${valor_total_formatado} ({num_extenso(valor_total_formatado)})').bold = True
format_paragraph(block_three_two_valor_total, 3, 1.5748, 18,18,18)

#paragrafo III-III
paragraph_three_three = document.add_paragraph()
if desconto == 'Sim':
    paragraph_three_three.add_run("DESCONTO").bold = True
    paragraph_three_three.add_run(f': Tendo em vista a parceria para com o cliente, a Roque Khouri & Pinheiro, por mera liberalidade e apenas no trabalho específico, concede o desconto de {desconto_percentual_formatado}% ({num_extenso_percentual(desconto_percentual_formatado)}) em todos os valores descritos, totalizando assim, R$ {total_final_formatado} ({num_extenso(total_final_formatado)}) pela prestação de serviços contratados.')
    format_paragraph(paragraph_three_three, 3, 1.5748, 18,18,18)

#paragrafo III-IV
paragraph_three_four = document.add_paragraph('Não estão incluídos na proposta ora apresentada eventuais custos com a contratação de advogados correspondentes fora de Brasília, bem como as despesas a serem incorridas em virtude da execução dos serviços, tais como, cópias reprográficas, custas judiciais, honorários periciais, emolumentos com autenticação de cópias e reconhecimento de firmas, obtenção de certidões, motoboys e deslocamentos à razão de R$ 1,00/km, entre outras despesas, as quais serão pagas diretamente por V.Sa. ou reembolsadas mediante a apresentação dos respectivos comprovantes.')
format_paragraph(paragraph_three_four, 3, 1.5748, 18,18,18)

#paragrafo III-V
paragraph_three_five = document.add_paragraph('Qualquer outro serviço ou indagação que não aqueles previstos no tópico I, serão estabelecidos os honorários de acordo com as horas efetivamente trabalhadas, mediante aprovação preliminar do interessado.')
format_paragraph(paragraph_three_five, 3, 1.5748, 18,18,18)


#############################################################################
# IV - DA CONFIDENCIALIDADE
#Tituolo
title_iv = document.add_heading('IV - DA CONFIDENCIALIDADE', level=2)
format_title_justified(title_iv)
paragraph = document.add_paragraph()
format_paragraph(paragraph, 3, 1.5748, 18,18,18)
paragraph.text = "O escritório e seus profissionais comprometem-se a: (i) tratar todas as informações que tiverem acesso por meio deste trabalho de forma confidencial durante o prazo de realização das atividades; e (ii) não utilizar qualquer informação confidencial para qualquer fim que não a realização dos trabalhos. Excetua-se do conceito de informação confidencial aquela que já for divulgada ou disponibilizada publicamente pelo interessado."

#paragrafo IV-I
paragraph = document.add_paragraph()
format_paragraph(paragraph, 3, 1.5748, 18,18,18) 
paragraph.text = "Atenciosamente,"


# Adicionar parágrafo centralizado
paragraph = document.add_paragraph()
paragraph.add_run('Roque Khouri & Pinheiro Advogados Associados \nPaulo R. Roque A. Khouri\nOAB/DF 10.671').bold = True
paragraph_format = paragraph.paragraph_format
paragraph_format.alignment = 1  # Centralizado
paragraph_format.space_before = Pt(64)

# Adicionar parágrafo para "De acordo:"
paragraph = document.add_paragraph()
paragraph.add_run("De acordo:________________").bold = True
paragraph_format = paragraph.paragraph_format
paragraph_format.space_before = Pt(40)


# Adicionar parágrafo para "Data:"
paragraph = document.add_paragraph()
paragraph.add_run("Data:_____________________").bold = True
paragraph_format = paragraph.paragraph_format
paragraph_format.space_before = Pt(32)
paragraph = document.add_paragraph()

############################################################################
with desenvolvimento:
    while True:
        if nome_cliente:
            break
        time.sleep(2)
    st.write(paragraph_date.text)
    st.write(title.text)
    st.write(p_de.text)
    st.write(f'**{paragraph_para.text}**')
    st.write(paragraph_ref.text)
    st.write('**texto padrao apresentação do escritorio**')
    st.write(title_one.text)
    # # Loop até que input_objeto não esteja vazio
    while True:
        if input_objeto:
            break
        time.sleep(2)  
    if len(desdobramentos) > 1:
        for texto in textos_paragrafos:
            st.write(texto)
    else:
        st.write(paragrah_padrao.text)
    while True:
        if resumo_objeto:
            break
        time.sleep(2)
    for item in itens_para_mostrar:
        st.write(item)
    st.write(paragraph_four.text)
    st.write(paragraph_five.text)
    st.write(title_two.text)
    st.write('**Texto padrão**')
    st.write(title_three.text)
    st.write(paragraph_three_one.text)
    st.write(paragraph_three_two.text)
    while True:
        if valor_aplicado:
            break
        time.sleep(3)
    st.write(block_three_two_resumo.text)
    st.write(block_three_two_hora.text)
    st.write(block_three_two_valor.text)
    st.write(block_three_two_valor_total.text)
    st.write(paragraph_three_three.text)
    st.write(paragraph_three_four.text)
    st.write(paragraph_three_five.text)
    st.write(title_iv.text)
    st.write("**texto padrão**")
    if st.button('Salvar'):
        # Salvar o documento
        document.save(f"proposta_consultivo_{nome_cliente}.docx")