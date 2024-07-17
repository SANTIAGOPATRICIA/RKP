import docx
import pandas as pd
import numpy as np
import streamlit as st
from streamlit_gsheets import GSheetsConnection
from st_pages import Page, Section, show_pages, add_page_title,add_indentation
from docx import Document
from docx.shared import Inches, Pt, Length, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from datetime import datetime
import locale
import time
import webbrowser



document = docx.Document()







# # Adicionar título ao doc
# title = document.add_heading('PROPOSTA PARA PRESTAÇÃO \nDE SERVIÇOS ADVOCATÍCIOS')

# valor_honorario_exito = document.add_paragraph(f'Honorários de Êxito: 10% () do benefício econômico¹ aferido ao final do processo.')
# valor_honorario_exito.add_footnote('Fica compreendido como benefício econômico todo e qualquer valor que a INTERESSADA receber em razão da propositura da ação ou valor que deixar de pagar.') # add a footnote
# # p_de = document.add_paragraph('teste de rodape')

# # p_de.add_footnote('Fica compreendido como benefício econômico todo e qualquer valor que a INTERESSADA receber em razão da propositura da ação ou valor que deixar de pagar.') # add a footnote

# document.save("teste-footnotes.docx")

# # Função para formatar o parágrafo
# def format_paragraph(paragraph, left_indent, right_indent, first_line_indent, space_before, space_after, line_spacing):
#     paragraph_format = paragraph.paragraph_format
#     paragraph_format.left_indent = docx.shared.Pt(left_indent)
#     paragraph_format.right_indent = docx.shared.Pt(right_indent)
#     paragraph_format.first_line_indent = docx.shared.Pt(first_line_indent)
#     paragraph_format.space_before = docx.shared.Pt(space_before)
#     paragraph_format.space_after = docx.shared.Pt(space_after)
#     paragraph_format.line_spacing = docx.shared.Pt(line_spacing)

# # Função para adicionar parágrafo com nota de rodapé
# def add_paragraph_with_footnote(doc, text_before, keyword, text_after, footnote_text, format_params):
#     # Adicionar parágrafo com partes separadas
#     paragraph = doc.add_paragraph()
#     paragraph.add_run(text_before)
#     run_keyword = paragraph.add_run(keyword)
#     # run_keyword.bold = True  # Exemplificando a formatação adicional no keyword
#     paragraph.add_run(f'¹{text_after}')

#     # Adicionar texto da nota de rodapé ao final do documento
#     footnote_paragraph = doc.add_paragraph()
#     footnote_paragraph.add_run(f'¹ {footnote_text}')
    
#     # Aplicar formatação
#     format_paragraph(paragraph, *format_params)
#     format_paragraph(footnote_paragraph, *format_params)

# # Criar o documento
# document = docx.Document()

# # Parâmetros de formatação (exemplo)
# format_params = (3, 0, 1.5748, 18, 18, 18)

# # Adicionar parágrafos com nota de rodapé
# add_paragraph_with_footnote(
#     document,
#     'Honorários de Êxito: {exito_percentual_formatado}% ({num_extenso_percentual(exito_percentual_formatado)}) do benefício ',
#     'econômico',
#     ' aferido ao final do processo.',
#     'Fica compreendido como benefício econômico todo e qualquer valor que a INTERESSADA receber em razão da propositura da ação ou valor que deixar de pagar.',
#     format_params
# )

# # Salvar o documento
# document.save('document_with_manual_note.docx')



from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.oxml.ns import qn

# Função para formatar o parágrafo
def format_paragraph(paragraph, left_indent, right_indent, first_line_indent, space_before, space_after, line_spacing):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Pt(left_indent)
    paragraph_format.right_indent = Pt(right_indent)
    paragraph_format.first_line_indent = Pt(first_line_indent)
    paragraph_format.space_before = Pt(space_before)
    paragraph_format.space_after = Pt(space_after)
    paragraph_format.line_spacing = Pt(line_spacing)

# Função para adicionar uma nota de rodapé com tamanho de fonte ajustado
def add_footnote_with_custom_font_size(paragraph, text_before, keyword, text_after, footnote_text):
    # Adicionar o texto ao parágrafo
    run = paragraph.add_run(f"{text_before}{keyword}{text_after}¹")
    
    # Criar a estrutura da nota de rodapé
    footnote = OxmlElement('w:footnote')
    footnote.set(qn('w:id'), '1')
    footnote_p = OxmlElement('w:p')
    footnote_r = OxmlElement('w:r')
    footnote_t = OxmlElement('w:t')
    footnote_t.text = footnote_text
    footnote_r.append(footnote_t)
    footnote_p.append(footnote_r)
    footnote.append(footnote_p)

    # Adicionar a nota de rodapé ao documento
    footnotes_part = paragraph.part._footnotes_part
    footnotes_part._element.append(footnote)

    # Aplicar formatação ao texto da nota de rodapé
    footnote_r_style = footnote_r.get_or_add_rPr()
    font_size = OxmlElement('w:sz')
    font_size.set(qn('w:val'), '20')  # 10 pt size
    footnote_r_style.append(font_size)

# Criar o documento
document = Document()

# Parâmetros de formatação (exemplo)
format_params = (3, 0, 1.5748, 18, 18, 18)

# Adicionar parágrafo com nota de rodapé personalizada
paragraph = document.add_paragraph()
add_footnote_with_custom_font_size(
    paragraph,
    'Honorários de Êxito: {exito_percentual_formatado}% ({num_extenso_percentual(exito_percentual_formatado)}) do benefício ',
    'econômico',
    ' aferido ao final do processo.',
    'Fica compreendido como benefício econômico todo e qualquer valor que a INTERESSADA receber em razão da propositura da ação ou valor que deixar de pagar.'
)

# Aplicar formatação ao parágrafo
format_paragraph(paragraph, *format_params)

# Salvar o documento
document.save('document_with_custom_footnote.docx')
