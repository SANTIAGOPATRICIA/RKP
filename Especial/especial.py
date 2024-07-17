import pandas as pd
import numpy as np
import streamlit as st
from streamlit_gsheets import GSheetsConnection
from st_pages import add_indentation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import locale
import time
from num2words import num2words
from utils.funcoes import format_paragraph, add_formatted_text, format_title_centered, \
    format_title_justified, num_extenso, data_extenso, fonte_name_and_size, add_section,\
    num_extenso_percentual, set_table_borders



st.set_page_config(layout="wide")

#identação
recuo = "&nbsp;" * 24
recuo_adicao = "&nbsp;" * 32

add_indentation()

# Define o local para português do Brasil
import locale
locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')

# Inicializar DataFrame vazio
df_inputs = pd.DataFrame(columns=[
    'objeto_consultivo', 
    'total-de-horas', 
    'valor-aplicado', 
    'valor-formatado',
    'subtotal',
    'objeto_contencioso',
    'pro_labore_inicial',
    'pro_labore_manutencao',
    'exito',
    'contencioso_desconto',
    'contencioso_valor_final'
])

# Store the initial value of widgets in session state
if "visibility" not in st.session_state:
    st.session_state.visibility = "visible"
    st.session_state.disabled = False
    st.session_state.horizontal = False


##Serviços
servicos = [
    "consultoria preliminar, com análise de informações e documentos relativos aos fatos delineados;",
    "até três reuniões (duração de uma hora) presenciais, ao telefone ou vídeo conferência;",
    "após a referida análise, confecção de nota técnica detalhada, a respeito do caso exposto em reunião preliminar, referente a estratégia de “blindagem e organização patrimonial”, respondendo as seguintes questões:" 
        ]

    #questões a serem respondidas
questoes = [
    'Como funciona uma estratégia de “blindagem patrimonial” ou ˜rearranjo patrimonial”? Quais são os objetivos e riscos envolvidos?',
    'O procedimento é 100% seguro? O que pode tornar uma reorganização patrimonial insegura? Quais são os riscos envolvidos?',
    'Existem riscos jurídicos com relação aos bens pessoais em decorrência da organização patrimonial? Há bens pertinentes a Interessada que estariam livres de riscos por dívidas contraídas por pessoa física ou pelas empresas?',
    'Por que é relevante conhecer as dívidas e obrigações envolvidas em uma reorganização patrimonial? No caso concreto, diante da documentação analisada, há dívidas e obrigações que chamam mais atenção?',
    'Da análise dos documentos e informações apresentadas, quais as repercussões e riscos trabalhistas envolvidos na situação exposta pelos Interessados? Há alguma alteração ou alerta a ser realizado para minimizar riscos trabalhistas?',
    'Quais são os riscos de operar como pessoa física e quais são os riscos de operar como pessoa jurídica? E quais são os benefícios em cada uma das opções?',
    'Da análise dos documentos e informações apresentadas, há sugestão de algum bem a ser protegido como bem de família? Por qual motivo? Quais os requisitos nesse sentido?',
    'Da análise dos documentos e informações apresentadas, quais as repercussões e riscos trabalhistas envolvidos na situação exposta pelos Interessados? Há alguma alteração ou alerta a ser realizado para minimizar riscos trabalhistas?',
    'Quais são os riscos, em linhas gerais, de um sócio/acionista de empresa limitada? E do sócio/acionista administrador?',
    'Quais são os riscos, em linhas gerais, de um sócio/acionista de uma S/A? E do sócio/acionista administrador?',
    'No caso concreto, quais os maiores riscos referentes a esses cargos vislumbrados pelo escritório? Há algum alerta importante que deve ser observado pelos interessados?',
    'Há alguma evidência sobre confusão patrimonial? Qual o risco nesse sentido?',
    'Diante do casamento do(a) Interessado(a), e na hipótese de haver dívidas em nome do conjuge, há algum risco específico atualmente?',
    'No que tange ao arranjo marital entre as partes, há alguma sugestão futura para minimizar riscos?',
    'Quais os principais riscos dos Interessados diante de todas as informações que foram prestadas?',
    'Qual os riscos enfrentados pelas partes diante de empreendimentos de risco',
    'Quais são as vantagens e riscos de uma Holding patrimonial?',
    'Considerando as informações prestadas pelos Interessados, quais as vantagens e riscos de uma holding pura? E quais as vantagens e riscos de uma holding mista?',
    'Em atenção aos impactos tributários de cada uma das empresas nas operações, há alguma saída para minimizar custos nesse sentido? Em caso positivo, há alguma divergência de interpretação entre os órgãos fiscalizadores ou algum risco que deva ser de conhecimento dos interessados?',
    'Considerando a intenção de minimizar os tributos incidentes, bem como garantir a maior proteção patrimonial em razão do casamento ou em caso de falecimento, o escritório sugere a reorganização dos ativos e empresas da Interessada? (caso positivo, poderá ser apresentado um quadro sugestivo da organização patrimonial com as repercussões tributárias pertinentes)',
    'Considerando a intenção de minimizar os tributos incidentes, bem como garantir a maior proteção patrimonial aos familiares e preparar eventual sucessão, o escritório sugere a reorganização dos ativos e empresas dos Interessados? (caso positivo, poderá ser apresentado um quadro sugestivo da organização patrimonial com as repercussões tributárias pertinentes)',
    'Quais as repercussões tributárias das sugestões do escritório?',
]

#documentos necessarios
documentos_necessarios = [
    'i. Listagem completa de todas as empresas que os Interessados possuem participação;',
    'ii. Cópia de todos os contratos sociais, estatutos, acordo de quotistas/acionistas atualizados das pessoas jurídicas que os Interessados fazem parte;',
    'iii. Existem sócios ocultos? Em caso positivo, há alguma formalização? ',
    'iv. Contratos de associações, parcerias e outras sociedades firmadas pelas pessoas jurídicas para empreendimentos e outros negócios.',
    'v. Existe alguma outra sociedade de fato existente, ou seja, que exista sem qualquer formalização?',
    'vi. Há alguma ação judicial em trâmite? Se sim, disponibilizar a cópia integral e/ou número das ações;',
    'vii. Detalhamento de eventuais dívidas, incluindo impostos, financiamentos etc.;',
    'viii. Certidões negativas (ou positivas) da Receita Federal, Receita Estadual e tribunais de Justiça dos sócios e das empresas (caso os interessados dispensem a apresentação dessas certidões, podem emitir uma declaração se responsabilizando pelas informações fornecidas);',
    'ix. Detalhamento de ativos (bens, direito, valores em conta, aplicações etc.) e eventuais créditos a receber (incluindo contratos firmados);',
    'x. Há bens em nome dos Interessados (pessoa física)?',
    'xi. Informações a respeito de eventuais garantias prestadas, informando o detalhamento dos bens e as movimentações atuais de todas as pessoas que participaram do negócio jurídico;',
    'xii. Listagem completa de todos os imóveis com a destinação específica atual de cada um deles e intenções futuras de cada um deles;',
    'xiii. Escrituras, inscrições de IPTU e certidões de ônus de todos os imóveis (tanto das pessoas jurídicas, quanto das pessoas físicas dos sócios);',
    'xiv. Listagem completa de todos os funcionários das empresas, incluindo a fonte pagadora das remunerações, resumo de atribuições e destinação dos serviços;',
    'xv. Qual a média do faturamento bruto de cada uma das empresas listadas e qual o enquadramento tributário atual delas? Há alguma perspectiva de alteração de lucro e faturamento?',
    'xvi. Relatório fático das situações não documentadas, mas que sejam relevantes para apreciação do jurídico;',
    'xvii. Relação de pessoas envolvidas e que desejam proteger (companheiro, filhos etc.)',
    'xviii. Outros documentos e esclarecimentos que julgarem necessários.'
]

lista_numerada = ['a)', 'b)', 'c)', 'd)', 'e)', 'f)', 'g)', 'h)', 'i)', 'j)', 'k)', 'l)', 'm)', 'n)', 'o)', 'p)', 'q)', 'r)', 's)', 't)']

#dividindo a tela em dados, desenvolvimento
dados, desenvolvimento = st.columns([2, 3])

with dados:
    
    st.write('**Informação para a proposta**')

    # Carregando a lista de clientes pela primeira vez
    conn = st.connection("gsheets", type=GSheetsConnection)
    existing_data = conn.read(worksheet="cliente", ttls=5, usecols=[1])
    lista_clientes = existing_data.sort_values(by='Nome')['Nome'].unique().tolist()

    # Adiciona uma opção para cadastrar novo cliente
    lista_clientes.append("--Novo cliente--")
    lista_clientes = sorted(lista_clientes)

    nome_cliente = st.selectbox(
            'Cliente',
            (lista_clientes),
            index=None,
            placeholder='Selecione o cliente'
        )
    

    if nome_cliente == '--Novo cliente--':
        form = st.form('Novo Cliente')
        nome_cliente = form.text_input('Cadastrar novo cliente')
        form.form_submit_button("Cadastrar")
    lista_clientes.append(nome_cliente)

    st.divider()
    input_objeto = st.text_area("Considerando que... (ENTER para quebra de linha)")

    perguntas_respostas = {
    '[nome_cliente]': nome_cliente,
    '[objeto_texto]': input_objeto,
    }

    paragrafos_objeto = perguntas_respostas['[objeto_texto]'].split("\n")

    #sintese do objeto
    sintese_objeto = st.text_area("Síntese do objeto")
    
    #multiselect das questoes
    questoes_selecionadas = st.multiselect('Questões', questoes)

    #Adicionar questões
    adicionar_questoes = st.text_area('Adicionar questões (ENTER para quebra de linha)')
    questoes_adicionais = adicionar_questoes.split("\n")
    
    #Append ao questões_selecionadas
    questoes_selecionadas += questoes_adicionais

    #prazo para entrega
    st.divider()
    prazo = st.selectbox("Prazo em dias para entrega do parecer jurídico", (10,15,30,60,90))

    

    #dos valores
    st.divider()
    hora_total = st.number_input(label='Total de horas:', step=10, key='hora_total_')
    hora_total = int(hora_total)
    #valor aplicado
    valor_aplicado = st.selectbox("Valor aplicado",
                            (1150.00, 850.00, 680.00, 580.00, 490.00, 290.00), key='valor_aplicado_')
    # Arredondar o valor para duas casas decimais e formatar como string
    valor_formatado = "{:.2f}".format(round(valor_aplicado, 2))
    #valor por extenso
    valor_por_extenso = num_extenso(valor_formatado)
    
    #valor do subtotal
    valor_total = (hora_total * valor_aplicado)
    valor_total_formatado = "{:.2f}".format(round(valor_total, 2))
    subtotal_extenso = num_extenso(valor_total_formatado)
    st.write(f'*Subtotal R${valor_total}*')

    #desconto
    desconto_percentual_consultivo = st.number_input('Percentual do desconto (%)', min_value=0.0, max_value=100.0, key='consultivo_desc')
    desconto_percentual_consultivo = float("{:.2f}".format(desconto_percentual_consultivo))
    desconto_percentual_formatado = "{:.2f}".format(round(desconto_percentual_consultivo, 2))

    if desconto_percentual_consultivo > 0:
        total_final = valor_total * ((100.00 - desconto_percentual_consultivo) / 100)
        total_final_formatado = "{:.2f}".format(round(total_final, 2))
        total_final_extenso = num_extenso(total_final)
        st.write(f'*O valor da proposta consultiva é de R${total_final_formatado}*')
    else:
        total_final = valor_total
        total_final_formatado = "{:.2f}".format(valor_total)  # Formatar o valor total mesmo sem desconto
        st.write(f'*O valor da proposta consultiva é de R${total_final_formatado}*')

    # Parcelamento
    parcelar = st.radio('Parcelar o valor?', ['Sim', 'Não'],
        key='parcelar',
        label_visibility=st.session_state.visibility,
        disabled=st.session_state.disabled,
        horizontal=st.session_state.horizontal,
        index=None)

    # Definir um valor padrão para parcelamento
    parcelamento = 1.0  # Se não for parcelado, será pago em uma única vez

    if parcelar == 'Sim':
        parcelamento = st.selectbox('Parcelamento do valor proposto', (2, 3, 4, 5, 6))
        valor_parcelado = total_final / parcelamento
        valor_parcelado_formatado = "{:.2f}".format(round(valor_parcelado, 2))
        st.write(f'O valor parcelado da proposta é de R${valor_parcelado_formatado}')

    st.divider()


# #####################################################################################
# Abrir documento com papel timbrado da RKP
document = Document(r".\docx\RKP-PapelTimbrado.docx")

fonte_name_and_size(document, 'Arial', 12)


# Adicionar uma sessão ao documento
section = add_section(document, 4,2.5,2,3)


# Data
data = datetime.now()
paragraph_date = document.add_paragraph(f'Brasília/DF, {data_extenso(data)}')
paragraph_format = paragraph_date.paragraph_format
paragraph_format.alignment = 2  # a direita

#############################################################################
# Adicionar título ao doc
title = document.add_heading('PROPOSTA PARA PRESTAÇÃO \nDE SERVIÇOS ADVOCATÍCIOS')
format_title_centered(title)
space_format = title.paragraph_format
space_format.space_before = Pt(48)


p_de = document.add_paragraph()
p_de.add_run('DE: ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS S/S').bold = True
p_de_format = p_de.paragraph_format
p_de_format.line_spacing = Pt(18)
p_de_format.space_before = Pt(148)
p_de_format.space_after = Pt(8)


paragraph_para = document.add_paragraph()
paragraph_para.add_run(f'PARA: {nome_cliente}  (Interessado(a))').bold = True
paragraph_format = paragraph_para.paragraph_format
paragraph_format.line_spacing = Pt(18)
paragraph_format.space_after = Pt(48)

paragraph_ref = document.add_paragraph()
paragraph_ref.text = 'Referência: PROPOSTA DE HONORÁRIOS ADVOCATÍCIOS'
paragraph_format = paragraph_ref.paragraph_format
paragraph_format.space_before = Pt(18)
paragraph_format.space_after = Pt(96)
paragraph_format.line_spacing = Pt(18)



paragraph = document.add_paragraph()
format_paragraph(paragraph, 3, 1.4764,0,48,16,20)
full_text= "ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS S/S, com sede no SIG - Quadra 01, Lote 495, Edifício Barão do Rio Branco, sala 244, Brasília-DF, CEP 70.610-410, telefones 3321-7043 e 3226-0137, inscrita no CNPJ sob o nº 03.899.920/0001- 81, registro na Ordem dos Advogados do Brasil – OAB/DF sob o número 616/00 – RS, endereço eletrônico www.khouriadvocacia.com.br, vem, mui respeitosamente, apresentar PROPOSTA DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS, nas condições a seguir."
# Texto que será negritado
bold_text = "ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS S/S"
# Adicionar o texto formatado ao parágrafo
add_formatted_text(paragraph, full_text, bold_text)

# I - DOS SERVIÇOS A SEREM DESENVOLVIDOS
title_one = document.add_heading('I - DOS SERVIÇOS A SEREM DESENVOLVIDOS', level=2)
format_title_justified(title_one)


# Objeto da proposta
textos_paragrafos = []
texto_padrao = []
if len(paragrafos_objeto) > 1:
            for i, paragrafo in enumerate(paragrafos_objeto):
                if i == len(paragrafos_objeto) - 1:
                    paragrafo_ = document.add_paragraph(f'Considerando que {paragrafos_objeto[-1]}, este jurídico elabora um esboço de proposta para a seguinte prestação de serviços advocatícios:')
                    format_paragraph(paragrafo_, 3, 0,1.5748, 18,18,18)
                else:
                    paragrafo_ = document.add_paragraph(paragrafo)
                    format_paragraph(paragrafo_, 3, 0,1.5748, 18,18,18)
                    textos_paragrafos.append(paragrafo_.text)
else:
   paragrafo_ = document.add_paragraph(f"Considerando que {input_objeto}, este jurídico elabora um esboço de proposta para a seguinte prestação de serviços advocatícios:")
   format_paragraph(paragrafo_, 3, 1.5748,0, 18,18,18)
   textos_paragrafos.append(paragrafo_.text)


# Cria uma nova lista com a mesma quantidade de elementos de `servicos`, preenchida com itens de `lista_numerada`
lista_numerada_servicos = lista_numerada[:len(servicos)]

# Serviços a serem prestados
for i in range(len(servicos)):
    paragrafo_ = document.add_paragraph(f'{lista_numerada[i]} {servicos[i]}')
    format_paragraph(paragrafo_, 3, 0,1.385827, 18, 18, 18)

#Questoes a serem respondidas

if len(questoes_selecionadas) > 1:
    for i, questao in enumerate(questoes_selecionadas):
        paragrafo_ = document.add_paragraph(f'{i+1}) {questao}')
        format_paragraph(paragrafo_, 3, 0, 1.85, 18,18,18)


#paragrafo I-IV
paragraph_four = document.add_paragraph()
format_paragraph(paragraph_four,3,1.5748, 0, 18,18,18)
paragraph_four.text =f'Para o cumprimento dos serviços, o escritório disponibilizará sua equipe técnica, sendo que haverá advogado responsável pelo acompanhamento direto da demanda. O parecer jurídico será entregue em até {prazo} dias úteis do recebimento de todos os documentos e/ou assinatura da presente proposta.'

#paragrafo I-V
paragraph_five = document.add_paragraph()
format_paragraph(paragraph_five,3,1.5748, 0, 18,18,18)
paragraph_five.text ='A Roque Khouri & Pinheiro Advogados Associados alerta que a obrigação do jurídico é de meio, não havendo garantias de nota técnica positiva, bem como que todo o trabalho será elaborado com base no direito vigente até o momento da proposta e, principalmente, com base nas informações e documentos que serão sempre fornecidos pelos Interessados. Neste ponto, ressalta-se que deverão ser apresentados vários documentos e informações antes da confecção do parecer, no intuito de garantir a efetiva prestação do serviço, entre eles:'

#paragrafo i_VI
for documento in documentos_necessarios:
    paragrafo_doc = document.add_paragraph(documento)
    format_paragraph(paragrafo_doc,3,0, 1.85,18,18,18)       

#paragrafo I_VII
paragraph_six = document.add_paragraph('Não estão incluídos na referida proposta qualquer outro serviço que não aqueles previstos nesse item, principalmente no que tange a operacionalização das sugestões que farão constar no parecer, mediação ou judicialização de qualquer demanda.')
format_paragraph(paragraph_six,3,1.5748,0, 18,18,18)

#############################################################################
# II - DA POLÍTICA GERAL DE VALORES - HONORÁRIOS
#titulo II
title_two = document.add_heading('II - DA POLÍTICA GERAL DE VALORES - HONORÁRIOS', level=2)
format_title_justified(title_two)
#Paragrafo II-I
paragraph_two_one = document.add_paragraph()
format_paragraph(paragraph_two_one,3, 1.5748, 0, 18,18,18)
paragraph_two_one.text = "Faz parte integrante de todas as nossas propostas de honorários os itens abaixo, componentes da nossa Política de Honorários de consultoria:" #\nTaxas horárias de honorários para projetos. Para projetos, nós cobramos valores de honorários de acordo com as seguintes taxas horárias:

#Paragrafo II-II
paragraph_two_two = document.add_paragraph()
format_paragraph(paragraph_two_two,3, 1.5748, 0, 18,18,18)
full_text= "Taxas horárias de honorários para projetos. Para projetos, nós cobramos valores de honorários de acordo com as seguintes taxas horárias:"
# Texto que será negritado
bold_text = "Taxas horárias de honorários para projetos."
# Adicionar o texto formatado ao parágrafo
add_formatted_text(paragraph_two_two, full_text, bold_text)

# Adicionar tabela
table = document.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style =  None #'LightShading-Accent3'
# Definir bordas da tabela
set_table_borders(table)
# get table data -------------
items = (
    ('Sócio Majoritário - Dr. Paulo Roque', 'R$1.150,00'),
    ('Sócia Nominal - Dra. Ângela Pinheiro', 'R$850,00'),
    ('Advogado Sênior', 'R$680,00'),
    ('Advogado Pleno', 'R$580,00'),
    ('Advogado Júnior', 'R$490,00'),
    ('Paralegal/Estagiário', 'R$290,00')
)

# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].text = 'Profissional'
heading_cells[1].text = 'Valor'


# Alinhar texto na primeira linha ao centro verticalmente
for cell in heading_cells:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Definir cor de fundo para a primeira linha
shading_elm_0 = OxmlElement('w:shd')
shading_elm_0.set(qn('w:fill'), 'B4FF00')  # RGB (180, 255, 0)
heading_cells[0]._element.get_or_add_tcPr().append(shading_elm_0)
shading_elm_1 = OxmlElement('w:shd')
shading_elm_1.set(qn('w:fill'), 'B4FF00')  # RGB (180, 255, 0)
heading_cells[1]._element.get_or_add_tcPr().append(shading_elm_1)


# add a data row for each item
for item in items:
    cells = table.add_row().cells
    cells[0].text = item[0]
    cells[1].text = item[1]
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Alinhar texto à direita na segunda coluna

# Definir bordas da tabela
set_table_borders(table)


#Paragrafo II-III
paragraph_two_three = document.add_paragraph()
format_paragraph(paragraph_two_three, 3, 1.5748, 0, 18,18,18)
full_text= "Reembolso de Despesas. As despesas incorridas no desenvolvimento dos trabalhos, como, por exemplo, despesas com ligações telefônicas, correios, couriers e outros meios de envio de documentos, com impressão de cópias e digitalização de documentos, com taxas governamentais, com viagens, táxis e outros deslocamentos, e, se aplicável, despesas com custas processuais e outras despesas relativas a processos arbitrais, judiciais e administrativos, e honorários de advogados correspondentes, serão reembolsadas, mediante a apresentação de planilha discriminada, e, se solicitado, dos respectivos comprovantes. Nenhuma despesa superior a R$ 1.000,00 (um mil Reais) será incorrida sem sua prévia aprovação por escrito."
# Texto que será negritado
bold_text = "Reembolso de Despesas."
# Adicionar o texto formatado ao parágrafo
add_formatted_text(paragraph_two_three, full_text, bold_text)


#Paragrafo II-IV
paragraph_two_four = document.add_paragraph()
format_paragraph(paragraph_two_four, 3, 1.5748, 0, 18,18,18)
full_text= "Interrupção ou Suspensão dos Trabalhos. Se por qualquer motivo os trabalhos forem eventualmente interrompidos ou suspensos, faremos o levantamento das horas trabalhadas e o valor de honorários pagos até então e, se houver saldo a ser pago, faremos o faturamento correspondente. Caso haja honorários a serem restituídos, a restituição será feita no mês seguinte ao da interrupção ou suspensão dos trabalhos e do valor a ser restituído serão descontados os tributos correspondentes pagos ou a pagar."
# Texto que será negritado
bold_text = "Interrupção ou Suspensão dos Trabalhos."
# Adicionar o texto formatado ao parágrafo
add_formatted_text(paragraph_two_four, full_text, bold_text)
#############################################################################

# III - DOS HONORÁRIOS ESPECÍFICOS
#titulo III
title_three = document.add_heading('III - DOS HONORÁRIOS ESPECÍFICOS', level=2)
format_title_justified(title_three)

#paragrafo III-I
paragraph_three_one = document.add_paragraph('Com o intuito de manter a proporcionalidade entre prestação de serviços e pagamento, os honorários advocatícios devidos em consequência da presente prestação de serviços seriam cobrados por meio do sistema de horas, ou seja, cada ato praticado, esse Jurídico seria remunerado de acordo com o tempo necessário para praticá-lo.')
format_paragraph(paragraph_three_one, 3, 1.5748,0, 18,18,18)


#paragrafo III-II
paragraph_three_two = document.add_paragraph('Para a prestação de serviços advocatícios listada no Tópico I, a Roque Khouri & Pinheiro estima os seguintes valores:')
format_paragraph(paragraph_three_two, 3, 1.5748,0,18,18,18)

#Paragrago III-valores
paragraph_valores = document.add_paragraph(f'{sintese_objeto}')
format_paragraph(paragraph_valores,  3, 0,1.5748, 18, 18, 18)
block_three_hora = document.add_paragraph(f'{hora_total}h estimada para a confecção e revisão')
format_paragraph(block_three_hora, 3, 0, 1.5748, 18,18,18)
block_three_valor_aplicado = document.add_paragraph(f"Valor da hora aplicada: R${valor_formatado} ({valor_por_extenso})")
format_paragraph(block_three_valor_aplicado, 3, 0, 1.5748, 18,18,18)
block_three_subtotal = document.add_paragraph(f'R${valor_total_formatado} ({subtotal_extenso}) estimada para a confecção e revisão')
format_paragraph(block_three_subtotal, 3, 0, 1.5748, 18,18,18)

#desconto
# Criação do parágrafo III-III
paragraph_three_three = document.add_paragraph()

if parcelamento > 1:
    paragraph_three_three.add_run("DESCONTO").bold = True
    paragraph_three_three.add_run(
        f""": Tendo em vista a parceria para com o cliente, a Roque Khouri & Pinheiro, por mera liberalidade e apenas no trabalho específico, concede o desconto de {desconto_percentual_formatado}% ({num_extenso_percentual(desconto_percentual_formatado)}) em todos os valores descritos, totalizando assim, R${total_final_formatado} ({num_extenso(total_final)}) pela prestação de serviços contratados, a ser pagos em {num2words(parcelamento, lang='pt_BR')} parcelas iguais de R$ {valor_parcelado_formatado} ({num_extenso(valor_parcelado_formatado)})."""
    )
    format_paragraph(paragraph_three_three, 3,  1.5748,0, 18, 18, 18)
else:
    paragraph_three_three.add_run("DESCONTO").bold = True
    paragraph_three_three.add_run(
        f""": Tendo em vista a parceria para com o cliente, a Roque Khouri & Pinheiro, por mera liberalidade e apenas no trabalho específico, concede o desconto de {desconto_percentual_formatado}% ({num_extenso_percentual(desconto_percentual_formatado)}) em todos os valores descritos, totalizando assim, R${total_final_formatado} ({num_extenso(total_final)}) pela prestação de serviços contratados.""")
    format_paragraph(paragraph_three_three, 3, 1.5748, 0, 18, 18, 18)

#paragrafo III-IV
paragraph_three_four = document.add_paragraph('Não estão incluídos na proposta ora apresentada eventuais custos com a contratação de advogados correspondentes fora de Brasília, bem como as despesas a serem incorridas em virtude da execução dos serviços, tais como, cópias reprográficas, custas judiciais, honorários periciais, emolumentos com autenticação de cópias e reconhecimento de firmas, obtenção de certidões, motoboys e deslocamentos à razão de R$ 1,00/km, entre outras despesas, as quais serão pagas diretamente por V.Sa. ou reembolsadas mediante a apresentação dos respectivos comprovantes.')
format_paragraph(paragraph_three_four, 3, 1.5748,0, 18,18,18)

#paragrafo III-IV.I
paragraph_three_four_one = document.add_paragraph("Eventuais despesas relativas a custas judiciais e extrajudiciais, como cópias, tributos, honorários periciais, bem como despesas com o eventual deslocamento e hospedagem de pessoal da Roque Khouri & Pinheiro Advogados Associados para fora de Brasília em razão da prestação de serviços serão de responsabilidade dos Interessados. Qualquer outro serviço ou indagação, incluindo também contatos informais por aplicativo de mensagem, também serão devidamente remunerados de acordo com as horas efetivamente trabalhadas.")
format_paragraph(paragraph_three_four_one, 3, 1.5748, 0, 18,18,18)


#paragrafo III-V
paragraph_three_five = document.add_paragraph('Qualquer outro serviço ou indagação que não aqueles previstos no tópico I, serão estabelecidos os honorários de acordo com as horas efetivamente trabalhadas, mediante aprovação preliminar do interessado.')
format_paragraph(paragraph_three_five, 3, 1.5748,0, 18,18,18)

#paragrafo III-VI
paragraph_three_six = document.add_paragraph('Havendo necessidade de operacionalização das sugestões envolvidas na nota técnica, incluindo confecção de qualquer documento, contratos ou atos jurídicos, propositura de ação judicial ou qualquer outro serviço não previsto no item I, deverá ser apresentado novo valor de honorários.')
format_paragraph(paragraph_three_six, 3, 1.5748,0, 18,18,18)


#############################################################################
# IV - DA CONFIDENCIALIDADE
#Titulo
title_iv = document.add_heading('IV - DA CONFIDENCIALIDADE', level=2)
format_title_justified(title_iv)
paragraph = document.add_paragraph()
format_paragraph(paragraph, 3, 1.5748, 0,18,18,18)
paragraph.text = "O escritório e seus profissionais comprometem-se a: (i) tratar todas as informações que tiverem acesso por meio deste trabalho de forma confidencial durante o prazo de realização das atividades; e (ii) não utilizar qualquer informação confidencial para qualquer fim que não a realização dos trabalhos. Excetua-se do conceito de informação confidencial aquela que já for divulgada ou disponibilizada publicamente pelo interessado."

#paragrafo IV-I
paragraph = document.add_paragraph()
format_paragraph(paragraph, 3, 1.5748,0, 18,18,18) 
paragraph.text = "Atenciosamente,"


# Adicionar parágrafo centralizado
paragraph = document.add_paragraph()
paragraph.add_run('Roque Khouri & Pinheiro Advogados Associados \nPaulo R. Roque A. Khouri\nOAB/DF 10.671').bold = True
paragraph_format = paragraph.paragraph_format
paragraph_format.alignment = 1  # Centralizado
paragraph_format.space_before = Pt(64)

# Adicionar parágrafo para "De acordo:"
paragraph = document.add_paragraph()
paragraph.add_run("De acordo:________________").bold = True
paragraph_format = paragraph.paragraph_format
paragraph_format.space_before = Pt(40)


# Adicionar parágrafo para "Data:"
paragraph = document.add_paragraph()
paragraph.add_run("Data:_____________________").bold = True
paragraph_format = paragraph.paragraph_format
paragraph_format.space_before = Pt(32)
paragraph = document.add_paragraph()

####################################################################################################

with desenvolvimento:
    while True:
        if nome_cliente:
            break
        time.sleep(2)
    st.write(paragraph_date.text)
    st.write(title.text)
    # st.write(p_de.text)
    st.write(f'**{paragraph_para.text}**')
    st.write(paragraph_ref.text)
    st.write('*texto padrao apresentação do escritorio*')
    st.write(title_one.text)
    if input_objeto:
        if len(paragrafos_objeto) > 1:
            for i, paragrafo in enumerate(paragrafos_objeto):
                if i == len(paragrafos_objeto) - 1:
                    st.write(f"Considerando que {paragrafo}, este jurídico elabora um esboço de proposta para a seguinte prestação de serviços advocatícios:")
                else:
                    st.write(f"Considerando que {paragrafo}")
        else:
            st.write(f"Considerando que {input_objeto}, este jurídico elabora um esboço de proposta para a seguinte prestação de serviços advocatícios:")
        
        #Serviços a serem prestados
        for i in range(len(servicos)):
            st.markdown(f"{recuo}{lista_numerada[i]} {servicos[i]}")
        
        #Questoes a serem respondidas
        if len(questoes_selecionadas) > 1:
            for i, questao in enumerate(questoes_selecionadas):
                st.markdown(f"{recuo_adicao} {i+1}) {questao}")
        
        #prazo de entrega
        st.write(paragraph_four.text)

        #texto padrão
        st.write(f'*texto padrão sobre o alerta do jurídico, apresentação de documentos pertinentes e aviso sobre a  não realização de objetos fora da propsota.*')
        st.write(f'*Texto padrão sobre a política de honorários*')
        st.write(" ")
        st.write(title_three.text)
        st.write(paragraph_three_two.text)
        st.markdown(f'{recuo} {paragraph_valores.text}')
        st.markdown(f'{recuo} {block_three_hora.text}')
        st.markdown(f'{recuo} {block_three_valor_aplicado.text}')
        st.markdown(f'{recuo} {block_three_subtotal.text}')
        st.markdown(f'{recuo} {paragraph_three_three.text}')
        st.write('*Texto padrão sobre custo de contratação de advogados correspondentes, despesas de custas*')
        st.write("*Texto padrão sobre a confidencialidade*")

    if st.button('Salvar'):
        # Salvar o documento
        document.save(f".\documentos_gerados\proposta_consultivo_especial_{nome_cliente}.docx")
