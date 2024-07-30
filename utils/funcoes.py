from docx.shared import Inches, Pt, Length, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT #WD_STYLE_TYPE
from docx import Document
import calendar
from num2words import num2words
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
import streamlit as st
from streamlit_gsheets import GSheetsConnection
from docx.shared import Inches
from docx.oxml.ns import qn


# Cor preta
black_color = RGBColor(0, 0, 0)


def atualizar_base_dados():
    """
    Função para ler os dados do Google Sheets e atualizar a lista de clientes
    """
    conn = st.connection("gsheets", type=GSheetsConnection)
    existing_data = conn.read(worksheet="cliente", ttls=5, usecols=[1])
    lista_clientes = existing_data.sort_values(by='Nome')['Nome'].unique().tolist()
    return lista_clientes

# Função para criar um estilo de lista numerada
# def create_list_number_style(doc):
#     styles = doc.styles
#     if 'List Number' not in styles:
#         style = styles.add_style('List Number', 2)  # 1 indica o estilo de parágrafo
        # style.font.name = 'Arial'
        # style.font.size = Document.shared.Pt(12)
        # Defina outras propriedades do estilo aqui, se necessário

# Função para adicionar bordas às células da tabela
def set_table_borders(table):
    tbl = table._tbl  # Acessar o elemento XML da tabela
    for cell in tbl.iter_tcs():
        tc_pr = cell.tcPr
        tc_borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Tamanho da borda (em pontos)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # Cor da borda
            tc_borders.append(border)
        tc_pr.append(tc_borders)

def format_title_centered(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.paragraph_format.space_after = Pt(24)

def format_title_justified(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(24)

# Função para formatar o parágrafo
def format_paragraph(paragraph, before=0, after=0, left=0, right=0):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.left_indent = Pt(left)
    paragraph.paragraph_format.right_indent = Pt(right)

# def add_formatted_text(paragraph, full_text, bold_text):
#     start = full_text.find(bold_text)
#     end = start + len(bold_text)
#     if start != -1:
#         paragraph.add_run(full_text[:start])
#         paragraph.add_run(full_text[start:end]).bold = True
#         paragraph.add_run(full_text[end:])
#     else:
#         paragraph.add_run(full_text)
def add_formatted_text(paragraph, full_text, bold_text_list=None, italic_text_list=None):
    if bold_text_list is None:
        bold_text_list = []
    if italic_text_list is None:
        italic_text_list = []

    current_pos = 0

    while current_pos < len(full_text):
        next_bold_pos = len(full_text)
        next_italic_pos = len(full_text)

        next_bold_text = None
        next_italic_text = None

        # Find the next position of any bold text
        for bold_text in bold_text_list:
            pos = full_text.find(bold_text, current_pos)
            if pos != -1 and pos < next_bold_pos:
                next_bold_pos = pos
                next_bold_text = bold_text

        # Find the next position of any italic text
        for italic_text in italic_text_list:
            pos = full_text.find(italic_text, current_pos)
            if pos != -1 and pos < next_italic_pos:
                next_italic_pos = pos
                next_italic_text = italic_text

        # Determine which comes first: bold or italic
        if next_bold_pos < next_italic_pos:
            if next_bold_pos > current_pos:
                paragraph.add_run(full_text[current_pos:next_bold_pos])
            paragraph.add_run(next_bold_text).bold = True
            current_pos = next_bold_pos + len(next_bold_text)
        elif next_italic_pos < next_bold_pos:
            if next_italic_pos > current_pos:
                paragraph.add_run(full_text[current_pos:next_italic_pos])
            paragraph.add_run(next_italic_text).italic = True
            current_pos = next_italic_pos + len(next_italic_text)
        else:
            # If there are no more bold or italic texts
            paragraph.add_run(full_text[current_pos:])
            break


def create_paragraph(document, text, bold_text=None, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, left_indent=0, space_before=0, space_after=0, line_spacing=1.15):
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, alignment, left_indent, space_before, space_after, line_spacing)
    if bold_text:
        add_formatted_text(paragraph, text, bold_text)
    else:
        paragraph.add_run(text)
    return paragraph

def add_section(document, top, left, rigth, bottom):
    sessao = document.sections[-1]
    sessao.top_margin = Cm(top)
    sessao.left_margin = Cm(left)
    sessao.rigth_margin = Cm(rigth)
    sessao.bottom_margin = Cm(bottom)
    


def fonte_name_and_size(document, font_name, font_size):
    paragrafo = document.add_paragraph()
    font = paragrafo.style.font
    font.name = font_name
    font.size = Pt(font_size)
    return paragrafo

def data_extenso(data):
    meses = [
        'janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
        'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro'
    ]
    dia = data.day
    mes = meses[data.month - 1]
    ano = data.year
    return f"{dia} de {mes} de {ano}"
   
def format_paragraph(paragraph, alignment,indent,recuo, space_before, space_after, line_spacing):
    # Formatar parágrafo
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = alignment
    paragraph_format.first_line_indent = Inches(indent)
    paragraph_format.left_indent = Inches(recuo)
    paragraph_format.space_before = Pt(space_before)
    paragraph_format.space_after = Pt(space_after)
    paragraph_format.line_spacing = Pt(line_spacing)


# Função para adicionar parágrafo com nota de rodapé
def add_paragraph_with_footnote(doc, text_before, keyword, text_after, footnote_text, format_params):
    # Adicionar parágrafo com partes separadas
    paragraph = doc.add_paragraph()
    paragraph.add_run(text_before)
    run_keyword = paragraph.add_run(keyword)
    # run_keyword.bold = True  # Exemplificando a formatação adicional no keyword
    paragraph.add_run(f'¹{text_after}')

    # Adicionar texto da nota de rodapé ao final do documento
    footnote_paragraph = doc.add_paragraph()
    footnote_paragraph.add_run(f'¹ {footnote_text}')
    
    # Aplicar formatação
    format_paragraph(paragraph, *format_params)
    format_paragraph(footnote_paragraph, *format_params)


# # Função para adicionar um trecho com formatação especial (por exemplo, negrito)
# def add_formatted_text(paragraph, full_text, bold_text):
#     # Dividir o texto completo em duas partes
#     first_part, second_part = full_text.split(bold_text, 1)
    
#     # Adicionar as partes ao parágrafo
#     paragraph.add_run(first_part)  # Primeira parte sem negrito
#     paragraph.add_run(bold_text).bold = True  # Parte a ser negritada
#     paragraph.add_run(second_part)  # Segunda parte sem negrito

# Função para adicionar múltiplos trechos em negrito ao parágrafo
def add_formatted_text(paragraph, full_text, bold_texts):
    remaining_text = full_text
    for bold_text in bold_texts:
        if bold_text in remaining_text:
            parts = remaining_text.split(bold_text, 1)
            paragraph.add_run(parts[0])  # Parte sem negrito
            paragraph.add_run(bold_text).bold = True  # Parte a ser negritada
            remaining_text = parts[1]
        else:
            paragraph.add_run(remaining_text)
            remaining_text = ''
    
    if remaining_text:
        paragraph.add_run(remaining_text)

#Função para formatar o titulo centralizado
def format_title_centered(title):
    title_format = title.runs[0].font
    title_format.name = 'Arial'  # Definir a fonte desejada
    title_format.size = Pt(12)   # Definir o tamanho da fonte
    title_format.color.rgb = black_color  # Definir a cor (preta)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centralizar o título
    title.space_before = Pt(128)

# #Função para formatar o titulo justificado
# def format_title_justified(title):
#     title_format = title.runs[0].font
#     title_format.name = 'Arial'  # Definir a fonte desejada
#     title_format.size = Pt(12)   # Definir o tamanho da fonte
#     title_format.color.rgb = black_color  # Definir a cor (preta)
    #Função para formatar o titulo justificado
def format_title_justified(title):
    for run in title.runs:
        run.bold = True
        title_format = run.font
        title_format.name = 'Arial'  # Definir a fonte desejada
        title_format.size = Pt(12)   # Definir o tamanho da fonte
        title_format.color.rgb = black_color  # Definir a cor (preta)
    title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justificar o título



def num_extenso_percentual(valor):
    """
    Tratamento do percentual por extenso
    """
    # Separa o valor em parte inteira e decimal
    parte_inteira, parte_decimal = str(valor).split('.')
    
    # Processa a parte inteira por extenso
    por_extenso_inteira = num2words(int(parte_inteira), lang='pt_BR')

    # Se houver centavos, processa-os por extenso
    if parte_decimal != '00' or parte_decimal == '0':
        por_extenso_decimal = num2words(int(parte_decimal), lang='pt_BR')
        # Formatação específica para valores monetários
        por_extenso = f'{por_extenso_inteira} ponto {por_extenso_decimal} por cento'
    else:
        por_extenso = f'{por_extenso_inteira} por cento'
    
    return por_extenso

#Tratamento do valor por extenso
def num_extenso(valor):
    # Separa o valor em parte inteira e decimal
    parte_inteira, parte_decimal = str(valor).split('.')
    
    # Processa a parte inteira por extenso
    por_extenso_inteira = num2words(int(parte_inteira), lang='pt_BR')

    # Se houver centavos, processa-os por extenso
    if parte_decimal != '00' or parte_decimal == '0':
        por_extenso_decimal = num2words(int(parte_decimal), lang='pt_BR')
        # Formatação específica para valores monetários
        por_extenso = f'{por_extenso_inteira} reais e {por_extenso_decimal} centavos'
    else:
        por_extenso = f'{por_extenso_inteira} reais'
    
    return por_extenso




# Definir a função de concordância nominal para parcelas
def obter_texto_parcelas(numero):
    if numero == 1:
        return 'uma parcela'
    elif numero == 2:
        return 'duas parcelas'
    else:
        return f"{num2words(numero, lang='pt_BR')} parcelas"