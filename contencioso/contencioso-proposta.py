import pandas as pd
import streamlit as st
# from streamlit_gsheets import GSheetsConnection
from st_pages import add_indentation
import docx
from docx import Document
from docx.shared import Pt, Inches
from datetime import datetime
import locale
import time
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
# from num2words import num2words
from tempfile import NamedTemporaryFile
from utils.funcoes import format_paragraph, add_formatted_text, format_title_centered, \
    format_title_justified, num_extenso, data_extenso, fonte_name_and_size, add_section,\
    num_extenso_percentual, set_table_borders, obter_texto_parcelas

##########################################################

## listas necessárias

#Lista dos TJs
TJs = [
    'Tribunal de Justiça do Acre (TJAC)','Tribunal de Justiça de Alagoas (TJAL)',
    'Tribunal de Justiça do Amapá (TJAP)', 'Tribunal de Justiça do Amazonas (TJAM)',
    'Tribunal de Justiça da Bahia (TJBA)', 'Tribunal de Justiça do Ceará (TJCE)',
    'Tribunal de Justiça do Distrito Federal e Territórios (TJDFT)', 'Tribunal de Justiça do Espírito Santo (TJES)',
    'Tribunal de Justiça de Goiás (TJGO)', 'Tribunal de Justiça do Maranhão (TJMA)',
    'Tribunal de Justiça de Mato Grosso (TJMT)', 'Tribunal de Justiça de Mato Grosso do Sul (TJMS)',
    'Tribunal de Justiça de Minas Gerais (TJMG)', 'Tribunal de Justiça do Pará (TJPA)',
    'Tribunal de Justiça da Paraíba (TJPB)', 'Tribunal de Justiça do Paraná (TJPR)',
    'Tribunal de Justiça de Pernambuco (TJPE)', 'Tribunal de Justiça do Piauí (TJPI)',
    'Tribunal de Justiça do Rio de Janeiro (TJRJ)', 'Tribunal de Justiça do Rio Grande do Norte (TJRN)',
    'Tribunal de Justiça do Rio Grande do Sul (TJRS)', 'Tribunal de Justiça de Rondônia (TJRO)',
    'Tribunal de Justiça de Roraima (TJRR)', 'Tribunal de Justiça de Santa Catarina (TJSC)',
    'Tribunal de Justiça de São Paulo (TJSP)', 'Tribunal de Justiça de Sergipe (TJSE)',
    'Tribunal de Justiça do Tocantins (TJTO)'
    ]

#lista numerada para replace
lista_numerada = ['a)', 'b)', 'c)', 'd)', 'e)', 'f)', 'g)', 'h)', 'i)', 'j)', 'k)', 'l)', 'm)', 'n)', 'o)', 'p)', 'q)', 'r)', 's)', 't)']


############################################################
recuo = "&nbsp;" * 24


# st.set_page_config(layout="wide")

add_indentation()

# Define o local para português do Brasil
try:
    locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
except locale.Error as e:
    print(f"Erro ao definir a localidade: {e}")

# Carregar o CSV existente ou criar um novo DataFrame
try:
    df_inputs = pd.read_csv('df_inputs.csv')
except FileNotFoundError:
    df_inputs = pd.DataFrame(columns=['nome_cliente', 'objeto_contencioso', 'instancia_superior', 'orgao', 'itens_atuacao',
                                      'pro_labore_inicial', 'pro_labore_inicial_desconto','pro_labore_manutencao', 'pro_labore_manutencao_valor_sm',
                                      'exito', 'exito_valor_teto', 'tempo_expectativa'])


# Store the initial value of widgets in session state
if "visibility" not in st.session_state:
    st.session_state.visibility = "visible"
    st.session_state.disabled = False
    st.session_state.horizontal = False

dados, desenvolvimento = st.columns([2, 3])

with dados:
    # st.table(df_inputs)
    st.write('**Informação proposta - contencioso**')

    # Carregando a lista de clientes pela primeira vez
    lista_clientes = pd.read_csv('clientes.csv')
    lista_clientes = lista_clientes.sort_values(by='Nome')['Nome'].unique().tolist()

    # Adiciona uma opção para cadastrar novo cliente
    lista_clientes.append("--Novo cliente--")
    lista_clientes = sorted(lista_clientes)

    nome_cliente = st.selectbox(
            'Cliente',
            (lista_clientes),
            index=None,
            placeholder='Selecione o cliente'
        )
    

    if nome_cliente == '--Novo cliente--':
        form = st.form('Novo Cliente')
        nome_cliente = form.text_input('Cadastrar novo cliente')
        form.form_submit_button("Cadastrar")
    lista_clientes.append(nome_cliente)

    st.divider()
    st.write("Objeto")
    input_contencioso_objeto = st.text_area(label=f"Conforme solicitação, apresentamos proposta de honorários para \
                                            atuação judicial, em defesa dos interesses de {nome_cliente}...")

    #Instancia
    instancia_ = st.selectbox("Instância", options=['primeira instância', 'segunda instância', 'tribunal superior'], index = None)
    if instancia_ == 'segunda instância':
        orgao_ = st.selectbox("Tribunal", options = TJs, index = None)
    elif instancia_ == 'tribunal superior':
        orgao_ = st.selectbox("Tribunal", options = ['STJ', "STF"], index = None)
    else:
        orgao_ = st.text_input(label = "Vara ou seção", placeholder = None)


    # Valores do contencioso
    st.divider()
    st.write('Pró-labore inicial')

    # Input para o valor do pro-labore inicial
    prolabore_inicial = st.number_input('Pró-labore inicial (R$)', key='prolabore_inicial')
    prolabore_inicial_formatado = "{:.2f}".format(round(prolabore_inicial, 2))

    #Parcelamento do prolabore inicial
    parcelamento = st.selectbox('Parcelamento', ['Regular', 'Entrada + parcelas'], index=None)

    numero_parcelas_formatado = ''
    valor_entrada_formatado = ''
    numero_parcelas = 0
    parcelamento_restante = 0

    if parcelamento != None:
        if parcelamento == 'Regular':
            numero_parcelas = st.selectbox('nº de parcelas', options=range(2,25))
            numero_parcelas_formatado = "{:.2f}".format(round(numero_parcelas, 2))
            valor_parcelamento = prolabore_inicial / numero_parcelas
            valor_parcelamento_formatado = "{:.2f}".format(round(valor_parcelamento, 2))
            st.write(f'O valor do parcelamento é de R$ {valor_parcelamento}')
        else:
            valor_entrada = st.number_input('Valor da entrada (R$)', min_value=1000)
            valor_entrada_formatado = "{:.2f}".format(round(valor_entrada, 2))
            saldo = prolabore_inicial - valor_entrada
            st.write(f'*O saldo é de R$ {saldo}*')
            parcelamento_restante = st.number_input('nº parcelas', min_value=2)
            # parcelamento_restante_formatado = "{.1f}".format(round(parcelamento_restante, 2))
            valor_parcelamento = saldo / parcelamento_restante
            valor_parcelamento_formatado = "{:.2f}".format(round(valor_parcelamento, 2))
            st.write(f'*O valor do parcelamento é de R$ {valor_parcelamento}*')
    
    # Selectbox para o tempo de isenção
    st.divider()
    st.write('Pró-labore de manutenção')
    prolabore_manutencao = st.selectbox("Tempo de isenção (meses)",
                                        (0, 6, 12, 18, 24, 30, 36, 42, 48, 54, 60),
                                        key='tempo_isencao')

    # Condicional para exibir o selectbox do pró-labore de manutenção
    prolabore_manutencao_valor = 0.0
    if prolabore_manutencao > 0:
        prolabore_manutencao_valor = st.selectbox("Proporção do salário mínimo para o pró-labore de manutenção",
                                                (0.5, 1, 1.5, 2, 2.5),
                                                key='valor_manutencao')


    # Input para o percentual do benefício econômico
    st.divider()
    st.write('Êxito')
    tipo_exito = st.selectbox('Tipo de êxito', options=['benefício econômico', 'outro'], key='exito', index=None)
    exito_percentual = 0.0
    exito_percentual_formatado = "0.00"
    exito_outro_texto = ""

    # if tipo_exito != None:
    if tipo_exito == 'benefício econômico':
        exito_percentual = st.number_input('Percentual do benefício econômico (%)', min_value=0.0, max_value=100.0, step=0.5, key='exito_percentual')
        exito_percentual_formatado = "{:.2f}".format(round(exito_percentual, 2))
    else:
        exito_percentual = st.number_input('Percentual (%)', min_value=0.0, max_value=100.0, step=0.5, key='exito_percentual')
        exito_percentual_formatado = "{:.2f}".format(round(exito_percentual, 2))
        exito_outro_texto = st.text_area('texto')

    # Input para o valor teto do êxito
    valor_teto_exito = st.number_input('Valor teto do êxito (R$)', key='valor_teto_exito')
    valor_teto_exito_formatado = "{:.2f}".format(round(valor_teto_exito, 2))

    #expectativa duração contrato
    st.divider()
    st.write('Estimativa duração do processo')
    expectativa_tempo = st.selectbox("Estimativa de duração do processo (meses)",
                                     (0, 6, 12, 18, 24, 30, 36, 42, 48, 54, 60),
                                     key='tempo_expectativa')


# #####################################################################################
# Abrir documento com papel timbrado da RKP
document = docx.Document(r"docx/RKP-PapelTimbrado.docx")

# Definir fonte e tamanho do documento
fonte_name_and_size(document, 'Arial', 12)


# Adicionar uma sessão ao documento
section = add_section(document, 5,3,2,2)


# Data
data = datetime.now()
paragraph_date = document.add_paragraph(f'Brasília/DF, {data_extenso(data)}')
paragraph_format = paragraph_date.paragraph_format
paragraph_format.alignment = 2  # a direita


#############################################################################
# Adicionar título ao doc
title = document.add_heading('PROPOSTA PARA PRESTAÇÃO \nDE SERVIÇOS ADVOCATÍCIOS')
format_title_centered(title)
space_format = title.paragraph_format
space_format.space_before = Pt(48)


p_de = document.add_paragraph()
p_de.add_run('DE: ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS S/S').bold = True
p_de_format = p_de.paragraph_format
p_de_format.line_spacing = Pt(18)
p_de_format.space_before = Pt(148)
p_de_format.space_after = Pt(8)


paragraph_para = document.add_paragraph()
paragraph_para.add_run(f'PARA: {nome_cliente} - (Interessado(a)').bold = True
paragraph_format = paragraph_para.paragraph_format
paragraph_format.line_spacing = Pt(18)
paragraph_format.space_after = Pt(48)

paragraph_ref = document.add_paragraph()
paragraph_ref.text = 'Referência: PROPOSTA DE HONORÁRIOS ADVOCATÍCIOS'
paragraph_format = paragraph_ref.paragraph_format
paragraph_format.space_before = Pt(18)
paragraph_format.space_after = Pt(96)
paragraph_format.line_spacing = Pt(18)



paragraph = document.add_paragraph()
format_paragraph(paragraph, 3, 1.5748,0, 48,16,20)
full_text= "ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS S/S, com sede no SIG - Quadra 01, Lote 495, Edifício Barão do Rio Branco, sala 244, Brasília-DF, CEP 70.610-410, telefones 3321-7043 e 3226-0137, inscrita no CNPJ sob o nº 03.899.920/0001- 81, registro na Ordem dos Advogados do Brasil – OAB/DF sob o número 616/00 – RS, endereço eletrônico www.khouriadvocacia.com.br, vem, mui respeitosamente, apresentar PROPOSTA DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS, nas condições a seguir."
# Texto que será negritado
bold_text = "ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS S/S"
# Adicionar o texto formatado ao parágrafo
add_formatted_text(paragraph, full_text, bold_text)

# I - DOS SERVIÇOS A SEREM DESENVOLVIDOS
title_one = document.add_heading('I - DOS SERVIÇOS A SEREM DESENVOLVIDOS', level=2)
format_title_justified(title_one)

#paragrafo objeto
paragraph_objeto = document.add_paragraph(f'Conforme solicitação, apresentamos proposta de honorários para atuação judicial, em defesa dos interesses de {nome_cliente} {orgao_} {input_contencioso_objeto}, para todo o acompanhamento junto ao poder judiciário, até o julgamento final em {instancia_}.')
format_paragraph(paragraph_objeto,3, 1.5748,0, 18,18,18)


#itens atuação
itens_atuacao = [
    "Providências preliminares de levantamento e análise de todas as informações e documentos relativos ao objeto da presente proposta, a fim de propiciar o embasamento jurídico necessário;",
    '',
    f"Atuação contenciosa com a confecção de petição inicial, ajuizamento de ação e acompanhamento de processo judicial até o seu julgamento final em sede de {instancia_};",
    f"Elaboração de todas as petições e recursos necessários (incluindo memoriais) ao acompanhamento da ação judicial até o seu julgamento final em sede de {instancia_};",
    "Diligências pessoais junto aos Tribunais, em especial despachos presenciais e telepresenciais com os Juízes, Desembargadores e Ministros responsáveis pelo julgamento, se cabível.",
    "Participações em reuniões com V.Sa. e demais profissionais envolvidos, incluindo em entendimentos entre as partes, se forem necessários;"
]


if instancia_ == 'segunda instância':
    itens_atuacao[1] += f'Acompanhamento do processo junto ao {orgao_}, até o julgamento final em {instancia_};'
else:
    del itens_atuacao[1]

# Cria uma nova lista com a mesma quantidade de elementos de `servicos`, preenchida com itens de `lista_numerada`
lista_numerada_servicos = lista_numerada[:len(itens_atuacao)]

# Serviços a serem prestados
for i in range(len(itens_atuacao)):
    paragrafo_ = document.add_paragraph(f'{lista_numerada[i]} {itens_atuacao[i]}')
    format_paragraph(paragrafo_, 3, 0, 1.5748, 18,18,18)


#paragrafo I-IV
paragraph_four = document.add_paragraph()
format_paragraph(paragraph_four,3, 1.5748, 0, 18,18,18)
paragraph_four.text ='Para o cumprimento dos serviços, o escritório disponibilizará sua equipe técnica, sendo que haverá advogado responsável pelo acompanhamento direto da demanda.'

#paragrafo I-V
paragraph_five = document.add_paragraph()
format_paragraph(paragraph_five,3, 1.5748,0, 18,18,18)
paragraph_five.text ='A Roque Khouri & Pinheiro Advogados Associados alerta que a análise e confecção de contrato é realizada com base no direito aplicável, jurisprudência atual e principalmente nas informações e documentos que serão sempre fornecidos pela Interessada.'

#############################################################################
# II - DA POLÍTICA GERAL DE VALORES - HONORÁRIOS
#titulo II
title_two = document.add_heading('II - DA POLÍTICA GERAL DE VALORES - HONORÁRIOS', level=2)
format_title_justified(title_two)
#Paragrafo II-I
paragraph_two_one = document.add_paragraph()
format_paragraph(paragraph_two_one,3, 1.5748, 0, 18,18,18)
paragraph_two_one.text = "Faz parte integrante de todas as nossas propostas de honorários os itens abaixo, componentes da nossa Política de Honorários de consultoria:" #\nTaxas horárias de honorários para projetos. Para projetos, nós cobramos valores de honorários de acordo com as seguintes taxas horárias:

#Paragrafo II-II
paragraph_two_two = document.add_paragraph()
format_paragraph(paragraph_two_two,3, 1.5748,0, 18,18,18)
full_text= "Taxas horárias de honorários para projetos. Para projetos, nós cobramos valores de honorários de acordo com as seguintes taxas horárias:"
# Texto que será negritado
bold_text = "Taxas horárias de honorários para projetos."
# Adicionar o texto formatado ao parágrafo
add_formatted_text(paragraph_two_two, full_text, bold_text)

# Adicionar tabela
table = document.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style =  None #'LightShading-Accent3'
# Definir bordas da tabela
set_table_borders(table)
# get table data -------------
items = (
    ('Sócio Majoritário - Dr. Paulo Roque', 'R$1.150,00'),
    ('Sócia Nominal - Dra. Ângela Pinheiro', 'R$850,00'),
    ('Advogado Sênior', 'R$680,00'),
    ('Advogado Pleno', 'R$580,00'),
    ('Advogado Júnior', 'R$490,00'),
    ('Paralegal/Estagiário', 'R$290,00')
)

# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].text = 'Profissional'
heading_cells[1].text = 'Valor'


# Alinhar texto na primeira linha ao centro verticalmente
for cell in heading_cells:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Definir cor de fundo para a primeira linha
shading_elm_0 = OxmlElement('w:shd')
shading_elm_0.set(qn('w:fill'), 'B4FF00')  # RGB (180, 255, 0)
heading_cells[0]._element.get_or_add_tcPr().append(shading_elm_0)
shading_elm_1 = OxmlElement('w:shd')
shading_elm_1.set(qn('w:fill'), 'B4FF00')  # RGB (180, 255, 0)
heading_cells[1]._element.get_or_add_tcPr().append(shading_elm_1)


# add a data row for each item
for item in items:
    cells = table.add_row().cells
    cells[0].text = item[0]
    cells[1].text = item[1]
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Alinhar texto à direita na segunda coluna

# Definir bordas da tabela
set_table_borders(table)


#Paragrafo II-III
paragraph_two_three = document.add_paragraph()
format_paragraph(paragraph_two_three, 3, 1.5748,0, 18,18,18)
full_text= "Reembolso de Despesas. As despesas incorridas no desenvolvimento dos trabalhos, como, por exemplo, despesas com ligações telefônicas, correios, couriers e outros meios de envio de documentos, com impressão de cópias e digitalização de documentos, com taxas governamentais, com viagens, táxis e outros deslocamentos, e, se aplicável, despesas com custas processuais e outras despesas relativas a processos arbitrais, judiciais e administrativos, e honorários de advogados correspondentes, serão reembolsadas, mediante a apresentação de planilha discriminada, e, se solicitado, dos respectivos comprovantes. Nenhuma despesa superior a R$ 1.000,00 (um mil Reais) será incorrida sem sua prévia aprovação por escrito."
# Texto que será negritado
bold_text = "Reembolso de Despesas."
# Adicionar o texto formatado ao parágrafo
add_formatted_text(paragraph_two_three, full_text, bold_text)


#Paragrafo II-IV
paragraph_two_four = document.add_paragraph()
format_paragraph(paragraph_two_four, 3, 1.5748,0, 18,18,18)
full_text= "Interrupção ou Suspensão dos Trabalhos. Se por qualquer motivo os trabalhos forem eventualmente interrompidos ou suspensos, faremos o levantamento das horas trabalhadas e o valor de honorários pagos até então e, se houver saldo a ser pago, faremos o faturamento correspondente. Caso haja honorários a serem restituídos, a restituição será feita no mês seguinte ao da interrupção ou suspensão dos trabalhos e do valor a ser restituído serão descontados os tributos correspondentes pagos ou a pagar."
# Texto que será negritado
bold_text = "Interrupção ou Suspensão dos Trabalhos."
# Adicionar o texto formatado ao parágrafo
add_formatted_text(paragraph_two_four, full_text, bold_text)
#############################################################################

# III - DOS HONORÁRIOS ESPECÍFICOS
#titulo III
title_three = document.add_heading('III - DOS HONORÁRIOS ESPECÍFICOS', level=2)
format_title_justified(title_three)


paragraph_three_one_one = document.add_paragraph('Os honorários advocatícios devidos em consequência da prestação de serviços previstas no item I seriam assim determinados:')
format_paragraph(paragraph_three_one_one, 3, 1.5748, 0,18,18,18)

#concordancia nominal das parcelas
parcelas_texto = obter_texto_parcelas(numero_parcelas)
# Atualizar 'parcelas_texto' caso o parcelamento seja 'Entrada + parcelas'
if parcelamento == 'Entrada + parcelas':
    parcelas_texto = obter_texto_parcelas(parcelamento_restante)

# valor atuação contencioso
#pro-labore inicial
if parcelamento == None:
    valor_prolabore_inicial = document.add_paragraph(f'a) Pró-labore inicial mínimo: R$ {prolabore_inicial_formatado} ({num_extenso(prolabore_inicial_formatado)});')
    format_paragraph(valor_prolabore_inicial, 3, 0, 1.5748, 18, 18, 18)
elif parcelamento == 'Regular':
    valor_prolabore_inicial = document.add_paragraph(f'a) Pró-labore inicial mínimo: R$ {prolabore_inicial_formatado} ({num_extenso(prolabore_inicial_formatado)}), podendo ser divido em {parcelas_texto} mensais consecutivas de R$ {valor_parcelamento}, a ser paga na assinatura deste contrato;')
    format_paragraph(valor_prolabore_inicial, 3, 0, 1.5748, 18, 18, 18)
else:
    valor_prolabore_inicial = document.add_paragraph(f'a) Pró-labore inicial mínimo: R$ {prolabore_inicial_formatado} ({num_extenso(prolabore_inicial_formatado)}), sendo a primeira parcela no valor de R$ {valor_entrada_formatado} ({num_extenso(valor_entrada_formatado)}) a ser paga na assinatura deste contrato, e {parcelas_texto} mensais consecutivas no valor de R$ {valor_parcelamento_formatado} ({num_extenso(valor_parcelamento_formatado)});')
    format_paragraph(valor_prolabore_inicial, 3, 0, 1.5748, 18, 18, 18)


# pro-labore de manutenção
if prolabore_manutencao == 0:
    valor_honorario_manutencao = document.add_paragraph(f"b) Honorário de manutenção: Isento.")
    format_paragraph(valor_honorario_manutencao, 3, 0, 1.5748, 18,18,18)
else:
    valor_honorario_manutencao = document.add_paragraph(f"b) Honorário de manutenção: Isento durante o período de {prolabore_manutencao} meses. Após este período, se o processo perdurar, será devido o valor de {prolabore_manutencao_valor} salário mínimo mensal;")
    format_paragraph(valor_honorario_manutencao, 3, 0, 1.5748, 18,18,18)

#exito
if tipo_exito == 'benefício econômico':
    valor_honorario_exito = document.add_paragraph(f'c) Honorários de Êxito: {exito_percentual_formatado}% ({num_extenso_percentual(exito_percentual_formatado)}) do benefício econômico aferido ao final do processo. Fica compreendido como benefício econômico todo e qualquer valor que a INTERESSADA receber em razão da propositura da ação ou valor que deixar de pagar.')
    # valor_honorario_exito.add_footnote('Fica compreendido como benefício econômico todo e qualquer valor que a INTERESSADA receber em razão da propositura da ação ou valor que deixar de pagar.') # add a footnote
    format_paragraph(valor_honorario_exito, 3, 0, 1.5748, 18,18,18)
else:
    valor_honorario_exito = document.add_paragraph(f'c) Honorários de Êxito: {exito_percentual_formatado}% ({num_extenso_percentual(exito_percentual_formatado)}) {exito_outro_texto}. Em caso de acordo parcial, fica estipulado que o êxito poderá ser compatível com a redução que vier a ser atingida.')
    # valor_honorario_exito.add_footnote('Em caso de acordo parcial, fica estipulado que o êxito poderá ser compatível com a redução que vier a ser atingida.') # add a footnote
    format_paragraph(valor_honorario_exito, 3, 0, 1.5748, 18,18,18)

#paragrafo III-IV
paragraph_three_four = document.add_paragraph('Não estão incluídos na proposta ora apresentada eventuais custos com a contratação de advogados correspondentes fora de Brasília, bem como as despesas a serem incorridas em virtude da execução dos serviços, tais como, cópias reprográficas, custas judiciais, honorários periciais, emolumentos com autenticação de cópias e reconhecimento de firmas, obtenção de certidões, motoboys e deslocamentos à razão de R$ 1,00/km, entre outras despesas, as quais serão pagas diretamente por V.Sa. ou reembolsadas mediante a apresentação dos respectivos comprovantes.')
format_paragraph(paragraph_three_four, 3, 1.5748, 0, 18,18,18)

#paragrafo III-IV.I
paragraph_three_four_one = document.add_paragraph("Eventuais despesas relativas a custas judiciais e extrajudiciais, como cópias, tributos, honorários periciais, bem como despesas com o eventual deslocamento e hospedagem de pessoal da Roque Khouri & Pinheiro Advogados Associados para fora de Brasília em razão da prestação de serviços serão de responsabilidade dos Interessados. Qualquer outro serviço ou indagação, incluindo também contatos informais por aplicativo de mensagem, também serão devidamente remunerados de acordo com as horas efetivamente trabalhadas.")
format_paragraph(paragraph_three_four_one, 3, 1.5748,0, 18,18,18)

#paragrafo III-IV.II
if valor_teto_exito > 0.0:
    paragraph_three_four_two = document.add_paragraph(f"Todos os valores aqui previstos serão devidamente atualizados anualmente pelo INPC ou índice que vier a substituí-lo. Todos os valores aqui previstos são devidos mesmo em caso de acordo. O valor do êxito fica limitado ao valor de R${valor_teto_exito_formatado}, devidamente atualizado.")
    format_paragraph(paragraph_three_four_two, 3, 1.5748,0, 18,18,18)
else:
    paragraph_three_four_two = document.add_paragraph("Todos os valores aqui previstos serão devidamente atualizados anualmente pelo INPC ou índice que vier a substituí-lo. Todos os valores aqui previstos são devidos mesmo em caso de acordo.")
    format_paragraph(paragraph_three_four_two, 3, 1.5748,0, 18,18,18)


#paragrafo III-V
paragraph_three_five = document.add_paragraph('Havendo necessidade de propositura de nova ação judicial que não aquela prevista no item I, deverá ser apresentado novo valor de honorários.')
format_paragraph(paragraph_three_five, 3, 1.5748, 0,18,18,18)


#############################################################################
# IV - DA CONFIDENCIALIDADE
#Tituolo
title_iv = document.add_heading('IV - DA CONFIDENCIALIDADE', level=2)
format_title_justified(title_iv)
paragraph = document.add_paragraph()
format_paragraph(paragraph, 3, 1.5748, 0, 18,18,18)
paragraph.text = "O escritório e seus profissionais comprometem-se a: (i) tratar todas as informações que tiverem acesso por meio deste trabalho de forma confidencial durante o prazo de realização das atividades; e (ii) não utilizar qualquer informação confidencial para qualquer fim que não a realização dos trabalhos. Excetua-se do conceito de informação confidencial aquela que já for divulgada ou disponibilizada publicamente pelo interessado."

#paragrafo IV-I
paragraph = document.add_paragraph()
format_paragraph(paragraph, 3, 1.5748, 0,18,18,18) 
paragraph.text = "Atenciosamente,"

# Adicionar imagem centralizada
image_paragraph = document.add_paragraph()
run = image_paragraph.add_run()
run.add_picture(r"img/arp.png", width=Inches(2.0))
image_paragraph.alignment = 1  # Centralizado
image_paragraph.space_before = Pt(64)

# Adicionar parágrafo centralizado
paragraph = document.add_paragraph()
paragraph.add_run('Roque Khouri & Pinheiro Advogados Associados \nPaulo R. Roque A. Khouri\nOAB/DF 10.671').bold = True
paragraph_format = paragraph.paragraph_format
paragraph_format.alignment = 1  # Centralizado
paragraph_format.space_before = Pt(4)


# Adicionar parágrafo para "De acordo:"
paragraph = document.add_paragraph()
paragraph.add_run("De acordo:________________").bold = True
paragraph_format = paragraph.paragraph_format
paragraph_format.space_before = Pt(40)


# Adicionar parágrafo para "Data:"
paragraph = document.add_paragraph()
paragraph.add_run("Data:_____________________").bold = True
paragraph_format = paragraph.paragraph_format
paragraph_format.space_before = Pt(32)
paragraph = document.add_paragraph()

####################################
with desenvolvimento:
    while True:
        if nome_cliente:
            break
        time.sleep(2)
    st.markdown(f"""
        <div style="text-align: right;">
            {paragraph_date.text}
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown(title.text)
    st.write(f'**{paragraph_para.text}**')
    st.write(paragraph_ref.text)
    st.write('*texto padrao apresentação do escritorio*')
    st.write(title_one.text)
    
    while True:
        if input_contencioso_objeto:
            break
        time.sleep(2)

    if input_contencioso_objeto:
        st.markdown(f"""
        <div style="text-align: justify;">
        {paragraph_objeto.text}
        </div>
        """, unsafe_allow_html=True)
        st.write("")
        st.markdown(f"""
        <div style="text-align: justify;">
        A atuação desse Jurídico compreenderá as seguintes atividades:
        </div>
        """, unsafe_allow_html=True)

    # Recuo para os itens
    
    for item in itens_atuacao:
        st.markdown(f"""
        <div style="text-align: justify;">
        {recuo}-  {item}
        </div>
        """, unsafe_allow_html=True)
        
    st.write("")
    st.write(f'*Texto padrão sobre: disposição de equipe, alerta sobre risco jurídico e política de honorários*')
    st.write("")
    st.markdown(f"""
    <div style="text-align: justify;">
    {title_three.text}
    </div>
    """, unsafe_allow_html=True)

    st.markdown(f"""
    <div style="text-align: justify;">
    {paragraph_three_one_one.text}
    </div>
    """, unsafe_allow_html=True)


    if prolabore_inicial > 0.0:
        st.markdown(f"""
        <div style="text-align: justify;">
        <p>{recuo}{valor_prolabore_inicial.text}</p>
        <p>{recuo}{valor_honorario_manutencao.text}</p>
        <p>{recuo}{valor_honorario_exito.text}</p>
        <p><i>Texto padrão sobre os eventuais custos com a contratação de advogads e despesas relativas a custas judiciais</i></p>
        <p>{paragraph_three_four_two.text}</p>
        # <p>{paragraph_three_four_two.text}</p>
        <p><i>Texto padrão sobre a necessidade de novo valor de honorários se propositura de nova ação judicial.</i></p>        
        <p>{title_iv.text}</p>        
        <p><i>Texto padrão sobre confidencialidade.</i></p>                
        </div>
        """, unsafe_allow_html=True)
        
    st.write("")
    st.write("")
    st.write("")

    # # Salvar documento em arquivo temporário e permitir download
    # with NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
    #     document.save(tmp_file.name)
    #     st.download_button(
    #         label="Baixar Documento",
    #         data=open(tmp_file.name, 'rb').read(),
    #         file_name=f'proposta_contencioso_{nome_cliente}.docx',
    #         mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    #     )
    # if st.button('Salvar dados'):
    #     novo_dado = {
    #         'nome_cliente': nome_cliente,
    #         'objeto_contencioso': input_contencioso_objeto,
    #         'instancia_superior': instancia_,
    #         'orgao': orgao_,
    #         'itens_atuacao': itens_atuacao,
    #         'pro_labore_inicial': prolabore_inicial_formatado,
    #         # 'pro_labore_inicial_desconto':prolabore_inicial_com_desconto_fomatado,
    #         'parcelamento': parcelamento,
    #         'numero_parcelas_formatado': numero_parcelas_formatado,
    #         'valor_entrada': valor_entrada_formatado,
    #         'parcelamento_restante': parcelamento_restante,
    #         'valor_parcelamento_formatado': valor_parcelamento_formatado,
    #         'pro_labore_manutencao': prolabore_manutencao,
    #         'pro_labore_manutencao_valor_sm': prolabore_manutencao_valor,
    #         'tipo_exito': tipo_exito,
    #         'exito_percentual_formatado': exito_percentual_formatado,
    #         'exito_texto': exito_outro_texto,
    #         'exito_valor_teto': valor_teto_exito_formatado,
    #         'tempo_expectativa': expectativa_tempo,
    #         }
        
    #     df_inputs = df_inputs.append(novo_dado, ignore_index=True)
    #     # Salvar o DataFrame atualizado no CSV
    #     df_inputs.to_csv('df_inputs.csv', index=False)

    #     # Limpar caracteres especiais no nome do cliente
    #     nome_cliente_formatado = re.sub(r'[^\w\s]', '_', nome_cliente)
    #     document.save(f".\documentos_gerados\proposta_contencioso_{nome_cliente_formatado}.docx")
    #     st.success('Dados salvos com sucesso!')
    
    #criar um dicionario com os dados de input
    novo_dado = {
        'nome_cliente': nome_cliente,
        'objeto_contencioso': input_contencioso_objeto,
        'instancia_superior': instancia_,
        'orgao': orgao_,
        'itens_atuacao': itens_atuacao,
        'pro_labore_inicial': prolabore_inicial_formatado,
        'parcelamento': parcelamento,
        'numero_parcelas_formatado': numero_parcelas_formatado,
        'valor_entrada': valor_entrada_formatado,
        'parcelamento_restante': parcelamento_restante,
        'valor_parcelamento_formatado': valor_parcelamento_formatado,
        'pro_labore_manutencao': prolabore_manutencao,
        'pro_labore_manutencao_valor_sm': prolabore_manutencao_valor,
        'tipo_exito': tipo_exito,
        'exito_percentual_formatado': exito_percentual_formatado,
        'exito_texto': exito_outro_texto,
        'exito_valor_teto': valor_teto_exito_formatado,
        'tempo_expectativa': expectativa_tempo,
    }

    # Salvar documento em arquivo temporário e permitir download
    with NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        document.save(tmp_file.name)
        st.download_button(
            label="Baixar Proposta",
            data=open(tmp_file.name, 'rb').read(),
            file_name=f'proposta_contencioso_{nome_cliente}.docx',
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
   
   
   ###################################################################################################################
   #CONTRATO PREENCHIDO COM OS DADOS DE INPUT DO USUARIO
    item_romano = ['i)', 'ii)', 'iii)', 'iv)', 'v)', 'vi)', 'vii)', 'viii', 'ix)', 'x)']
    
    # if novo_dado:
    # st.write(novo_dado)
    #criar documento em branco
    document = docx.Document()

    #Definir fonte e tamanho do documento
    fonte_name_and_size(document, 'Arial', 12)


    # # Adicionar uma sessão ao documento
    section = add_section(document, 3,3,2,2)

    # Adicionar título ao doc
    title = document.add_heading('CONTRATO DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS')
    format_title_centered(title)
    space_format = title.paragraph_format
    space_format.space_before = Pt(16)

    #############################################
    #cONTRATADO
    contratado = document.add_paragraph()
    full_text = 'CONTRATADO: ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS, pessoa jurídica de direito privado, CNPJ n.º 03.899.920/0001-81, neste ato representado por seu sócio administrador PAULO R. ROQUE A. KHOURI, OAB/DF 10.671.'
    bold_text = ['CONTRATADO: ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS']
    add_formatted_text(contratado, full_text, bold_text)
    format_paragraph(contratado, 3, 0, 3.14961, 18,18,18)

    #Contratante
    contratante = document.add_paragraph()
    contratante.add_run(f'CONTRATANTE: {novo_dado["nome_cliente"]}').bold = True
    contratante.add_run(', (complementar com as informações do cliente)')
    format_paragraph(contratante, 3, 0, 3.14961, 18,18,18)

    # #Objeto
    objeto_1 = document.add_paragraph(style='List Number')
    objeto_1.add_run('DO OBJETO').bold = True
    format_paragraph(objeto_1, 3, 0,0, 18,18,18)

    objeto_1_1 = document.add_paragraph()
    full_text = f'1.1. O presente tem por objeto a prestação de serviços advocatícios pelo escritório CONTRATADO ao CONTRATANTE, em defesa dos interesses de {novo_dado["nome_cliente"]} {novo_dado["objeto_contencioso"]}.'
    bold_text = [f'{novo_dado["nome_cliente"]}']
    format_paragraph(objeto_1_1, 3, 0, 0,18,18,18)
    add_formatted_text(objeto_1_1, full_text, bold_text)

    objeto_1_2 = document.add_paragraph()
    full_text = f'1.2. A atuação desse Jurídico compreenderá as seguintes atividades: '
    bold_text = ['1.2.']
    add_formatted_text(objeto_1_2, full_text, bold_text)
    format_paragraph(objeto_1_2, 3, 0, 0,18,18,18)


    # Transformar a string em uma lista de itens
    itens_lista = (novo_dado["itens_atuacao"])
    # itens_lista = itens_lista.split(',')
    

    # Cria uma nova lista com a mesma quantidade de elementos de `servicos`, preenchida com itens de `lista_numerada`
    lista_numerada_servicos = item_romano[:len(itens_lista)]
    
    # Serviços a serem prestados
    for i in range(len(itens_lista)):
        paragrafo_ = document.add_paragraph(f'{lista_numerada_servicos[i]} {itens_lista[i]}')
        format_paragraph(paragrafo_, 3, 0,1.385827, 18, 18, 18)



    objeto_1_3 = document.add_paragraph()
    full_text = '1.3. Para o cumprimento dos serviços, o escritório disponibilizará sua equipe técnica, sendo que haverá advogado designado para manter o contato com o cliente, como ponto focal de atendimento.'
    bold_text = ['1.3.']
    add_formatted_text(objeto_1_3, full_text, bold_text)
    format_paragraph(objeto_1_3, 3, 0, 0,18,18,18)

    objeto_1_4 = document.add_paragraph()
    format_paragraph(objeto_1_4, 3,0,0, 18, 18, 18)
    full_text= "1.4. Os serviços são prestados pela equipe de profissionais do CONTRATADO, que se compromete a prestá-los com ética e profissionalismo, sendo certo que se trata de obrigação de meio e não de resultado."
    bold_text = ['1.4.', 'CONTRATADO', 'obrigação de meio e não de resultado']
    add_formatted_text(objeto_1_4, full_text, bold_text)


    objeto_1_5 = document.add_paragraph()
    full_text = '1.5. O CONTRATANTE declara ter sido alertado de não haver garantias de êxito, sobretudo em razão das fases avançadas do processo, e, por essa razão, a atuação do CONTRATADO é limitada ao que já foi alegado no processo até a presente fase;'
    bold_text = ['1.5.', 'CONTRATANTE', 'CONTRATADO']
    add_formatted_text(objeto_1_5, full_text, bold_text)
    format_paragraph(objeto_1_5, 3,0,0, 18, 18, 18)

    objeto_1_6 = document.add_paragraph()
    full_text = '1.6. O CONTRATANTE declara ter sido alertado de que a análise e confecção de petições serão elaboradas com base no direito, jurisprudência do momento da produção e principalmente das informações e documentos que serão sempre fornecidos pelo CONTRATANTE.'
    bold_text = ['1.6.', 'CONTRATANTE', 'CONTRATADO']
    add_formatted_text(objeto_1_6, full_text, bold_text)
    format_paragraph(objeto_1_6, 3,0,0, 18, 18, 18)


    objeto_1_7 = document.add_paragraph()
    full_text = '1.7. O CONTRATANTE declara ter sido alertado que qualquer ação judicial implica em risco, estando os interessados cientes principalmente da possibilidade de condenação em honorários advocatícios, conforme previsto no Código de Processo Civil (10% a 20% sobre o valor da causa atualizado), além de custas, multas processuais, perícias.'
    bold_text = ['1.7.', 'CONTRATANTE' ]
    add_formatted_text(objeto_1_7, full_text, bold_text)
    format_paragraph(objeto_1_7, 3,0,0, 18, 18, 18)

    objeto_1_8 = document.add_paragraph()
    full_text = '1.8. Não estão incluídos na referida proposta qualquer outro serviço que não aqueles previstos nesse item, principalmente o acompanhamento e peticionamento em outros processos ou atuação extrajudicial, ainda que relativa aos processos mencionados no presente instrumento.'
    bold_text = ['1.8.']
    add_formatted_text(objeto_1_8, full_text, bold_text)
    format_paragraph(objeto_1_8, 3,0,0, 18, 18, 18)

    #############################################
    # do pagamento
    pagamento = document.add_paragraph(style = 'List Number')
    pagamento.add_run('DO PAGAMENTO').bold = True
    format_paragraph(pagamento, 3, 0,0, 18,18,18)


    pag_2_1 = document.add_paragraph()
    full_text = '2.1. Os honorários advocatícios devidos em consequência da prestação de serviços previstas na cláusula anterior serão assim pactuados:'
    bold_text = ['2.1.']
    add_formatted_text(pag_2_1, full_text, bold_text)
    format_paragraph(pag_2_1, 3,0,0, 18, 18, 18)

    #concordancia nominal das parcelas
    parcelas_texto = obter_texto_parcelas(novo_dado["numero_parcelas_formatado"])


    # Atualizar 'parcelas_texto' caso o parcelamento seja 'Entrada + parcelas'
    if novo_dado["parcelamento"] == 'Entrada + parcelas':
        parcelas_texto = obter_texto_parcelas(novo_dado["parcelamento_restante"])

    if novo_dado["parcelamento"] == None:
        valor_prolabore_inicial = document.add_paragraph()
        full_text = f'a) Pró-labore inicial mínimo: R$ {novo_dado["pro_labore_inicial"]} ({num_extenso(novo_dado["pro_labore_inicial"])});'
        bold_text = ['a)']
        add_formatted_text(valor_prolabore_inicial, full_text, bold_text)
        format_paragraph(valor_prolabore_inicial, 3, 0, 1.385827, 18, 18, 18)
    elif novo_dado["parcelamento"] == 'Regular':
        valor_prolabore_inicial = document.add_paragraph()
        full_text = f'a) Pró-labore inicial mínimo: R$ {novo_dado["pro_labore_inicial"]} ({num_extenso(novo_dado["pro_labore_inicial"])}), podendo ser divido em {parcelas_texto} mensais consecutivas de R$ {novo_dado["valor_parcelamento_formatado"]}, a ser paga na assinatura deste contrato;'
        bold_text = ['a)']
        add_formatted_text(valor_prolabore_inicial, full_text, bold_text)
        format_paragraph(valor_prolabore_inicial, 3, 0, 1.385827, 18, 18, 18)
    # else:
    #     valor_prolabore_inicial = document.add_paragraph()
    #     full_text = f'a) Pró-labore inicial mínimo: R$ {novo_dado["pro_labore_inicial"]} ({num_extenso(novo_dado["pro_labore_inicial"])}), sendo a primeira parcela no valor de R$ {cliente[9]} ({num_extenso(cliente[9])}) a ser paga na assinatura deste contrato, e {parcelas_texto} mensais consecutivas no valor de R$ {cliente[11]} ({num_extenso(cliente[11])});'
    #     bold_text = ['a)']
    #     add_formatted_text(valor_prolabore_inicial, full_text, bold_text)
    #     format_paragraph(valor_prolabore_inicial, 3, 0, 1.385827, 18, 18, 18)


    #pro-labore de manutenção
    if novo_dado["pro_labore_manutencao"] == 0:
        pag2_1_b = document.add_paragraph()
        full_text = 'b) Honorário de manutenção: Isento.'
        bold_text = ['b)']
        add_formatted_text(pag2_1_b, full_text, bold_text)
        format_paragraph(pag2_1_b, 3, 0,1.385827, 18, 18, 18)
    else:
        pag2_1_b = document.add_paragraph()
        full_text = f'b) Honorário de manutenção: Isento durante o período de {novo_dado["pro_labore_manutencao"]} meses.  Após este período, se o processo perdurar, será devido o valor de {novo_dado["pro_labore_manutencao_valor_sm"]} salário mínimo mensal;'
        bold_text = ['b)']
        add_formatted_text(pag2_1_b, full_text, bold_text)
        format_paragraph(pag2_1_b, 3, 0,1.385827, 18, 18, 18)

    # #exito
    if novo_dado["tipo_exito"] == 'benefício econômico':
        pag2_1_c = document.add_paragraph()
        full_text = f'c) Honorários de Êxito: {novo_dado["exito_percentual_formatado"]}% ({num_extenso_percentual(novo_dado["exito_percentual_formatado"])}) do benefício econômico¹ aferido ao final do processo.'
        bold_text = ['c)']
        add_formatted_text(pag2_1_c, full_text, bold_text)
        pag2_1_c.add_footnote('Fica compreendido como benefício econômico todo e qualquer valor que a INTERESSADA receber em razão da propositura da ação ou valor que deixar de pagar.') # add a footnote
        format_paragraph(pag2_1_c, 3, 0, 1.385827, 18,18,18)
    else:
        pag2_1_c = document.add_paragraph()
        full_text = f'c) Honorários de Êxito: {novo_dado["exito_percentual_formatado"]}% ({num_extenso_percentual(novo_dado["exito_percentual_formatado"])}) do {novo_dado["tipo_exito"]}.'
        bold_text = ['c)']
        add_formatted_text(pag2_1_c, full_text, bold_text)
        pag2_1_c.add_footnote(f'{novo_dado["exito_texto"]}.') # add a footnote
        format_paragraph(pag2_1_c, 3, 0, 1.385827, 18,18,18)
    
    pag2_2 = document.add_paragraph()
    full_text = '2.2) Em caso de acordo parcial, fica estipulado que o êxito poderá ser compatível coma redução que vier a ser atingida. '
    bold_text = ['2.2)']
    add_formatted_text(pag2_2, full_text, bold_text)
    format_paragraph(pag2_2, 3,0,0, 18, 18, 18)


    pag2_3 = document.add_paragraph()
    full_text = '2.3) Em caso de atraso, serão devidos juros de mora de 1% ao mês, correção monetária pro rata die pelo INPC e multa moratória de 2%, a incidir de maneira retroativa a cada um dos vencimentos.'
    bold_text_list = ['2.3)']
    add_formatted_text(pag2_3, full_text, bold_text_list)
    format_paragraph(pag2_3,3,0,0, 18, 18, 18)

    pag2_4 = document.add_paragraph()
    full_text = '2.4) Sem prejuízo dos valores propostos acima, em caso de êxito parcial ou total do CONTRATANTE, 100% do valor dos eventuais honorários sucumbenciais eventualmente fixados na ação judicial serão devidos ao CONTRATADO, sendo que os honorários de sucumbência não se confundem com os honorários de êxito contratuais, visto que os honorários sucumbenciais são devidos pela outra parte enquanto os honorários contratuais são devidos pelo CONTRATANTE.'
    bold_text = ['2.4)', '2.5)']
    add_formatted_text(pag2_4, full_text, bold_text)
    format_paragraph(pag2_4, 3,0,0, 18, 18, 18)

    pag2_5 = document.add_paragraph()
    full_text = '2.5) Não estão incluídos na proposta ora apresentada eventuais custos com a contratação de advogados correspondentes fora de Brasília-DF, bem como as despesas a serem incorridas em virtude da execução dos serviços, tais como, cópias reprográficas, custas judiciais, honorários periciais, contratação de assistente técnico, emolumentos com autenticação de cópias e reconhecimento de firmas, obtenção de certidões, motoboys, as quais serão pagas diretamente pelo CONTRATANTE ou reembolsadas mediante a apresentação dos respectivos comprovantes.'
    bold_text = ['2.5)']
    add_formatted_text(pag2_5, full_text, bold_text)
    format_paragraph(pag2_5, 3,0,0, 18, 18, 18)

    pag2_6 = document.add_paragraph()
    full_text = '2.6) Todos os valores previstos nesse instrumento serão automaticamente corrigidos monetariamente de forma anual pelo Índice do INPC ou índice que vier a substituí-lo.'
    bold_text = ['2.6)']
    add_formatted_text(pag2_6, full_text, bold_text)
    format_paragraph(pag2_6, 3,0,0, 18, 18, 18)

    #############################################
    #das obrigações
    obrigacoes = document.add_paragraph(style = 'List Number')
    obrigacoes.add_run('OBRIGAÇÕES DAS PARTES').bold = True
    format_paragraph(obrigacoes, 3, 0,0, 18,18,18)


    ob3_1 = document.add_paragraph()
    full_text = '3.1) Obrigações do CONTRATANTE:'
    bold_text = ['3.1)']
    add_formatted_text(ob3_1, full_text, bold_text)
    format_paragraph(ob3_1, 3,0,0, 18, 18, 18)

    lista_obrigacoes = [
        'a) Pagar pontualmente os honorários do CONTRATADO, atendidas as condições da Cláusula 2;',
        'b) Enviar ao CONTRATADO, previamente e, em caso de novos documentos, tão logo os receba, todos as informações e documentos que possam influenciar na referida análise;',
        'c) Informar ao CONTRATADO todos os documentos solicitados e necessários para a execução do presente contrato, no prazo máximo de 72 horas;', 
        'd) Informar imediatamente ao CONTRATADO eventuais mudanças de endereço, telefone e qualquer outro meio de contato;', 
        'e) confirmar ao CONTRATADO o recebimento de todos os e-mails por estesenviados quando os receber;'
    ]

    for obrigacao in lista_obrigacoes:
        obrigacao_p = document.add_paragraph(obrigacao)
        format_paragraph(obrigacao_p, 3, 0,1.385827, 18, 18, 18)

    ob3_2 = document.add_paragraph()
    full_text = '3.2) Obrigações do CONTRATADO:'
    bold_text = ['3.2)']
    add_formatted_text(ob3_2, full_text, bold_text)
    format_paragraph(ob3_2, 3,0,0, 18, 18, 18)

    lista_obrigacoes_contratado = [
        'a) executar o objeto do contrato com responsabilidade técnica e eficácia;',
        'b) comunicar de imediato ao CONTRATANTE qualquer mudança de seu endereço ou telefone, inclusive endereço eletrônico (e-mail);',
        'c) confirmar ao CONTRATANTE o recebimento de todos os e-mails por estes enviados quando os receber;',
        'd) manter, durante toda a execução do contrato, todas as condições de habilitação jurídica, técnica e de regularidade fiscal;',
        'e) recolher todos os tributos e taxas incidentes sobre a contratação.'
    ]


    for obrigacao in lista_obrigacoes_contratado:
        obrigacao_c = document.add_paragraph(obrigacao)
        format_paragraph(obrigacao_c, 3, 0,1.385827, 18, 18, 18)

    #############################################
    # da vigência
    vigencia = document.add_paragraph(style = 'List Number')
    vigencia.add_run('DA VIGÊNCIA').bold = True
    format_paragraph(vigencia, 3, 0,0, 18,18,18)


    vig4_1 = document.add_paragraph()
    full_text = f'4.1) A prestação de serviços objeto do presente contrato vigerá até o julgamento final em sede de {novo_dado["orgao"]}.'
    bold_text = ['4.1)']
    add_formatted_text(vig4_1, full_text, bold_text)
    format_paragraph(vig4_1, 3,0,0, 18, 18, 18)


    if int(novo_dado["tempo_expectativa"]) > 0:
        vig4_2 = document.add_paragraph()
        full_text = f'4.2) As partes concordam que têm a expectativa de duração do processo de {int(novo_dado["tempo_expectativa"])} meses. Caso o tempo supere esse prazo as partes se comprometem em renegociar o contrato em comum acordo.'
        bold_text = ['4.2)']
        add_formatted_text(vig4_2, full_text, bold_text)
        format_paragraph(vig4_2, 3,0,0, 18, 18, 18)

        vig4_3 = document.add_paragraph()
        full_text = f'4.3) Em caso de rescisão antecipada do contrato, permanecerão devidos os honorários aqui previstos, incluindo de êxito.'
        bold_text = ['4.3)']
        add_formatted_text(vig4_3, full_text, bold_text)
        format_paragraph(vig4_3, 3,0,0, 18, 18, 18)
    else:
        vig4_2 = document.add_paragraph()
        full_text = f'4.2) Em caso de rescisão antecipada do contrato, permanecerão devidos os honoráriosaqui previstos, incluindo de êxito.'
        bold_text = ['4.2)']
        add_formatted_text(vig4_2, full_text, bold_text)
        format_paragraph(vig4_2, 3,0,0, 18, 18, 18)


    #############################################
    #Outras informações
    outras_infor = document.add_paragraph(style = 'List Number')
    outras_infor.add_run('OUTRAS INFORMAÇÕES').bold = True
    format_paragraph(outras_infor, 3, 0,0, 18,18,18)


    outras_info5_1 = document.add_paragraph()
    full_text = '5.1) Todas as dúvidas e informações sobre a referida prestação de serviços poderão ser dirimidas, preferencialmente, por e-mail, telefone, mensagens ou WhatsApp, desde que em horário comercial (09:00h às 12:00h e 14:00h às 19:00h).'
    bold_text = ['5.1)']
    add_formatted_text(outras_info5_1, full_text, bold_text)
    format_paragraph(outras_info5_1, 3,0,0, 18, 18, 18)

    outras_info5_2 = document.add_paragraph()
    full_text = '5.2) As partes se comprometem a (i) tratar todas as informações que tiverem acesso por meio deste trabalho de forma confidencial durante o prazo de realização das atividades; e (ii) não utilizar qualquer informação confidencial para qualquer fim que não a realização dos trabalhos. Excetua-se do conceito de informação confidencial aquela que já for divulgada ou disponibilizada publicamente pelo CONTRATANTE.'
    bold_text = ['5.2)']
    add_formatted_text(outras_info5_2, full_text, bold_text)
    format_paragraph(outras_info5_2, 3,0,0, 18, 18, 18)

    outras_info5_3 = document.add_paragraph()
    full_text = '5.3) O CONTRATANTE está autorizado a substabelecer os poderes outorgados, desde com reserva de poderes.'
    bold_text = ['5.3)']
    add_formatted_text(outras_info5_3, full_text, bold_text)
    format_paragraph(outras_info5_3, 3,0,0, 18, 18, 18)

    #############################################
    #do foro
    foro = document.add_paragraph(style = 'List Number')
    foro.add_run('DO FORO').bold = True
    format_paragraph(foro, 3, 0,0, 18,18,18)


    foro6_1 = document.add_paragraph()
    full_text = '6.1) Fica eleito o foro de Brasília (DF), com a exclusão de qualquer outro, ainda que privilegiado, para dirimir quaisquer litígios referentes ao cumprimento das obrigações ora assumidas.'
    bold_text = ['6.1)']
    add_formatted_text(foro6_1, full_text, bold_text)
    format_paragraph(foro6_1, 3,0,0, 18, 18, 18)

    #############################################
    #Data, assinantes
    #data
    data = datetime.now()
    paragraph_date = document.add_paragraph(f'Brasília/DF, {data_extenso(data)}')
    paragraph_format = paragraph_date.paragraph_format
    paragraph_format.alignment = 1  # centralizado?
    paragraph_format.space_before = Pt(32)
    #paragraph_date = WD_ALIGN_PARAGRAPH.CENTER

    paragraph = document.add_paragraph()
    paragraph.add_run(f'{novo_dado["nome_cliente"].upper()}\nCONTRATANTE').bold = True
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = 1  # Centralizado
    paragraph_format.space_after = Pt(64)
    paragraph_format.space_before = Pt(64)

    # Adicionar imagem centralizada
    image_paragraph = document.add_paragraph()
    run = image_paragraph.add_run()
    run.add_picture(r"img/arp.png", width=Inches(2.0))
    image_paragraph.alignment = 1  # Centralizado
    # image_paragraph.space_before = Pt(4)

    #assinatura RKP
    paragraph = document.add_paragraph()
    paragraph.add_run('ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS\nCONTRATADO').bold = True
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = 1  # Centralizado

    

    with NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        document.save(tmp_file.name)
        st.download_button(
            label="Baixar Contrato",
            data=open(tmp_file.name, 'rb').read(),
            file_name=f'contrato_contencioso_{novo_dado["nome_cliente"]}.docx',
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            key='proposta_contencioso'
                )


    
    


# def save_to_db(novo_dado):
#     conn = sqlite3.connect('database.db')
#     c = conn.cursor()
#     c.execute('''
#         INSERT INTO propostas (
#             nome_cliente, objeto_contencioso, instancia_superior, orgao, itens_atuacao,
#             pro_labore_inicial, parcelamento, numero_parcelas_formatado, valor_entrada,
#             parcelamento_restante, valor_parcelamento_formatado, pro_labore_manutencao,
#             pro_labore_manutencao_valor_sm, tipo_exito, exito_percentual_formatado,
#             exito_texto, exito_valor_teto, tempo_expectativa
#         ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
#     ''', (
#         novo_dado['nome_cliente'], novo_dado['objeto_contencioso'], novo_dado['instancia_superior'],
#         novo_dado['orgao'], ','.join(novo_dado['itens_atuacao']),  # Converting list to string
#         novo_dado['pro_labore_inicial'], novo_dado['parcelamento'], novo_dado['numero_parcelas_formatado'],
#         novo_dado['valor_entrada'], novo_dado['parcelamento_restante'], novo_dado['valor_parcelamento_formatado'],
#         novo_dado['pro_labore_manutencao'], novo_dado['pro_labore_manutencao_valor_sm'],
#         novo_dado['tipo_exito'], novo_dado['exito_percentual_formatado'], novo_dado['exito_texto'],
#         novo_dado['exito_valor_teto'], novo_dado['tempo_expectativa']
#     ))
#     conn.commit()
#     conn.close()


# with NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
#     document.save(tmp_file.name)
#     st.download_button(
#         label="Baixar Documento",
#         data=open(tmp_file.name, 'rb').read(),
#         file_name=f'proposta_contencioso_{nome_cliente}.docx',
#         mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
#         key='proposta_contencioso'
#             )

# if st.button('Salvar dados'):
#     novo_dado = {
#         'nome_cliente': nome_cliente,
#         'objeto_contencioso': input_contencioso_objeto,
#         'instancia_superior': instancia_,
#         'orgao': orgao_,
#         'itens_atuacao': itens_atuacao,
#         'pro_labore_inicial': prolabore_inicial_formatado,
#         'parcelamento': parcelamento,
#         'numero_parcelas_formatado': numero_parcelas_formatado,
#         'valor_entrada': valor_entrada_formatado,
#         'parcelamento_restante': parcelamento_restante,
#         'valor_parcelamento_formatado': valor_parcelamento_formatado,
#         'pro_labore_manutencao': prolabore_manutencao,
#         'pro_labore_manutencao_valor_sm': prolabore_manutencao_valor,
#         'tipo_exito': tipo_exito,
#         'exito_percentual_formatado': exito_percentual_formatado,
#         'exito_texto': exito_outro_texto,
#         'exito_valor_teto': valor_teto_exito_formatado,
#         'tempo_expectativa': expectativa_tempo,
#     }
#     save_to_db(novo_dado)
#     st.success("Dados salvos com sucesso!")
