import pandas as pd
import streamlit as st
from streamlit_gsheets import GSheetsConnection
import docx
from docx import Document
from docx.shared import Pt
from datetime import datetime
from num2words import num2words
from ast import literal_eval
import re
from tempfile import NamedTemporaryFile
from utils.funcoes import format_paragraph, add_formatted_text, format_title_centered, \
    num_extenso, data_extenso, fonte_name_and_size, add_section,\
    num_extenso_percentual, obter_texto_parcelas, load_data



#############################################
#Listas necessárias

item_romano = ['i)', 'ii)', 'iii)', 'iv)', 'v)', 'vi)', 'vii)', 'viii', 'ix)', 'x)']

#identação
recuo = "&nbsp;" * 24
recuo_adicao = "&nbsp;" * 32

#############################################
data = load_data()
lista_clientes = [cliente[1] for cliente in data]
cliente_selecionado = st.selectbox('Contrato do cliente', lista_clientes)

# Obter o índice do cliente selecionado
indice_cliente = lista_clientes.index(cliente_selecionado)


# Obter os dados do cliente selecionado usando o índice
cliente = data[indice_cliente]

if data:
    # cliente = data[-1]
    #criar documento em branco
    document = docx.Document()

    #Definir fonte e tamanho do documento
    fonte_name_and_size(document, 'Arial', 12)


    # Adicionar uma sessão ao documento
    section = add_section(document, 3,3,2,2)

    # Adicionar título ao doc
    title = document.add_heading('CONTRATO DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS')
    format_title_centered(title)
    space_format = title.paragraph_format
    space_format.space_before = Pt(16)

    #############################################
    #cONTRATADO
    contratado = document.add_paragraph()
    full_text = 'CONTRATADO: ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS, pessoa jurídica de direito privado, CNPJ n.º 03.899.920/0001-81, neste ato representado por seu sócio administrador PAULO R. ROQUE A. KHOURI, OAB/DF 10.671.'
    bold_text = ['CONTRATADO: ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS']
    add_formatted_text(contratado, full_text, bold_text)
    format_paragraph(contratado, 3, 0, 3.14961, 18,18,18)

    #Contratante
    contratante = document.add_paragraph()
    contratante.add_run(f'CONTRATANTE: {cliente[1]}' ).bold = True
    contratante.add_run(', (complementar com as informações do cliente)')
    format_paragraph(contratante, 3, 0, 3.14961, 18,18,18)

    # #Objeto
    objeto_1 = document.add_paragraph(style='List Number')
    objeto_1.add_run('DO OBJETO').bold = True
    format_paragraph(objeto_1, 3, 0,0, 18,18,18)

    objeto_1_1 = document.add_paragraph()
    full_text = f'1.1. O presente tem por objeto a prestação de serviços advocatícios pelo escritório CONTRATADO ao CONTRATANTE, em defesa dos interesses de {cliente[1]} {cliente[2]}.'
    bold_text = [f'{cliente[1]}']
    format_paragraph(objeto_1_1, 3, 0, 0,18,18,18)
    add_formatted_text(objeto_1_1, full_text, bold_text)

    objeto_1_2 = document.add_paragraph()
    full_text = f'1.2. A atuação desse Jurídico compreenderá as seguintes atividades: '
    bold_text = ['1.2.']
    add_formatted_text(objeto_1_2, full_text, bold_text)
    format_paragraph(objeto_1_2, 3, 0, 0,18,18,18)


    # Transformar a string em uma lista de itens
    itens_lista = (cliente[5])
    itens_lista = itens_lista.split(';,')
    

    # Cria uma nova lista com a mesma quantidade de elementos de `servicos`, preenchida com itens de `lista_numerada`
    lista_numerada_servicos = item_romano[:len(itens_lista)]
    
    # Serviços a serem prestados
    for i in range(len(itens_lista)):
        paragrafo_ = document.add_paragraph(f'{lista_numerada_servicos[i]} {itens_lista[i]}')
        format_paragraph(paragrafo_, 3, 0,1.385827, 18, 18, 18)



    objeto_1_3 = document.add_paragraph()
    full_text = '1.3. Para o cumprimento dos serviços, o escritório disponibilizará sua equipe técnica, sendo que haverá advogado designado para manter o contato com o cliente, como ponto focal de atendimento.'
    bold_text = ['1.3.']
    add_formatted_text(objeto_1_3, full_text, bold_text)
    format_paragraph(objeto_1_3, 3, 0, 0,18,18,18)

    objeto_1_4 = document.add_paragraph()
    format_paragraph(objeto_1_4, 3,0,0, 18, 18, 18)
    full_text= "1.4. Os serviços são prestados pela equipe de profissionais do CONTRATADO, que se compromete a prestá-los com ética e profissionalismo, sendo certo que se trata de obrigação de meio e não de resultado."
    bold_text = ['1.4.', 'CONTRATADO', 'obrigação de meio e não de resultado']
    add_formatted_text(objeto_1_4, full_text, bold_text)


    objeto_1_5 = document.add_paragraph()
    full_text = '1.5. O CONTRATANTE declara ter sido alertado de não haver garantias de êxito, sobretudo em razão das fases avançadas do processo, e, por essa razão, a atuação do CONTRATADO é limitada ao que já foi alegado no processo até a presente fase;'
    bold_text = ['1.5.', 'CONTRATANTE', 'CONTRATADO']
    add_formatted_text(objeto_1_5, full_text, bold_text)
    format_paragraph(objeto_1_5, 3,0,0, 18, 18, 18)

    objeto_1_6 = document.add_paragraph()
    full_text = '1.6. O CONTRATANTE declara ter sido alertado de que a análise e confecção de petições serão elaboradas com base no direito, jurisprudência do momento da produção e principalmente das informações e documentos que serão sempre fornecidos pelo CONTRATANTE.'
    bold_text = ['1.6.', 'CONTRATANTE', 'CONTRATADO']
    add_formatted_text(objeto_1_6, full_text, bold_text)
    format_paragraph(objeto_1_6, 3,0,0, 18, 18, 18)


    objeto_1_7 = document.add_paragraph()
    full_text = '1.7. O CONTRATANTE declara ter sido alertado que qualquer ação judicial implica em risco, estando os interessados cientes principalmente da possibilidade de condenação em honorários advocatícios, conforme previsto no Código de Processo Civil (10% a 20% sobre o valor da causa atualizado), além de custas, multas processuais, perícias.'
    bold_text = ['1.7.', 'CONTRATANTE' ]
    add_formatted_text(objeto_1_7, full_text, bold_text)
    format_paragraph(objeto_1_7, 3,0,0, 18, 18, 18)

    objeto_1_8 = document.add_paragraph()
    full_text = '1.8. Não estão incluídos na referida proposta qualquer outro serviço que não aqueles previstos nesse item, principalmente o acompanhamento e peticionamento em outros processos ou atuação extrajudicial, ainda que relativa aos processos mencionados no presente instrumento.'
    bold_text = ['1.8.']
    add_formatted_text(objeto_1_8, full_text, bold_text)
    format_paragraph(objeto_1_8, 3,0,0, 18, 18, 18)

    #############################################
    # do pagamento
    pagamento = document.add_paragraph(style = 'List Number')
    pagamento.add_run('DO PAGAMENTO').bold = True
    format_paragraph(pagamento, 3, 0,0, 18,18,18)


    pag_2_1 = document.add_paragraph()
    full_text = '2.1. Os honorários advocatícios devidos em consequência da prestação de serviços previstas na cláusula anterior serão assim pactuados:'
    bold_text = ['2.1.']
    add_formatted_text(pag_2_1, full_text, bold_text)
    format_paragraph(pag_2_1, 3,0,0, 18, 18, 18)

    #concordancia nominal das parcelas
    parcelas_texto = obter_texto_parcelas(cliente[8])


    # Atualizar 'parcelas_texto' caso o parcelamento seja 'Entrada + parcelas'
    if cliente[7] == 'Entrada + parcelas':
        parcelas_texto = obter_texto_parcelas(cliente[8])

    if cliente[7] == None:
        valor_prolabore_inicial = document.add_paragraph()
        full_text = f'a) Pró-labore inicial mínimo: R$ {cliente[6]} ({num_extenso(cliente[6])});'
        bold_text = ['a)']
        add_formatted_text(valor_prolabore_inicial, full_text, bold_text)
        format_paragraph(valor_prolabore_inicial, 3, 0, 1.385827, 18, 18, 18)
    elif cliente[7] == 'Regular':
        valor_prolabore_inicial = document.add_paragraph()
        full_text = f'a) Pró-labore inicial mínimo: R$ {cliente[6]} ({num_extenso(cliente[6])}), podendo ser divido em {parcelas_texto} mensais consecutivas de R$ {cliente[11]}, a ser paga na assinatura deste contrato;'
        bold_text = ['a)']
        add_formatted_text(valor_prolabore_inicial, full_text, bold_text)
        format_paragraph(valor_prolabore_inicial, 3, 0, 1.385827, 18, 18, 18)
    else:
        valor_prolabore_inicial = document.add_paragraph()
        full_text = f'a) Pró-labore inicial mínimo: R$ {cliente[6]} ({num_extenso(cliente[6])}), sendo a primeira parcela no valor de R$ {cliente[9]} ({num_extenso(cliente[9])}) a ser paga na assinatura deste contrato, e {parcelas_texto} mensais consecutivas no valor de R$ {cliente[11]} ({num_extenso(cliente[11])});'
        bold_text = ['a)']
        add_formatted_text(valor_prolabore_inicial, full_text, bold_text)
        format_paragraph(valor_prolabore_inicial, 3, 0, 1.385827, 18, 18, 18)


    #pro-labore de manutenção
    if cliente[12] == 0:
        pag2_1_b = document.add_paragraph()
        full_text = 'b) Honorário de manutenção: Isento.'
        bold_text = ['b)']
        add_formatted_text(pag2_1_b, full_text, bold_text)
        format_paragraph(pag2_1_b, 3, 0,1.385827, 18, 18, 18)
    else:
        pag2_1_b = document.add_paragraph()
        full_text = f'b) Honorário de manutenção: Isento durante o período de {cliente[12]} meses.  Após este período, se o processo perdurar, será devido o valor de {cliente[13]} salário mínimo mensal;'
        bold_text = ['b)']
        add_formatted_text(pag2_1_b, full_text, bold_text)
        format_paragraph(pag2_1_b, 3, 0,1.385827, 18, 18, 18)

    # #exito
    if cliente[14] == 'benefício econômico':
        pag2_1_c = document.add_paragraph()
        full_text = f'c) Honorários de Êxito: {cliente[15]}% ({num_extenso_percentual(cliente[15])}) do benefício econômico¹ aferido ao final do processo.'
        bold_text = ['c)']
        add_formatted_text(pag2_1_c, full_text, bold_text)
        pag2_1_c.add_footnote('Fica compreendido como benefício econômico todo e qualquer valor que a INTERESSADA receber em razão da propositura da ação ou valor que deixar de pagar.') # add a footnote
        format_paragraph(pag2_1_c, 3, 0, 1.385827, 18,18,18)
    else:
        pag2_1_c = document.add_paragraph()
        full_text = f'c) Honorários de Êxito: {cliente[15]}% ({num_extenso_percentual(cliente[15])}) do {cliente[16]}.'
        bold_text = ['c)']
        add_formatted_text(pag2_1_c, full_text, bold_text)
        pag2_1_c.add_footnote(f'{cliente[16]}.') # add a footnote
        format_paragraph(pag2_1_c, 3, 0, 1.385827, 18,18,18)
    
    pag2_2 = document.add_paragraph()
    full_text = '2.2) Em caso de acordo parcial, fica estipulado que o êxito poderá ser compatível coma redução que vier a ser atingida. '
    bold_text = ['2.2)']
    add_formatted_text(pag2_2, full_text, bold_text)
    format_paragraph(pag2_2, 3,0,0, 18, 18, 18)


    pag2_3 = document.add_paragraph()
    full_text = '2.3) Em caso de atraso, serão devidos juros de mora de 1% ao mês, correção monetária pro rata die pelo INPC e multa moratória de 2%, a incidir de maneira retroativa a cada um dos vencimentos.'
    bold_text_list = ['2.3)']
    add_formatted_text(pag2_3, full_text, bold_text_list)
    format_paragraph(pag2_3,3,0,0, 18, 18, 18)

    pag2_4 = document.add_paragraph()
    full_text = '2.4) Sem prejuízo dos valores propostos acima, em caso de êxito parcial ou total do CONTRATANTE, 100% do valor dos eventuais honorários sucumbenciais eventualmente fixados na ação judicial serão devidos ao CONTRATADO, sendo que os honorários de sucumbência não se confundem com os honorários de êxito contratuais, visto que os honorários sucumbenciais são devidos pela outra parte enquanto os honorários contratuais são devidos pelo CONTRATANTE.'
    bold_text = ['2.4)', '2.5)']
    add_formatted_text(pag2_4, full_text, bold_text)
    format_paragraph(pag2_4, 3,0,0, 18, 18, 18)

    pag2_5 = document.add_paragraph()
    full_text = '2.5) Não estão incluídos na proposta ora apresentada eventuais custos com a contratação de advogados correspondentes fora de Brasília-DF, bem como as despesas a serem incorridas em virtude da execução dos serviços, tais como, cópias reprográficas, custas judiciais, honorários periciais, contratação de assistente técnico, emolumentos com autenticação de cópias e reconhecimento de firmas, obtenção de certidões, motoboys, as quais serão pagas diretamente pelo CONTRATANTE ou reembolsadas mediante a apresentação dos respectivos comprovantes.'
    bold_text = ['2.5)']
    add_formatted_text(pag2_5, full_text, bold_text)
    format_paragraph(pag2_5, 3,0,0, 18, 18, 18)

    pag2_6 = document.add_paragraph()
    full_text = '2.6) Todos os valores previstos nesse instrumento serão automaticamente corrigidos monetariamente de forma anual pelo Índice do INPC ou índice que vier a substituí-lo.'
    bold_text = ['2.6)']
    add_formatted_text(pag2_6, full_text, bold_text)
    format_paragraph(pag2_6, 3,0,0, 18, 18, 18)

    #############################################
    #das obrigações
    obrigacoes = document.add_paragraph(style = 'List Number')
    obrigacoes.add_run('OBRIGAÇÕES DAS PARTES').bold = True
    format_paragraph(obrigacoes, 3, 0,0, 18,18,18)


    ob3_1 = document.add_paragraph()
    full_text = '3.1) Obrigações do CONTRATANTE:'
    bold_text = ['3.1)']
    add_formatted_text(ob3_1, full_text, bold_text)
    format_paragraph(ob3_1, 3,0,0, 18, 18, 18)

    lista_obrigacoes = [
        'a) Pagar pontualmente os honorários do CONTRATADO, atendidas as condições da Cláusula 2;',
        'b) Enviar ao CONTRATADO, previamente e, em caso de novos documentos, tão logo os receba, todos as informações e documentos que possam influenciar na referida análise;',
        'c) Informar ao CONTRATADO todos os documentos solicitados e necessários para a execução do presente contrato, no prazo máximo de 72 horas;', 
        'd) Informar imediatamente ao CONTRATADO eventuais mudanças de endereço, telefone e qualquer outro meio de contato;', 
        'e) confirmar ao CONTRATADO o recebimento de todos os e-mails por estesenviados quando os receber;'
    ]

    for obrigacao in lista_obrigacoes:
        obrigacao_p = document.add_paragraph(obrigacao)
        format_paragraph(obrigacao_p, 3, 0,1.385827, 18, 18, 18)

    ob3_2 = document.add_paragraph()
    full_text = '3.2) Obrigações do CONTRATADO:'
    bold_text = ['3.2)']
    add_formatted_text(ob3_2, full_text, bold_text)
    format_paragraph(ob3_2, 3,0,0, 18, 18, 18)

    lista_obrigacoes_contratado = [
        'a) executar o objeto do contrato com responsabilidade técnica e eficácia;',
        'b) comunicar de imediato ao CONTRATANTE qualquer mudança de seu endereço ou telefone, inclusive endereço eletrônico (e-mail);',
        'c) confirmar ao CONTRATANTE o recebimento de todos os e-mails por estes enviados quando os receber;',
        'd) manter, durante toda a execução do contrato, todas as condições de habilitação jurídica, técnica e de regularidade fiscal;',
        'e) recolher todos os tributos e taxas incidentes sobre a contratação.'
    ]


    for obrigacao in lista_obrigacoes_contratado:
        obrigacao_c = document.add_paragraph(obrigacao)
        format_paragraph(obrigacao_c, 3, 0,1.385827, 18, 18, 18)

    #############################################
    # da vigência
    vigencia = document.add_paragraph(style = 'List Number')
    vigencia.add_run('DA VIGÊNCIA').bold = True
    format_paragraph(vigencia, 3, 0,0, 18,18,18)


    vig4_1 = document.add_paragraph()
    full_text = f'4.1) A prestação de serviços objeto do presente contrato vigerá até o julgamento final em sede de {cliente[3]}.'
    bold_text = ['4.1)']
    add_formatted_text(vig4_1, full_text, bold_text)
    format_paragraph(vig4_1, 3,0,0, 18, 18, 18)


    if int(cliente[18]) > 0:
        vig4_2 = document.add_paragraph()
        full_text = f'4.2) As partes concordam que têm a expectativa de duração do processo de {int(cliente[18])} meses. Caso o tempo supere esse prazo as partes se comprometem em renegociar o contrato em comum acordo.'
        bold_text = ['4.2)']
        add_formatted_text(vig4_2, full_text, bold_text)
        format_paragraph(vig4_2, 3,0,0, 18, 18, 18)

        vig4_3 = document.add_paragraph()
        full_text = f'4.3) Em caso de rescisão antecipada do contrato, permanecerão devidos os honorários aqui previstos, incluindo de êxito.'
        bold_text = ['4.3)']
        add_formatted_text(vig4_3, full_text, bold_text)
        format_paragraph(vig4_3, 3,0,0, 18, 18, 18)
    else:
        vig4_2 = document.add_paragraph()
        full_text = f'4.2) Em caso de rescisão antecipada do contrato, permanecerão devidos os honoráriosaqui previstos, incluindo de êxito.'
        bold_text = ['4.2)']
        add_formatted_text(vig4_2, full_text, bold_text)
        format_paragraph(vig4_2, 3,0,0, 18, 18, 18)


    #############################################
    #Outras informações
    outras_infor = document.add_paragraph(style = 'List Number')
    outras_infor.add_run('OUTRAS INFORMAÇÕES').bold = True
    format_paragraph(outras_infor, 3, 0,0, 18,18,18)


    outras_info5_1 = document.add_paragraph()
    full_text = '5.1) Todas as dúvidas e informações sobre a referida prestação de serviços poderão ser dirimidas, preferencialmente, por e-mail, telefone, mensagens ou WhatsApp, desde que em horário comercial (09:00h às 12:00h e 14:00h às 19:00h).'
    bold_text = ['5.1)']
    add_formatted_text(outras_info5_1, full_text, bold_text)
    format_paragraph(outras_info5_1, 3,0,0, 18, 18, 18)

    outras_info5_2 = document.add_paragraph()
    full_text = '5.2) As partes se comprometem a (i) tratar todas as informações que tiverem acesso por meio deste trabalho de forma confidencial durante o prazo de realização das atividades; e (ii) não utilizar qualquer informação confidencial para qualquer fim que não a realização dos trabalhos. Excetua-se do conceito de informação confidencial aquela que já for divulgada ou disponibilizada publicamente pelo CONTRATANTE.'
    bold_text = ['5.2)']
    add_formatted_text(outras_info5_2, full_text, bold_text)
    format_paragraph(outras_info5_2, 3,0,0, 18, 18, 18)

    outras_info5_3 = document.add_paragraph()
    full_text = '5.3) O CONTRATANTE está autorizado a substabelecer os poderes outorgados, desde com reserva de poderes.'
    bold_text = ['5.3)']
    add_formatted_text(outras_info5_3, full_text, bold_text)
    format_paragraph(outras_info5_3, 3,0,0, 18, 18, 18)

    #############################################
    #do foro
    foro = document.add_paragraph(style = 'List Number')
    foro.add_run('DO FORO').bold = True
    format_paragraph(foro, 3, 0,0, 18,18,18)


    foro6_1 = document.add_paragraph()
    full_text = '6.1) Fica eleito o foro de Brasília (DF), com a exclusão de qualquer outro, ainda que privilegiado, para dirimir quaisquer litígios referentes ao cumprimento das obrigações ora assumidas.'
    bold_text = ['6.1)']
    add_formatted_text(foro6_1, full_text, bold_text)
    format_paragraph(foro6_1, 3,0,0, 18, 18, 18)

    #############################################
    #Data, assinantes
    #data
    data = datetime.now()
    paragraph_date = document.add_paragraph(f'Brasília/DF, {data_extenso(data)}')
    paragraph_format = paragraph_date.paragraph_format
    paragraph_format.alignment = 1  # centralizado?
    paragraph_format.space_before = Pt(32)
    #paragraph_date = WD_ALIGN_PARAGRAPH.CENTER

    paragraph = document.add_paragraph()
    paragraph.add_run(f'{cliente[1].upper()}\nCONTRATANTE').bold = True
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = 1  # Centralizado
    paragraph_format.space_after = Pt(64)
    paragraph_format.space_before = Pt(64)


    paragraph = document.add_paragraph()
    paragraph.add_run('ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS\nCONTRATADO').bold = True
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = 1  # Centralizado

    
    ####################################################################################################################
    st.write('')
    st.write('')
    st.write('')
    if cliente_selecionado is not None:
        st.markdown(f'#### {recuo_adicao}{title.text}')
        st.write(contratante.text)


        st.markdown(f"""
            <div style="text-align: justify;">
            <p><b>{objeto_1.text}</b></p>
            <p>{objeto_1_1.text}</p>
            <p>{objeto_1_2.text}</p>
            </div>
            """, unsafe_allow_html=True)

        # Loop sobre os itens na lista e realizar as operações necessárias
        for item in itens_lista:
            st.markdown(f"""
            <div style="text-align: justify;">
            <p>{recuo} - {item.strip()}</p> 
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown(f"""
            <div style="text-align: justify;">
            <p><i>texto padrão  sobre serviços, equipe, riscos</i></p>
            <p><b>{pagamento.text}</b></p>
            <p>{pag_2_1.text} </p>
            <p>{recuo}{valor_prolabore_inicial.text}</p>
            <p>{recuo}{pag2_1_b.text}</p>
            <p>{recuo}{pag2_1_c.text}</p>
            <p><b>{obrigacoes.text}</b></p>
            <p><i>texto padrão</i><p>
            <p><b>{vigencia.text}</b></p>
            <p>{vig4_1.text}</p>
            </div>
            """, unsafe_allow_html=True)
        
        if int(cliente[18]) > 0:
            st.markdown(f"""
            <div style="text-align: justify;">
            <p>{vig4_2.text} </p>
            <p>{vig4_3.text} </p>
            </div>
            """, unsafe_allow_html=True)
        else:
                    st.markdown(f"""
            <div style="text-align: justify;">
            <p>{vig4_2.text} </p>
            </div>
            """, unsafe_allow_html=True)

        st.markdown(f"""
        <div style="text-align: justify;">
        <p><b>{outras_infor.text}</b> </p>
        <p><i>texto padrão</i><p>
        <p><b>{foro.text}</b></p>
        <p><i>texto padrão</i><p>
        </div>
        """, unsafe_allow_html=True)

        st.write("")
        st.write("")
        st.write("")

        # Salvar documento em arquivo temporário e permitir download
        with NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            document.save(tmp_file.name)
            st.download_button(
                label="Baixar Documento",
                data=open(tmp_file.name, 'rb').read(),
                file_name=f'contrato_contencioso_{cliente[1]}.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
