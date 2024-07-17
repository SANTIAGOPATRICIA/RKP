import pandas as pd
import locale
import streamlit as st
from streamlit_gsheets import GSheetsConnection
from docx import Document
from datetime import datetime
import calendar
from num2words import num2words
from st_pages import Page, Section, show_pages, add_page_title, add_indentation



# Define o local para português do Brasil
locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')

# Expande a largura da tela
st.set_page_config(layout="wide")
add_indentation() 

########################################
#Funções
#Tratamento do valor por extenso
def num_extenso(valor:str):
    """
    Escreve por extenso, os valores monetários

    valor(str) = valor em reais
    """
    # Separa o valor em parte inteira e decimal
    parte_inteira, parte_decimal = str(valor).split('.')
    
    # Processa a parte inteira por extenso
    por_extenso_inteira = num2words(int(parte_inteira), lang='pt_BR')

    # Se houver centavos, processa-os por extenso
    if parte_decimal and parte_decimal != '0':
        por_extenso_decimal = num2words(int(parte_decimal), lang='pt_BR')
        # Formatação específica para valores monetários
        por_extenso = f'{por_extenso_inteira} reais e {por_extenso_decimal} centavos'
    else:
        por_extenso = f'{por_extenso_inteira} reais'
    
    return por_extenso

#Tratamento data por extenso
def data_extenso(date:str):
    """
    Escreve a data por extenso
    date(date) = data
    """
    # Obtém o nome do mês por extenso
    mes_por_extenso = calendar.month_name[date.month]
    # Formata a data por extenso
    data_por_extenso = f"{date.day} de {mes_por_extenso} de {date.year}"
    return data_por_extenso

#########################################
#Lista dos TJs
TJs = [
    'Tribunal de Justiça do Acre (TJAC)','Tribunal de Justiça de Alagoas (TJAL)',
    'Tribunal de Justiça do Amapá (TJAP)', 'Tribunal de Justiça do Amazonas (TJAM)',
    'Tribunal de Justiça da Bahia (TJBA)', 'Tribunal de Justiça do Ceará (TJCE)',
    'Tribunal de Justiça do Distrito Federal e Territórios (TJDFT)', 'Tribunal de Justiça do Espírito Santo (TJES)',
    'Tribunal de Justiça de Goiás (TJGO)', 'Tribunal de Justiça do Maranhão (TJMA)',
    'Tribunal de Justiça de Mato Grosso (TJMT)', 'Tribunal de Justiça de Mato Grosso do Sul (TJMS)',
    'Tribunal de Justiça de Minas Gerais (TJMG)', 'Tribunal de Justiça do Pará (TJPA)',
    'Tribunal de Justiça da Paraíba (TJPB)', 'Tribunal de Justiça do Paraná (TJPR)',
    'Tribunal de Justiça de Pernambuco (TJPE)', 'Tribunal de Justiça do Piauí (TJPI)',
    'Tribunal de Justiça do Rio de Janeiro (TJRJ)', 'Tribunal de Justiça do Rio Grande do Norte (TJRN)',
    'Tribunal de Justiça do Rio Grande do Sul (TJRS)', 'Tribunal de Justiça de Rondônia (TJRO)',
    'Tribunal de Justiça de Roraima (TJRR)', 'Tribunal de Justiça de Santa Catarina (TJSC)',
    'Tribunal de Justiça de São Paulo (TJSP)', 'Tribunal de Justiça de Sergipe (TJSE)',
    'Tribunal de Justiça do Tocantins (TJTO)'
    ]
#########################################

#dividindo a pagina em duas colunas
dados, desenvolvimento = st.columns([2,3])
# dados.metric(label='dados', value=123)
# desenvolvimento.metric(label='desenvolvimento', value=456)

#abrindo o documento
documento = Document(r'docx\Modelo-Contrato-Segunda-PF-Usucapiao.docx')
with dados:

    # st.title('Contrato, Segunda, PF e usucapião')
    st.markdown('**Informação do contratante**')
    contratante_ = st.text_input(label='Contratante', placeholder = None)
    nacionalidade_ = st.text_input(label='Nacionalidade', placeholder = None)
    estado_civil_ = st.text_input(label="Estado civil", placeholder = None)
    profissao_ = st.text_input(label="Profissão", placeholder = None)
    rg_ = st.text_input(label="RG com orgão emissor", placeholder = None)
    cpf_ = st.text_input(label="CPF", placeholder = None)
    email_ = st.text_input(label="e-mail", placeholder = None)
    endereco_ = st.text_input(label="Endereço", placeholder = None)
    municipio_uf_ = st.text_input(label="Municipio/UF", placeholder = None)
    #Dados para o documento
    st.divider()
    st.markdown('**Conteúdo do documento**')
    objeto_ = st.text_area(label="Objeto", )
    instancia_ = st.selectbox("Instância", options=['primeira instância', 'segunda instância', 'tribunal superior'], index = None)
    if instancia_ == 'segunda instância':
        orgao_ = st.selectbox("Tribunal", options = TJs, index = None)
    elif instancia_ == 'tribunal superior':
        orgao_ = st.selectbox("Tribunal", options = ['STJ', "STF"], index = None)
    else:
        orgao_ = st.text_input(label = "Vara ou seção", placeholder = None)
    proLabore_inicial_ =  st.number_input("Pró-labore inicial", format = "%.2f", value = 0.00, step = 1000.00)
    manutencao_isencao_ = st.selectbox("Honorário de manutenção - tempo de isenção em meses", options = ['0','3','6','12','18','24','36'])
    manutencao_salario_ = st.selectbox("equivalencia salarial", options = ['0,5', '1', '1,5'], index = None)
    exito_valor_ = st.number_input("Êxito", format = "%.2f", value = 0.00, step = 1000.00)
    data_documento = st.date_input(label='Data do documento', value='default_value_today', format='DD/MM/YYYY')

    #Armazenar os dados do usuario em dicionário
    references = {}
    references = {
    "<CONTRATANTE>": contratante_,
    "<NACIONALIDADE>": nacionalidade_,
    "<ESTADO_CIVIL>": estado_civil_,
    "<PROFISSAO>": profissao_,
    "<RG>": rg_,
    "<CPF>": cpf_,
    "<EMAIL>": email_,
    "<ENDERECO>": endereco_,
    "<CIDADE_UF>": municipio_uf_,
    "<OBJETO>": objeto_,
    "<INSTANCIA>": instancia_,
    "<ORGAO>": orgao_,
    "<VALOR>": float(proLabore_inicial_),
    "<VALOR_EXTENSO>": num_extenso(proLabore_inicial_),
    "<QTDE_MESES>": manutencao_isencao_,
    "<QDE_SM>": manutencao_salario_,
    "<EXITO_VALOR_MIN>": float(exito_valor_),
    "<EXITO_MIN_EXTENSO>": num_extenso(exito_valor_),
    "<DATA>": data_extenso(data_documento)
    }

    
    if st.button('Visualizar'):
        for paragrafo in documento.paragraphs:
            for codigo, valor in references.items():
                # Converte valores numéricos em strings antes de substituir
                if isinstance(valor, (int, float)):
                    valor = str(valor)
                paragrafo.text = paragrafo.text.replace(codigo, valor)
      
with desenvolvimento: 
    for paragraph in documento.paragraphs:
        st.write(paragraph.text)

    if st.button('Salvar docx'):
        documento.save(f'docx\CONTRATO -{contratante_}.docx')




    # conn = st.connection("gsheets", type=GSheetsConnection)
    # existing_data = conn.read(worksheet="cliente", ttls=5)
    # st.dataframe(existing_data)

# Lista do sidebar para o tipo de documento
# tipo_documento = st.sidebar.selectbox("Tipo de documento a gerar", options=['Contrato', 'Proposta', 'Aditivo'])

# if tipo_documento == "Contrato":
#     # Opção de instância para contratos
#     vigencia = st.sidebar.selectbox("Instância", options=['Primeira Instância', 'Segunda Instância', 'Terceira Instância', 'Corte Superior'])

# # Tipo de pessoa
# pessoa = st.sidebar.selectbox("Tipo de Pessoa", options=["Física", "Jurídica"])

# # Assunto do objeto
# objeto = st.sidebar.selectbox("Assunto do objeto", options=['Divórcio', 'Inventário', 'Pensão alimentícia', 'Usucapião'])


# # Botão de submissão na barra lateral
# submit_button = st.sidebar.button('Iniciar documento')

# # Variável para controlar se os campos de entrada devem ser exibidos
# show_inputs = True

# # Lista das possibilidades
# pf_contrato_primeira_divorcio = ['Contrato', 'Segunda Instância', 'Física', 'Usucapião']

# # Verifica se o botão de submissão foi clicado
# if submit_button:
#     input_usuario = [tipo_documento, vigencia if tipo_documento == 'Contrato' else None, pessoa, objeto]
    
#     # Verifica se o input do usuário corresponde a uma das possibilidades
#     if all(pf_contrato_primeira_divorcio == input_usuario for possibilidade, resposta in zip(pf_contrato_primeira_divorcio, input_usuario)):
#         # Renderiza os campos de entrada de dados
#         st.title('Contrato, Segunda, PF e usucapião')
#         st.markdown('1. Carregar o arquivo')
#         st.markdown('2. Colher as informações dos clientes')
#         contratante_ = st.text_input(label='Contratante', placeholder = None)
#         nacionalidade_ = st.text_input(label='Nacionalidade', placeholder = None)
#         estado_civil_ = st.text_input(label="Estado civil", placeholder = None)
#         profissao_ = st.text_input(label="Profissão", placeholder = None)
#         rg_ = st.text_input(label="RG com orgão emissor", placeholder = None)
#         cpf_ = st.text_input(label="CPF", placeholder = None)
#         email_ = st.text_input(label="e-mail", placeholder = None)
#         endereco_ = st.text_input(label="Endereço", placeholder = None)
#         municipio_uf_ = st.text_input(label="Municipio/UF", placeholder = None)
#         objeto_ = st.text_area(label="Objeto", placeholder = None)
#         orgao_ = st.text_input(label="Nome órgão", placeholder = None)
#         proLabore_inicial_ =  st.number_input("Pró-labore incial:", format="%.2f", value=0.00, step=1000.00)
        
#         st.markdown('3. Mostrar o documento na tela')
#         st.markdown('4. Salvar documento em docx')
        
#         # Defina show_inputs como False para evitar que os campos de entrada sejam exibidos novamente
#         show_inputs = False
#     else:
#         st.write('Page under construction. 😀😀😀')

# Renderiza os campos de entrada de dados se show_inputs for True
# if show_inputs:
#     st.title('Bem-vindo!')
#     st.write('Clique no botão na barra lateral para iniciar o documento.')

# st.title('Criar Contratos/Propostas')
# st.markdown('*MVP* - Criar Contratos - PF/Pensão alimentícia')

# conn = st.connection("gsheets", type=GSheetsConnection)
# existing_data = conn.read(worksheet="cliente", ttls=5)

# #Lista do dropbox
# tipo_pessoa = ["Física", "Jurídica"]
# tipo_documento = ['Contrato','Proposta', 'Aditivo']
# instancia = ['Primeira Instância', 'Segunda Instância', 'Terceira Instância', 'Corte Superior']
# objeto_materia = ['Acompanhamento processual','Divórcio', 'Inventário', 'Pensão alimentícia', 'Usucapião']

# #Iniciando o filtro do tipo de documento a gerar

# # input do usuario
# input_usuario = []

# #Tipo de documento 
# tipo_documento = st.selectbox("Tipo de documento a gerar", options=tipo_documento, index=None, placeholder="Escolha uma opção")
# # input do usuario
# input_usuario.append(tipo_documento)

# if tipo_documento == "Contrato":
#     vigencia = st.selectbox("Instância", options=instancia, index=None, placeholder="Escolha uma opção")
#     input_usuario.append(vigencia)
# #Natureza da pessoa
# pessoa = st.selectbox("Tipo de Pessoa", options=tipo_pessoa, index=None, placeholder="Escolha uma opção")
# input_usuario.append(pessoa)    
# assunto_objeto = st.selectbox("Assunto do objeto", options=objeto_materia, index=None, placeholder="Escolha uma opção")
# input_usuario.append(assunto_objeto)
# # data_documento = st.date_input(label='Data do documento', value='default_value_today', format='DD/MM/YYYY')
# submit_button = st.button('Iniciar documento')


# #Lista das possibilidades 
# pf_contrato_primeira_divorcio = ['Contrato', 'Primeira Instância', 'Física', 'Divórcio']

# if submit_button:
#     if all(pf_contrato_primeira_divorcio == input_usuario for possibilidade, resposta in zip(pf_contrato_primeira_divorcio, input_usuario)):
#         st.write('Parabéns, você acertou todas as questões!')
#         st.experimental_rerun()
#     else:
#         st.write('Você errou em pelo menos uma questão.')





# #Quantidade de clientes para inputar nome e os respectivos dados
# qde_clientes = st.number_input(label='Número de contratantes', min_value=1)

# clientes_data = []  # Lista para armazenar os dados dos clientes

# for x in range(qde_clientes):
#     chave_cliente = f'cliente_{x}'
#     chave_cpf = f'cpf_{x}'
#     chave_endereco = f'endereco_{x}'
    
#     if qde_clientes > 1:
#         cliente = st.text_input(label=f"Cliente {x+1}", key=chave_cliente)
#         documentoCPF = st.text_input(label=f"CPF {x+1}", key=chave_cpf)
#         endereco = st.text_input(label=f"Endereço {x+1}", key=chave_endereco)
#     else:
#         cliente = st.text_input(label=f"Cliente", key=chave_cliente)
#         documentoCPF = st.text_input(label=f"CPF", key=chave_cpf)
#         endereco = st.text_input(label=f"Endereço", key=chave_endereco)
    
#     # Salva os dados do cliente atual em um dicionário
#     cliente_data = {
#         'Cliente': cliente,
#         'CPF': documentoCPF,
#         'Endereço': endereco
#     }
    
#     # Adiciona os dados do cliente à lista
#     clientes_data.append(cliente_data)

# if pessoa == 'Jurídica':
#     representante = st.text_input(label='Nome do Representante', placeholder = None)
#     representanteCPF = st.text_input(label='CPF do representante', placeholder = None)
#     representanteEmail = st.text_input(label='e-mail do representante', placeholder = None)

# if submit_button:
#     novo_doc = pd.DataFrame(clientes_data)  # Converte a lista de dicionários em um DataFrame
#     st.dataframe(novo_doc)




