import pandas as pd
import streamlit as st
from streamlit_gsheets import GSheetsConnection
from st_pages import add_indentation
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime
import locale
import time
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from tempfile import NamedTemporaryFile
from utils.funcoes import format_paragraph, add_formatted_text, format_title_centered, \
    format_title_justified, num_extenso, data_extenso, fonte_name_and_size, add_section,\
    create_paragraph, atualizar_base_dados, num_extenso_percentual, set_table_borders #, create_list_number_style, create_numbered_list_style

# st.set_page_config(layout="wide")

add_indentation()

# Define o local para português do Brasil
try:
    locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
except locale.Error as e:
    print(f"Erro ao definir a localidade: {e}")



# Inicializar DataFrame vazio
df_inputs = pd.DataFrame(columns=[
    'objeto_consultivo', 
    'total-de-horas', 
    'valor-aplicado', 
    'valor-formatado',
    'subtotal',
    'objeto_contencioso',
    'pro_labore_inicial',
    'pro_labore_manutencao',
    'exito',
    'contencioso_desconto',
    'contencioso_valor_final'
])

# Store the initial value of widgets in session state
if "visibility" not in st.session_state:
    st.session_state.visibility = "visible"
    st.session_state.disabled = False
    st.session_state.horizontal = False

#####################################################################################

lista_numerada = ['a)', 'b)', 'c)', 'd)', 'e)', 'f)', 'g)', 'h)', 'i)', 'j)', 'k)', 'l)', 'm)', 'n)', 'o)', 'p)', 'q)', 'r)', 's)', 't)']

#####################################################################################

recuo = "&nbsp;" * 24


dados, desenvolvimento = st.columns([2, 3])

with dados:
    
    st.write('**Informação para a proposta**')

    # Carregando a lista de clientes pela primeira vez
    lista_clientes = pd.read_csv('clientes.csv')
    lista_clientes = lista_clientes.sort_values(by='Nome')['Nome'].unique().tolist()

    # Adiciona uma opção para cadastrar novo cliente
    lista_clientes.append("--Novo cliente--")
    lista_clientes = sorted(lista_clientes)

    nome_cliente = st.selectbox(
            'Cliente',
            (lista_clientes),
            index=None,
            placeholder='Selecione o cliente'
        )
    

    if nome_cliente == '--Novo cliente--':
        form = st.form('Novo Cliente')
        nome_cliente = form.text_input('Cadastrar novo cliente')
        form.form_submit_button("Cadastrar")
    lista_clientes.append(nome_cliente)

    st.divider()
    st.write("Objetos")
    input_consultivo_objeto = st.text_area(label="Objeto do consultivo:")
    input_contencioso_objeto = st.text_area(label="Objeto do contencioso:")


    #Complemento dos objetos para a descrição das atividades a serem realizadas
    st.divider()
    st.write("Atividades")
    #alterar atividade do consultivo
    alterar_atividade_consutltivo =st.radio('Alterar texto do objeto do consultivo?', ['Sim', 'Não'],
        key='atividade_consultivo',
        label_visibility=st.session_state.visibility,
        disabled=st.session_state.disabled,
        horizontal=st.session_state.horizontal,
        index=None)
    if alterar_atividade_consutltivo == 'Sim':
        atividade_consultivo_alterada = st.text_area('Alterar texto atividade consutlivo:')

    #alterar atividade do contencioso
    alterar_atividade_contencioso = st.radio('Alterar texto do objeto do contencioso?', ['Sim', 'Não'],
        key='atividade_contencioso',
        label_visibility=st.session_state.visibility,
        disabled=st.session_state.disabled,
        horizontal=st.session_state.horizontal,
        index=None)
    if alterar_atividade_contencioso == 'Sim':
        atividade_contencioso_alterada = st.text_area('Alterar texto atividade contencioso:')
    #tempo máximo para as reuniões
    tempo_max_reunioes = st.number_input(label='Tempo (em horas) máximo para reuniões:', step=10, key="tempo_max")
    #inserir mais atividades 

    # inserir_mais_atividades = st.radio('Inserir mais atividades a serem desenvolvidas?', ['Sim', 'Não'],
    #     key='adicionar_atividade_',
    #     label_visibility=st.session_state.visibility,
    #     disabled=st.session_state.disabled,
    #     horizontal=st.session_state.horizontal,
    #     index=None)


    #Valores do consultivo
    st.divider()
    st.write('Valor do consultivo')
    #hora total
    hora_total = st.number_input(label='Total de horas:', step=10, key='hora_total_')
    hora_total = int(hora_total)
    #valor aplicado
    valor_aplicado = st.selectbox("Valor aplicado",
                            (1150.00, 850.00, 680.00, 580.00, 490.00, 290.00), key='valor_aplicado_')
    # Arredondar o valor para duas casas decimais e formatar como string
    valor_formatado = "{:.2f}".format(round(valor_aplicado, 2))
    #valor por extenso
    valor_por_extenso = num_extenso(valor_formatado)
    
    #valor do subtotal
    valor_total = (hora_total * valor_aplicado)
    valor_total_formatado = "{:.2f}".format(round(valor_total, 2))
    subtotal_extenso = num_extenso(valor_total_formatado)
    st.write(f'*Subtotal R${valor_total}*')

    #desconto
    desconto_percentual_consultivo = st.number_input('Percentual do desconto do consultivo (%)', min_value=0.0, max_value=100.0, key='consultivo_desc')
    desconto_percentual_consultivo = float("{:.2f}".format(desconto_percentual_consultivo))
    desconto_percentual_formatado = "{:.2f}".format(round(desconto_percentual_consultivo, 2))

    if desconto_percentual_consultivo > 0:
        total_final = valor_total*((100.00-desconto_percentual_consultivo)/100)
        total_final_formatado = "{:.2f}".format(round(total_final, 2))
        total_final_extenso = num_extenso(total_final_formatado)
        st.write(f'*O valor da proposta consultiva é de R${total_final}*')
    else:
        st.write(f'*O valor da proposta consultiva é de R${valor_total_formatado}*')
    
    
    st.divider()
    # Valores do contencioso
    st.write('Valor do contencioso')

    # Input para o valor do pro-labore inicial
    prolabore_inicial = st.number_input('Valor do pro-labore inicial (R$)', key='prolabore_inicial')
    prolabore_inicial_formatado = "{:.2f}".format(round(prolabore_inicial, 2))
    # Input para o percentual do desconto do contencioso
    desconto_contencioso = st.number_input('Percentual do desconto do contencioso (%)', min_value=0.0, max_value=100.0, step=10.0, key='contencioso_desc')
    #valor contencioso com desconto
    #iniciando o valor com desconto como sendo 0.0
    prolabore_inicial_com_desconto = 0.0

    if desconto_contencioso> 0.0:
        prolabore_inicial_com_desconto = prolabore_inicial*((100.00-desconto_contencioso)/100)
        st.write(f'*O valor do pró-labore inicial com desconto é de R$ {prolabore_inicial_com_desconto}*')
    else:
        st.write(f'*O valor do pró-laobre inicial é de R${prolabore_inicial_formatado}.*')
    
    # Selectbox para o tempo de isenção
    prolabore_manutencao = st.selectbox("Tempo de isenção (meses)",
                                        (0, 6, 12, 18, 24, 30, 36, 42, 48, 56, 60),
                                        key='tempo_isencao')

    # Condicional para exibir o selectbox do pró-labore de manutenção
    prolabore_manutencao_valor = 0.0
    if prolabore_manutencao > 0:
        prolabore_manutencao_valor = st.selectbox("Proporção do salário mínimo para o pró-labore de manutenção",
                                                (0.5, 1, 1.5, 2, 2.5),
                                                key='valor_manutencao')

    # Input para o percentual do benefício econômico
    exito_percentual = st.number_input('Percentual do benefício econômico (%)', min_value=0.0, max_value=100.0, step=0.5, key='exito_percentual')
    exito_percentual_formatado = "{:.2f}".format(round(exito_percentual, 2))
    # Input para o valor teto do êxito
    valor_teto_exito = st.number_input('Valor teto do êxito (R$)', key='valor_teto_exito')
    valor_teto_exito_formatado = "{:.2f}".format(round(valor_teto_exito, 2))



# #####################################################################################
# Abrir documento com papel timbrado da RKP
document = Document(r"docx/RKP-PapelTimbrado.docx")

# Definir fonte e tamanho do documento
fonte_name_and_size(document, 'Arial', 12)


# Adicionar uma sessão ao documento
section = add_section(document, 5,3,2,2)


# Data
data = datetime.now()
paragraph_date = document.add_paragraph(f'Brasília/DF, {data_extenso(data)}')
paragraph_format = paragraph_date.paragraph_format
paragraph_format.alignment = 2  # a direita


#############################################################################
# Adicionar título ao doc
title = document.add_heading('PROPOSTA PARA PRESTAÇÃO \nDE SERVIÇOS ADVOCATÍCIOS')
format_title_centered(title)
space_format = title.paragraph_format
space_format.space_before = Pt(48)


p_de = document.add_paragraph()
p_de.add_run('DE: ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS S/S').bold = True
p_de_format = p_de.paragraph_format
p_de_format.line_spacing = Pt(18)
p_de_format.space_before = Pt(148)
p_de_format.space_after = Pt(8)


paragraph_para = document.add_paragraph()
paragraph_para.add_run(f'PARA: {nome_cliente} - Interessado(a)').bold = True
paragraph_format = paragraph_para.paragraph_format
paragraph_format.line_spacing = Pt(18)
paragraph_format.space_after = Pt(48)

paragraph_ref = document.add_paragraph()
paragraph_ref.text = 'Referência: PROPOSTA DE HONORÁRIOS ADVOCATÍCIOS'
paragraph_format = paragraph_ref.paragraph_format
paragraph_format.space_before = Pt(18)
paragraph_format.space_after = Pt(96)
paragraph_format.line_spacing = Pt(18)



paragraph = document.add_paragraph()
format_paragraph(paragraph, 3, 1.5748, 0, 48,16,20)
full_text= "ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS S/S, com sede no SIG - Quadra 01, Lote 495, Edifício Barão do Rio Branco, sala 244, Brasília-DF, CEP 70.610-410, telefones 3321-7043 e 3226-0137, inscrita no CNPJ sob o nº 03.899.920/0001- 81, registro na Ordem dos Advogados do Brasil – OAB/DF sob o número 616/00 – RS, endereço eletrônico www.khouriadvocacia.com.br, vem, mui respeitosamente, apresentar PROPOSTA DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS, nas condições a seguir."
# Texto que será negritado
bold_text = "ROQUE KHOURI & PINHEIRO ADVOGADOS ASSOCIADOS S/S"
# Adicionar o texto formatado ao parágrafo
add_formatted_text(paragraph, full_text, bold_text)

# I - DOS SERVIÇOS A SEREM DESENVOLVIDOS
title_one = document.add_heading('I - DOS SERVIÇOS A SEREM DESENVOLVIDOS', level=2)
format_title_justified(title_one)

#itens atuação
itens_atuacao = [
    "Providências preliminares de levantamento e análise de todas as informações e documentos relativos ao objeto da presente proposta, a fim de propiciar o embasamento jurídico necessário;",
    f"Atuação consultiva: {input_consultivo_objeto};",
    f"Atuação contenciosa: {input_contencioso_objeto};",
    "Participações em reuniões com V.Sa. e demais profissionais envolvidos, objetivando a negociação entre as partes; ",
    "Atuação contenciosa com a confecção de petição inicial, ajuizamento de ação e acompanhamento de processo judicial até o seu julgamento final em sede de 2ª Instância;",
    "Elaboração de todas as petições e recursos necessários (incluindo memoriais) ao acompanhamento da ação judicial até o seu julgamento final em sede de 2ª Instância;",
    "Diligências pessoais junto aos Tribunais, em especial despachos presenciais e telepresenciais com os Juízes, Desembargadores e Ministros responsáveis pelo julgamento, se cabível. "
]

if alterar_atividade_consutltivo == 'Sim':
    itens_atuacao[1] = atividade_consultivo_alterada

if alterar_atividade_contencioso == "Sim":
    itens_atuacao[2] = atividade_contencioso_alterada

if tempo_max_reunioes > 0:
    itens_atuacao[3] = f'Participações em reuniões com V.Sa. e demais profissionais envolvidos, objetivando a negociação entre as partes, limitada a {tempo_max_reunioes} horas;'

if input_contencioso_objeto:
    paragrah_padrao = document.add_paragraph(f"Conforme solicitação, apresentamos proposta de honorários para atuação consultiva, referente à {input_consultivo_objeto}, ou caso não tenha êxito, atuação judicial contenciosa {input_contencioso_objeto}")
    format_paragraph(paragrah_padrao,3, 1.5748, 0,18,18,18)

    paragraph_atividades = document.add_paragraph('A atuação desse Jurídico compreenderá as seguintes atividades:')
    format_paragraph(paragraph_atividades, 3, 1.5748,0,18,18,18)
        
    for i in range(len(itens_atuacao)):
        paragraph_itens_atuacao = document.add_paragraph(f'{lista_numerada[i]} {itens_atuacao[i]}')
        format_paragraph(paragraph_itens_atuacao, 3, 0,1.77165, 18, 18, 18)


else:   
    paragrah_padrao = document.add_paragraph(f"Conforme solicitação, apresentamos proposta de honorários para atuação consultiva, referente à {input_consultivo_objeto}, ou caso não tenha êxito, atuação judicial contenciosa, em defesa dos interesses de {nome_cliente}")
    format_paragraph(paragrah_padrao,3, 1.5748, 0, 18,18,18)
    paragraph_atividades = document.add_paragraph('A atuação desse Jurídico compreenderá as seguintes atividades:')
    format_paragraph(paragraph_atividades, 3, 1.5748,0,18,18,18)
    
    # del itens_atuacao[2]
    itens_atuacao[2] = 'Atuação contenciosa no acompanhamento  judicial até o julgamento final em 2ª Instância'
    
    for i in range(len(itens_atuacao)):
        paragraph_itens_atuacao = document.add_paragraph(f'{lista_numerada[i]} {itens_atuacao[i]}')
        format_paragraph(paragraph_itens_atuacao, 3, 0,1.77165, 18, 18, 18)


#paragrafo I-IV
paragraph_four = document.add_paragraph()
format_paragraph(paragraph_four,3, 1.5748,0, 18,18,18)
paragraph_four.text ='Para o cumprimento dos serviços, o escritório disponibilizará sua equipe técnica, sendo que haverá advogado responsável pelo acompanhamento direto da demanda.'

#paragrafo I-V
paragraph_five = document.add_paragraph()
format_paragraph(paragraph_five,3, 1.5748,0, 18,18,18)
paragraph_five.text ='A Roque Khouri Pinheiro Advogados Associados alerta que qualquer ação judicial implica em risco, estando o interessado ciente principalmente da possibilidade de condenação em honorários advocatícios, conforme previsto no Código de Processo Civil (10% a 20% sobre o valor da causa atualizado), principalmente diante das especificidades da fase atual em que o processo se encontra. Alerta também se tratar de análise e de confecção de peças a serem elaboradas com base no direito, jurisprudência atual e principalmente das informações e documentos que serão sempre fornecidos pelo Interessado.'

#############################################################################
# II - DA POLÍTICA GERAL DE VALORES - HONORÁRIOS
#titulo II
title_two = document.add_heading('II - DA POLÍTICA GERAL DE VALORES - HONORÁRIOS', level=2)
format_title_justified(title_two)
#Paragrafo II-I
paragraph_two_one = document.add_paragraph()
format_paragraph(paragraph_two_one,3, 1.5748, 0,18,18,18)
paragraph_two_one.text = "Faz parte integrante de todas as nossas propostas de honorários os itens abaixo, componentes da nossa Política de Honorários de consultoria:" #\nTaxas horárias de honorários para projetos. Para projetos, nós cobramos valores de honorários de acordo com as seguintes taxas horárias:

#Paragrafo II-II
paragraph_two_two = document.add_paragraph()
format_paragraph(paragraph_two_two,3, 1.5748, 0,18,18,18)
full_text= "Taxas horárias de honorários para projetos. Para projetos, nós cobramos valores de honorários de acordo com as seguintes taxas horárias:"
# Texto que será negritado
bold_text = "Taxas horárias de honorários para projetos."
# Adicionar o texto formatado ao parágrafo
add_formatted_text(paragraph_two_two, full_text, bold_text)

# Adicionar tabela
table = document.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style =  None #'LightShading-Accent3'
# Definir bordas da tabela
set_table_borders(table)
# get table data -------------
items = (
    ('Sócio Majoritário - Dr. Paulo Roque', 'R$1.150,00'),
    ('Sócia Nominal - Dra. Ângela Pinheiro', 'R$850,00'),
    ('Advogado Sênior', 'R$680,00'),
    ('Advogado Pleno', 'R$580,00'),
    ('Advogado Júnior', 'R$490,00'),
    ('Paralegal/Estagiário', 'R$290,00')
)

# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].text = 'Profissional'
heading_cells[1].text = 'Valor'


# Alinhar texto na primeira linha ao centro verticalmente
for cell in heading_cells:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Definir cor de fundo para a primeira linha
shading_elm_0 = OxmlElement('w:shd')
shading_elm_0.set(qn('w:fill'), 'B4FF00')  # RGB (180, 255, 0)
heading_cells[0]._element.get_or_add_tcPr().append(shading_elm_0)
shading_elm_1 = OxmlElement('w:shd')
shading_elm_1.set(qn('w:fill'), 'B4FF00')  # RGB (180, 255, 0)
heading_cells[1]._element.get_or_add_tcPr().append(shading_elm_1)


# add a data row for each item
for item in items:
    cells = table.add_row().cells
    cells[0].text = item[0]
    cells[1].text = item[1]
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Alinhar texto à direita na segunda coluna

# Definir bordas da tabela
set_table_borders(table)


#Paragrafo II-III
paragraph_two_three = document.add_paragraph()
format_paragraph(paragraph_two_three, 3, 1.5748, 0,18,18,18)
full_text= "Reembolso de Despesas. As despesas incorridas no desenvolvimento dos trabalhos, como, por exemplo, despesas com ligações telefônicas, correios, couriers e outros meios de envio de documentos, com impressão de cópias e digitalização de documentos, com taxas governamentais, com viagens, táxis e outros deslocamentos, e, se aplicável, despesas com custas processuais e outras despesas relativas a processos arbitrais, judiciais e administrativos, e honorários de advogados correspondentes, serão reembolsadas, mediante a apresentação de planilha discriminada, e, se solicitado, dos respectivos comprovantes. Nenhuma despesa superior a R$ 1.000,00 (um mil Reais) será incorrida sem sua prévia aprovação por escrito."
# Texto que será negritado
bold_text = "Reembolso de Despesas."
# Adicionar o texto formatado ao parágrafo
add_formatted_text(paragraph_two_three, full_text, bold_text)


#Paragrafo II-IV
paragraph_two_four = document.add_paragraph()
format_paragraph(paragraph_two_four, 3, 1.5748,0, 18,18,18)
full_text= "Interrupção ou Suspensão dos Trabalhos. Se por qualquer motivo os trabalhos forem eventualmente interrompidos ou suspensos, faremos o levantamento das horas trabalhadas e o valor de honorários pagos até então e, se houver saldo a ser pago, faremos o faturamento correspondente. Caso haja honorários a serem restituídos, a restituição será feita no mês seguinte ao da interrupção ou suspensão dos trabalhos e do valor a ser restituído serão descontados os tributos correspondentes pagos ou a pagar."
# Texto que será negritado
bold_text = "Interrupção ou Suspensão dos Trabalhos."
# Adicionar o texto formatado ao parágrafo
add_formatted_text(paragraph_two_four, full_text, bold_text)
#############################################################################

# III - DOS HONORÁRIOS ESPECÍFICOS
#titulo III
title_three = document.add_heading('III - DOS HONORÁRIOS ESPECÍFICOS', level=2)
format_title_justified(title_three)

#paragrafo III-I
paragraph_three_one = document.add_paragraph('Com o intuito de manter a proporcionalidade entre prestação de serviços e pagamento, os honorários advocatícios devidos em consequência da presente prestação de serviços seriam cobrados por meio do sistema de horas, ou seja, cada ato praticado, esse Jurídico seria remunerado de acordo com o tempo necessário para praticá-lo.')
format_paragraph(paragraph_three_one, 3, 1.5748, 0,18,18,18)

paragraph_three_one_one = document.add_paragraph('Os honorários advocatícios devidos em consequência da prestação de serviços previstas no item I seriam assim determinados:')
format_paragraph(paragraph_three_one_one, 3, 1.5748, 0,18,18,18)

#valor atuação consultivo
valor_atuacao_consultivo = document.add_paragraph(f"a) {input_consultivo_objeto}")
format_paragraph(valor_atuacao_consultivo, 3,0, 1.5748,18,18,18)
valor_atuacao_consultivo_hora = document.add_paragraph(f'{hora_total}h estimada para a confecção e revisão;')
format_paragraph(valor_atuacao_consultivo_hora, 3, 0, 1.5748, 18,18,18)
valor_consultivo_valor_aplicado = document.add_paragraph(f"Valor da hora aplicada: R${valor_formatado} ({num_extenso(valor_formatado)});")
format_paragraph(valor_consultivo_valor_aplicado, 3, 0, 1.5748, 18,18,18)
valor_consultivo = document.add_paragraph(f'Valor total: R${valor_total_formatado} ({subtotal_extenso})')
format_paragraph(valor_consultivo, 3, 0, 1.5748, 18,18,18)

#desconto do consultivo
paragraph_desconto_consultivo = document.add_paragraph()
if desconto_percentual_consultivo > 0.0:
    paragraph_desconto_consultivo.add_run("DESCONTO").bold = True
    paragraph_desconto_consultivo.add_run(f': Tendo em vista a parceria para com o cliente, a Roque Khouri & Pinheiro, por mera liberalidade e apenas no trabalho específico, concede o desconto de {desconto_percentual_formatado}% ({num_extenso_percentual(desconto_percentual_formatado)}) em todos os valores descritos, totalizando assim, R${total_final_formatado} ({num_extenso(total_final_formatado)}) pela prestação de serviços contratados. Os honorários
previstos nos itens B2 e B3 serão devidos normalmente.')
    format_paragraph(paragraph_desconto_consultivo, 3, 1.5748, 0, 18,18,18)


#valor atuação contencioso
valor_atuacao_contencioso = document.add_paragraph('b) Atuação judicial contenciosa (caso seja necessário)')
format_paragraph(valor_atuacao_contencioso, 3, 0, 1.5748, 18,18,18)
valor_prolabore_inicial = document.add_paragraph(f'Pró-labore inicial: R${prolabore_inicial} ({num_extenso(prolabore_inicial_formatado)});')
format_paragraph(valor_prolabore_inicial, 3, 0,1.5748, 18,18,18)
valor_honorario_manutencao = document.add_paragraph(f"Honorário de manutenção: Isento durante {prolabore_manutencao} meses. Após este período, se o processo perdurar, será devido o valor de {prolabore_manutencao_valor} salário mínimo mensal;")
format_paragraph(valor_honorario_manutencao, 3, 0,1.5748, 18,18,18)
valor_honorario_exito = document.add_paragraph(f'Honorários de Êxito: {exito_percentual_formatado}% ({num_extenso_percentual(exito_percentual_formatado)}) do benefício econômico aferido ao final do processo. Fica compreendido como benefício econômico todo e qualquer valor que a INTERESSADA receber em razão da propositura da ação ou valor que deixar de pagar.')
# valor_honorario_exito.add_footnote('Fica compreendido como benefício econômico todo e qualquer valor que a INTERESSADA receber em razão da propositura da ação ou valor que deixar de pagar.') # add a footnote
format_paragraph(valor_honorario_exito, 3, 0,1.5748, 18,18,18)

#desconto do contencioso
paragraph_desconto_contencioso = document.add_paragraph()
if desconto_percentual_consultivo > 0.0:
    paragraph_desconto_contencioso.add_run("DESCONTO").bold = True
    paragraph_desconto_contencioso.add_run(f': Tendo em vista a parceria para com o cliente, a Roque Khouri & Pinheiro, por mera liberalidade e apenas no trabalho específico, concede o desconto de {desconto_percentual_formatado}% ({num_extenso_percentual(desconto_percentual_formatado)}) apenas no valor referente ao pró-labore inicial, totalizando assim, R${total_final_formatado} ({num_extenso(total_final_formatado)}) pela prestação de serviços judiciais contenciosos contratados.')
    format_paragraph(paragraph_desconto_contencioso, 3, 1.5748, 0,18,18,18)


#paragrafo III-IV
paragraph_three_four = document.add_paragraph('Não estão incluídos na proposta ora apresentada eventuais custos com a contratação de advogados correspondentes fora de Brasília, bem como as despesas a serem incorridas em virtude da execução dos serviços, tais como, cópias reprográficas, custas judiciais, honorários periciais, emolumentos com autenticação de cópias e reconhecimento de firmas, obtenção de certidões, motoboys e deslocamentos à razão de R$ 1,00/km, entre outras despesas, as quais serão pagas diretamente por V.Sa. ou reembolsadas mediante a apresentação dos respectivos comprovantes.')
format_paragraph(paragraph_three_four, 3, 1.5748, 0,18,18,18)

#paragrafo III-IV.I
paragraph_three_four_one = document.add_paragraph("Eventuais despesas relativas a custas judiciais e extrajudiciais, como cópias, tributos, honorários periciais, bem como despesas com o eventual deslocamento e hospedagem de pessoal da Roque Khouri & Pinheiro Advogados Associados para fora de Brasília em razão da prestação de serviços serão de responsabilidade dos Interessados. Qualquer outro serviço ou indagação, incluindo também contatos informais por aplicativo de mensagem, também serão devidamente remunerados de acordo com as horas efetivamente trabalhadas.")
format_paragraph(paragraph_three_four_one, 3, 1.5748, 0,18,18,18)

#paragrafo III-IV.II
if valor_teto_exito > 0.0:
    paragraph_three_four_two = document.add_paragraph(f"Todos os valores aqui previstos serão devidamente atualizados anualmente pelo INPC ou índice que vier a substituí-lo. Todos os valores aqui previstos são devidos mesmo em caso de acordo. O valor do êxito fica limitado ao valor de R${valor_teto_exito_formatado}, devidamente atualizado.")
    format_paragraph(paragraph_three_four_two, 3, 1.5748, 0,18,18,18)
else:
    paragraph_three_four_two = document.add_paragraph("Todos os valores aqui previstos serão devidamente atualizados anualmente pelo INPC ou índice que vier a substituí-lo. Todos os valores aqui previstos são devidos mesmo em caso de acordo.")
    format_paragraph(paragraph_three_four_two, 3, 1.5748,0, 18,18,18)


#paragrafo III-V
paragraph_three_five = document.add_paragraph('Havendo necessidade de propositura de nova ação judicial que não aquela prevista no item I, deverá ser apresentado novo valor de honorários.')
format_paragraph(paragraph_three_five, 3, 1.5748, 0,18,18,18)


#############################################################################
# IV - DA CONFIDENCIALIDADE
#Tituolo
title_iv = document.add_heading('IV - DA CONFIDENCIALIDADE', level=2)
format_title_justified(title_iv)
paragraph = document.add_paragraph()
format_paragraph(paragraph, 3, 1.5748, 0,18,18,18)
paragraph.text = "O escritório e seus profissionais comprometem-se a: (i) tratar todas as informações que tiverem acesso por meio deste trabalho de forma confidencial durante o prazo de realização das atividades; e (ii) não utilizar qualquer informação confidencial para qualquer fim que não a realização dos trabalhos. Excetua-se do conceito de informação confidencial aquela que já for divulgada ou disponibilizada publicamente pelo interessado."

#paragrafo IV-I
paragraph = document.add_paragraph()
format_paragraph(paragraph, 3, 1.5748,0, 18,18,18) 
paragraph.text = "Atenciosamente,"


# Adicionar parágrafo centralizado
paragraph = document.add_paragraph()
paragraph.add_run('Roque Khouri & Pinheiro Advogados Associados \nPaulo R. Roque A. Khouri\nOAB/DF 10.671').bold = True
paragraph_format = paragraph.paragraph_format
paragraph_format.alignment = 1  # Centralizado
paragraph_format.space_before = Pt(64)

# Adicionar parágrafo para "De acordo:"
paragraph = document.add_paragraph()
paragraph.add_run("De acordo:________________").bold = True
paragraph_format = paragraph.paragraph_format
paragraph_format.space_before = Pt(40)


# Adicionar parágrafo para "Data:"
paragraph = document.add_paragraph()
paragraph.add_run("Data:_____________________").bold = True
paragraph_format = paragraph.paragraph_format
paragraph_format.space_before = Pt(32)
paragraph = document.add_paragraph()

####################################
with desenvolvimento:
    while True:
        if nome_cliente:
            break
        time.sleep(2)
    st.markdown(f"""
        <div style="text-align: right;">
            {paragraph_date.text}
        </div>
        """, unsafe_allow_html=True)
    st.markdown(title.text)
    st.write(f'**{paragraph_para.text}**')
    st.write(paragraph_ref.text)
    st.write('*texto padrao apresentação do escritorio*')
    st.write(title_one.text)
    
    while True:
        if input_consultivo_objeto:
            break
        time.sleep(2)

    if input_contencioso_objeto:
        st.markdown(f"""
        <div style="text-align: justify;">
            Conforme solicitação, apresentamos proposta de honorários para atuação consultiva, \
                referente a <b>{input_consultivo_objeto}</b>, ou, caso não tenha êxito, atuação judicial contenciosa\
                <b>{input_contencioso_objeto}</b>.
            </div>
            """, unsafe_allow_html=True)


    else:
        st.markdown(f"""
    <div style="text-align: justify;">
        Conforme solicitação, apresentamos proposta de honorários para atuação consultiva, \
                referente a <b>{input_consultivo_objeto}</b>, ou, caso não tenha êxito, atuação judicial contenciosa\
                em defesa dos interesses de <b>{nome_cliente}</b>.
        </div>
        """, unsafe_allow_html=True)
    
    st.write("")
    st.markdown(f"""
    <div style="text-align: justify;">
        {paragraph_atividades.text}
    </div>
    """, unsafe_allow_html=True)
    
    # Recuo para os itens
    for item in itens_atuacao:
        st.markdown(f"""
    <div style="text-align: justify;">
        {recuo}-  {item}
    </div>
    """, unsafe_allow_html=True)    
        
    
    st.write("")
    st.write(f'*Texto padrão sobre: disposição de equipe, alerta sobre risco jurídico e política de honorários*')
    st.write(title_three.text)
    if hora_total > 0:
        st.write(f'*Texto padrão sobre a cobrança pelo sistema de horas*')
        st.markdown(f"""
    <div style="text-align: justify;">
        {paragraph_three_one_one.text}
    </div>
    """, unsafe_allow_html=True)    
        
        st.markdown(f"{recuo}{valor_atuacao_consultivo.text}")
        st.markdown(f"{recuo}{valor_atuacao_consultivo_hora.text}")
        st.markdown(f"{recuo}{valor_consultivo_valor_aplicado.text}")
        st.markdown(f"{recuo}{valor_consultivo.text}")
        st.markdown(f"{recuo}{paragraph_desconto_consultivo.text}")
    
    if prolabore_inicial > 0.0:
        st.markdown(f"{recuo}{valor_atuacao_contencioso.text}")
        st.markdown(f"{recuo}{valor_prolabore_inicial.text}")
        st.markdown(f"{recuo}{valor_honorario_manutencao.text}")
        st.markdown(f"{recuo}{valor_honorario_exito.text}")
        st.markdown(f"{recuo}{paragraph_desconto_contencioso.text}")

    st.write('*Texto padrão sobre os eventuais custos com a contratação de advogads e despesas relativas a custas judiciais*')
    st.markdown(f"""
<div style="text-align: justify;">
    {paragraph_three_four_two.text}
</div>
""", unsafe_allow_html=True)    
            
    st.write("*Texto padrão sobre a necessidade de novo valor de honorários se propositura de nova ação judicial*")
    st.write(title_iv.text)
    st.write("")
    st.markdown(f"""
    <div style="text-align: justify;">
        <i>Texto padrão sobre confidencialidade.</i>
        </div>
        """, unsafe_allow_html=True)

        
    st.write("")
    st.write("")
    st.write("")

    # Salvar documento em arquivo temporário e permitir download
    with NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        document.save(tmp_file.name)
        st.download_button(
            label="Baixar Documento",
            data=open(tmp_file.name, 'rb').read(),
            file_name=f'proposta_consultivo_contencioso_{nome_cliente}.docx',
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

